"""
飞书机器人集成模块 v1.1
────────────────────────────────────────────────────
工作流：发消息给飞书机器人 → 自动抓取/解析内容 → AI生成要讯 → 回复卡片
支持：文章链接（自动抓取）/ 文章正文（直接生成）

环境变量配置：
  FEISHU_APP_ID            飞书应用 App ID
  FEISHU_APP_SECRET        飞书应用 App Secret
  FEISHU_VERIFY_TOKEN      飞书事件订阅 Verification Token（负载字段校验）
  FEISHU_ENCRYPT_KEY       飞书事件订阅 Encrypt Key（生产签名必填）
  FEISHU_TENANT_KEY        允许的租户 Key（生产事件身份绑定必填）
  FEISHU_EVENT_LEASE_SECONDS  后台任务租约秒数（默认 900，范围 30-3600）
  FEISHU_WEBHOOK_ALLOW_TOKEN_ONLY=1  仅本地开发兼容旧 token-only 回调
────────────────────────────────────────────────────
"""
import hmac
import json
import logging
import os
import re
import threading
import time
from concurrent.futures import ThreadPoolExecutor
from io import BytesIO
import requests
from flask import Blueprint, request, jsonify
from docx import Document as DocxDocument
from docx.shared import Pt, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from feishu_common import (
    ascii_fallback_name as _ascii_fallback_name,
    sanitize_feishu_filename as _sanitize_feishu_filename,
)
from feishu_webhook_security import (
    WebhookMisconfigured,
    WebhookRejected,
    MAX_WEBHOOK_BODY_BYTES,
    acquire_event_lease,
    decrypt_event_payload,
    submit_leased_event,
    token_only_development_enabled,
    validate_event_identity,
    verify_signed_request,
)
from protected_secrets import (
    FEISHU_CONFIG_FIELDS,
    FEISHU_SECRET_FIELDS,
    FeishuSecretStore,
    ProtectedSecretError,
    ROTATION_NOTICE,
)
from state import CONFIG_DIR

feishu_bp = Blueprint("feishu_bot", __name__)
logger = logging.getLogger(__name__)

# 全局线程池（替代无限 daemon thread，防止线程爆炸）
_worker_pool = ThreadPoolExecutor(max_workers=8, thread_name_prefix="feishu_bot")
_feishu_config_lock = threading.RLock()

# ── 飞书应用配置（运行时可通过 /api/feishu/config 接口更新）────
_FEISHU_CONFIG_FILE = os.path.join(CONFIG_DIR, ".feishu_config.json")
_FEISHU_ROTATION_REQUIRED = False
_FEISHU_ROTATION_FINGERPRINT_KEY = os.urandom(32)
_FEISHU_MIGRATED_SECRET_FINGERPRINT = None
_FEISHU_STRIPPED_SECRET_FIELDS = frozenset({
    "verify_token",
    "encrypt_key",
    "tenant_key",
})


def _feishu_secret_fingerprint(config: dict) -> bytes:
    serialized = json.dumps(
        [
            config.get(field, "").strip()
            if field in _FEISHU_STRIPPED_SECRET_FIELDS
            else config.get(field, "")
            for field in FEISHU_SECRET_FIELDS
        ],
        ensure_ascii=False,
        separators=(",", ":"),
    ).encode("utf-8")
    return hmac.digest(
        _FEISHU_ROTATION_FINGERPRINT_KEY,
        serialized,
        "sha256",
    )


def _matches_migrated_feishu_secrets(config: dict) -> bool:
    if _FEISHU_MIGRATED_SECRET_FINGERPRINT is None:
        return False
    return hmac.compare_digest(
        _feishu_secret_fingerprint(config),
        _FEISHU_MIGRATED_SECRET_FINGERPRINT,
    )


def _environment_feishu_config() -> dict:
    return {
        "app_id":       os.environ.get("FEISHU_APP_ID", ""),
        "app_secret":   os.environ.get("FEISHU_APP_SECRET", ""),
        "verify_token": os.environ.get("FEISHU_VERIFY_TOKEN", ""),
        "encrypt_key":  os.environ.get("FEISHU_ENCRYPT_KEY", ""),
        "tenant_key":   os.environ.get("FEISHU_TENANT_KEY", ""),
    }


def _environment_manages_feishu_secrets() -> bool:
    environment = _environment_feishu_config()
    return any(environment.get(field) for field in FEISHU_SECRET_FIELDS)


def _load_feishu_config(*, protector=None) -> dict:
    global _FEISHU_MIGRATED_SECRET_FINGERPRINT, _FEISHU_ROTATION_REQUIRED
    environment = _environment_feishu_config()
    loaded = FeishuSecretStore(
        _FEISHU_CONFIG_FILE,
        protector=protector,
    ).load()
    base = {field: "" for field in FEISHU_CONFIG_FIELDS}
    if loaded is not None:
        base.update(loaded.values)
        _FEISHU_ROTATION_REQUIRED = loaded.rotation_required
        if loaded.rotation_required:
            _FEISHU_MIGRATED_SECRET_FINGERPRINT = _feishu_secret_fingerprint(
                loaded.values,
            )
            logger.warning("feishu_bot: %s", ROTATION_NOTICE)
        else:
            _FEISHU_MIGRATED_SECRET_FINGERPRINT = None
    else:
        _FEISHU_ROTATION_REQUIRED = False
        _FEISHU_MIGRATED_SECRET_FINGERPRINT = None
    # Environment variables are the only supported secret source on non-Windows
    # and intentionally take precedence over local protected state everywhere.
    base.update({key: value for key, value in environment.items() if value})
    return base


def _save_feishu_config(
    config=None,
    *,
    protector=None,
    rotation_required=None,
):
    if config is None:
        config = FEISHU_CONFIG
    if rotation_required is None:
        rotation_required = _FEISHU_ROTATION_REQUIRED
    FeishuSecretStore(
        _FEISHU_CONFIG_FILE,
        protector=protector,
    ).save(config, rotation_required=rotation_required)

FEISHU_CONFIG = _load_feishu_config()

FEISHU_API = "https://open.feishu.cn/open-apis"

# ── Token 缓存 ────────────────────────────────────────────────
_token_lock  = threading.Lock()
_token_cache = {"token": "", "expire": 0}

# ════════════════════════════════════════════════════════════
# 飞书 API 工具函数
# ════════════════════════════════════════════════════════════

def get_tenant_token() -> str:
    """获取 tenant_access_token，带缓存、锁和重试。
    失败时清空缓存防止中毒，指数退避重试 3 次。"""
    with _token_lock:
        if time.time() < _token_cache["expire"] - 120:
            return _token_cache["token"]
        last_err = None
        for attempt in range(3):
            try:
                resp = requests.post(
                    f"{FEISHU_API}/auth/v3/tenant_access_token/internal",
                    json={"app_id": FEISHU_CONFIG["app_id"],
                          "app_secret": FEISHU_CONFIG["app_secret"]},
                    timeout=10,
                )
                data = resp.json()
                if data.get("code") != 0:
                    raise RuntimeError(f"飞书Token获取失败: {data.get('msg')}")
                _token_cache["token"] = data["tenant_access_token"]
                _token_cache["expire"] = time.time() + data.get("expire", 7200)
                return _token_cache["token"]
            except Exception as e:
                last_err = e
                _token_cache["token"] = ""
                _token_cache["expire"] = 0
                logger.warning(
                    "tenant_token 获取失败 attempt=%d error_type=%s",
                    attempt + 1,
                    type(e).__name__,
                )
                if attempt < 2:
                    time.sleep(1 * (attempt + 1))
        raise RuntimeError(f"飞书Token获取失败（3次重试后放弃）: {last_err}")


def _post(path: str, payload: dict) -> dict:
    """通用飞书 API POST，带返回值校验"""
    token = get_tenant_token()
    r = requests.post(
        f"{FEISHU_API}{path}",
        headers={"Authorization": f"Bearer {token}",
                 "Content-Type": "application/json"},
        json=payload,
        timeout=20,
    )
    data = r.json()
    code = data.get("code", -1)
    if code != 0:
        logger.warning("飞书 API %s 返回错误 code=%s", path, code)
    return data


def send_text(chat_id: str, text: str) -> bool:
    """向群/会话发送纯文本消息，返回是否成功"""
    try:
        data = _post("/im/v1/messages?receive_id_type=chat_id", {
            "receive_id": chat_id,
            "msg_type":   "text",
            "content":    json.dumps({"text": text}),
        })
        return data.get("code") == 0
    except Exception as e:
        logger.error("send_text 失败 error_type=%s", type(e).__name__)
        return False


def send_card(chat_id: str, card: dict) -> bool:
    """向群/会话发送卡片消息，返回是否成功"""
    try:
        data = _post("/im/v1/messages?receive_id_type=chat_id", {
            "receive_id": chat_id,
            "msg_type":   "interactive",
            "content":    json.dumps(card),
        })
        return data.get("code") == 0
    except Exception as e:
        logger.error("send_card 失败 error_type=%s", type(e).__name__)
        return False


# ════════════════════════════════════════════════════════════
# 卡片构建
# ════════════════════════════════════════════════════════════

def _build_brief_card(brief_text: str, source_info: dict) -> dict:
    """将要讯正文 + 来源信息组装成飞书互动卡片"""
    title  = (source_info.get("title") or "防务要讯")[:60]
    source = source_info.get("source", "")
    url    = source_info.get("url", "")
    now    = time.localtime()
    ts     = f"{now.tm_mon:02d}月{now.tm_mday:02d}日 {now.tm_hour:02d}:{now.tm_min:02d}"

    # 飞书卡片单元素内容限 5000 字符，分段展示
    preview = brief_text[:1500] + ("…" if len(brief_text) > 1500 else "")

    elements = [
        {"tag": "div", "text": {"tag": "lark_md",
            "content": f"**📰 素材标题**　{title}"}},
        {"tag": "div", "text": {"tag": "lark_md",
            "content": f"**📍 信源**　{source or '用户导入'}　　**🕐**　{ts}"}},
        {"tag": "hr"},
        {"tag": "div", "text": {"tag": "lark_md",
            "content": f"**📋 要讯全文**\n\n{preview}"}},
    ]

    if url:
        elements.append({
            "tag": "action",
            "actions": [{
                "tag":  "button",
                "text": {"tag": "plain_text", "content": "🔗 查看原文"},
                "type": "default",
                "url":  url,
            }]
        })

    return {
        "config": {"wide_screen_mode": True},
        "header": {
            "title":    {"tag": "plain_text", "content": "🛡️ 防务要讯已生成"},
            "template": "blue",
        },
        "elements": elements,
    }


def _build_error_card(msg: str) -> dict:
    """错误提示卡片"""
    return {
        "config": {"wide_screen_mode": True},
        "header": {
            "title":    {"tag": "plain_text", "content": "❌ 生成失败"},
            "template": "red",
        },
        "elements": [
            {"tag": "div", "text": {"tag": "lark_md", "content": msg}},
            {"tag": "div", "text": {"tag": "lark_md",
                "content": "**提示**：检查链接是否可访问，或AI Key是否已在系统中配置"}},
        ],
    }


# ════════════════════════════════════════════════════════════
# DOCX 要讯排版（严格按模板1格式）
# ════════════════════════════════════════════════════════════

def _set_font(run, font_name: str, size_pt: float = 16, bold: bool = True):
    """设置 run 的中西文字体"""
    run.font.name = font_name
    run.font.size = Pt(size_pt)
    run.font.bold = bold
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.get_or_add_rFonts()
    rFonts.set(qn('w:eastAsia'), font_name)


def _add_para(doc, align=None, left_indent=None, first_indent=None):
    """添加段落并设置统一行距 + 缩进"""
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    pf.line_spacing = Emu(363220)    # 固定行距 28.6pt
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    if align is not None:
        p.alignment = align
    if left_indent is not None:
        pf.left_indent = Emu(left_indent)
    if first_indent is not None:
        pf.first_line_indent = Emu(first_indent)
    return p


def _parse_brief_sections(text: str) -> dict:
    """解析已通过门禁的六部分要讯；缺项时失败关闭。"""
    lines = text.strip().split('\n')
    sec = {'event_time': '', 'value_point': '', 'title': '',
           'body': '', 'source': '', 'reporter': ''}

    remaining = []
    for line in lines:
        s = line.strip()
        if not s:
            continue
        if s.startswith('事件时间'):
            sec['event_time'] = s
        elif '价' in s and '值' in s and '点' in s and '：' in s:
            sec['value_point'] = s
        elif s.startswith('（信息来源') or s.startswith('(信息来源'):
            sec['source'] = s
        elif s.startswith('报送人'):
            sec['reporter'] = s
        else:
            remaining.append(s)

    non_empty = [l for l in remaining if l]
    if non_empty:
        longest_idx = max(range(len(non_empty)), key=lambda i: len(non_empty[i]))
        sec['body'] = non_empty[longest_idx]
        title_parts = [non_empty[i] for i in range(len(non_empty)) if i != longest_idx]
        sec['title'] = ''.join(title_parts)

    missing = [name for name, value in sec.items() if not value]
    if missing:
        raise ValueError("要讯结构不完整，不能生成DOCX：" + "、".join(missing))
    return sec


def _generate_brief_docx(brief_text: str) -> bytes:
    """将要讯纯文本生成为严格排版的 .docx（按模板1格式）"""
    sec = _parse_brief_sections(brief_text)
    doc = DocxDocument()

    # ── 页面设置：A4 ──
    section = doc.sections[0]
    section.page_width  = Emu(7560310)     # A4 宽
    section.page_height = Emu(10692130)    # A4 高
    section.top_margin    = Emu(914400)    # 上 2.54cm
    section.bottom_margin = Emu(914400)    # 下 2.54cm
    section.left_margin   = Emu(1143000)   # 左 3.18cm
    section.right_margin  = Emu(1143000)   # 右 3.18cm

    # ── P0: 事件时间 ──
    # "事件时间：" 黑体 + 日期 楷体_GB2312
    p = _add_para(doc)
    et = sec['event_time']
    if '：' in et:
        label, date_val = et.split('：', 1)
        r1 = p.add_run(label + '：')
        _set_font(r1, '黑体')
        r2 = p.add_run(date_val)
        _set_font(r2, '楷体_GB2312')
    else:
        r = p.add_run(et)
        _set_font(r, '黑体')

    # ── P1: 价值点 ──
    # "价 值 点：" 黑体 + 内容 楷体_GB2312，悬挂缩进
    p = _add_para(doc, left_indent=1019810, first_indent=-1019810)
    vp = sec['value_point']
    if '：' in vp:
        label, content = vp.split('：', 1)
        r1 = p.add_run(label + '：')
        _set_font(r1, '黑体')
        r2 = p.add_run(content)
        _set_font(r2, '楷体_GB2312')
    else:
        r = p.add_run(vp)
        _set_font(r, '黑体')

    # ── P2: 空行（标题前） ──
    _add_para(doc, align=WD_ALIGN_PARAGRAPH.CENTER,
              left_indent=1606550, first_indent=-1612900)

    # ── P3: 标题 ──
    # 方正小标宋简体 22pt 居中
    p = _add_para(doc, align=WD_ALIGN_PARAGRAPH.CENTER,
                  left_indent=1606550, first_indent=-1612900)
    r = p.add_run(sec['title'])
    _set_font(r, '方正小标宋简体', 22)

    # ── P4: 空行（标题后） ──
    p = _add_para(doc, align=WD_ALIGN_PARAGRAPH.CENTER,
                  left_indent=1606550, first_indent=-1612900)
    r = p.add_run('      ')
    _set_font(r, '方正小标宋简体', 22)

    # ── P5: 正文 ──
    # 仿宋_GB2312 16pt 首行缩进2字符
    p = _add_para(doc, first_indent=408305)
    r = p.add_run(sec['body'])
    _set_font(r, '仿宋_GB2312')

    # ── P6: 信息来源 ──
    if sec['source']:
        p = _add_para(doc, first_indent=408305)
        r = p.add_run(sec['source'])
        _set_font(r, '楷体_GB2312')

    # ── P7: 报送人 ──
    p = _add_para(doc, align=WD_ALIGN_PARAGRAPH.CENTER)
    r = p.add_run(sec['reporter'])
    _set_font(r, '楷体_GB2312')

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ════════════════════════════════════════════════════════════
# 飞书文件上传 + 发送
# ════════════════════════════════════════════════════════════

def upload_file(file_name: str, file_data: bytes, file_type: str = "stream") -> str:
    """上传文件到飞书，返回 file_key。失败时用 ASCII fallback 重试。"""
    token = get_tenant_token()
    cleaned_name = _sanitize_feishu_filename(file_name)
    lower = cleaned_name.lower()
    ext_map = {
        ".docx": ("doc", "application/vnd.openxmlformats-officedocument.wordprocessingml.document"),
        ".doc":  ("doc", "application/msword"),
        ".xlsx": ("xls", "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"),
        ".xls":  ("xls", "application/vnd.ms-excel"),
        ".pptx": ("ppt", "application/vnd.openxmlformats-officedocument.presentationml.presentation"),
        ".ppt":  ("ppt", "application/vnd.ms-powerpoint"),
        ".pdf":  ("pdf", "application/pdf"),
        ".txt":  ("stream", "text/plain"),
        ".md":   ("stream", "text/markdown"),
    }
    mime = "application/octet-stream"
    matched_ext = ".bin"
    for ext, (ft, mt) in ext_map.items():
        if lower.endswith(ext):
            file_type = ft
            mime = mt
            matched_ext = ext
            break
    ascii_placeholder = f"upload{matched_ext}"

    def _do_upload(send_name: str) -> dict:
        resp = requests.post(
            f"{FEISHU_API}/im/v1/files",
            headers={"Authorization": f"Bearer {token}"},
            data={"file_type": file_type, "file_name": send_name},
            files={"file": (ascii_placeholder, file_data, mime)},
            timeout=30,
        )
        return resp.json()

    data = _do_upload(cleaned_name)
    last_err = data.get("msg", "")
    if data.get("code") == 0:
        return data["data"]["file_key"]

    logger.warning(
        "飞书上传第一次失败 error_type=%s，尝试 ASCII fallback",
        type(last_err).__name__,
    )
    fallback_name = _ascii_fallback_name(matched_ext, original=file_name)
    data = _do_upload(fallback_name)
    if data.get("code") == 0:
        logger.info("飞书上传 ASCII fallback 成功: %s", fallback_name)
        return data["data"]["file_key"]

    raise RuntimeError(
        f"飞书文件上传失败: {data.get('msg', last_err)} "
        f"(file_type={file_type}, size={len(file_data)}B, "
        f"tried_name={cleaned_name!r}, fallback={fallback_name!r})"
    )


def send_file(chat_id: str, file_key: str):
    """向群/会话发送文件消息"""
    _post("/im/v1/messages?receive_id_type=chat_id", {
        "receive_id": chat_id,
        "msg_type":   "file",
        "content":    json.dumps({"file_key": file_key}),
    })


# ════════════════════════════════════════════════════════════
# 消息处理逻辑
# ════════════════════════════════════════════════════════════

def _extract_url(text: str) -> str | None:
    """从文本中提取第一个 HTTP URL"""
    m = re.search(r'https?://[^\s<>"{}|\\^`\[\]]+', text)
    return m.group(0) if m else None


def _process_async(chat_id: str, text: str):
    """在子线程中处理消息，生成并回复要讯（避免飞书 3s 超时）"""
    try:
        from app import (
            _extract_url_content,
            _build_brief_user_prompt_imported,
            _call_ai,
            _validate_brief_text,
            _brief_source_context,
            _brief_parse_date_value,
            SYSTEM_PROMPT_BRIEF_WRITE,
            AI_CONFIG,
        )

        if not AI_CONFIG.get("api_key"):
            send_text(chat_id,
                "❌ AI API Key 未配置。\n"
                "请先在电脑端追踪系统的「🤖 AI分析 → ⚙️ 配置AI」中设置 Key。")
            return

        url = _extract_url(text)

        if url:
            # ── URL 模式 ──────────────────────────────────
            # SSRF 防护：禁止访问私有/本地地址
            from app import _is_ssrf_safe
            safe, reason = _is_ssrf_safe(url)
            if not safe:
                send_text(chat_id, f"❌ URL不安全，拒绝访问：{reason}")
                return
            send_text(chat_id, f"⏳ 正在抓取链接，请稍候（约10-30秒）…")
            extracted = _extract_url_content(url)
            body = extracted.get("body", "")
            if not body or len(body) < 50:
                send_card(chat_id, _build_error_card(
                    f"无法提取页面正文（字符数：{len(body)}）\n链接：{url}"))
                return
            if not _brief_parse_date_value(extracted.get("pub_date")):
                send_text(chat_id, "❌ 页面未提取到可核实的发文日期，无法生成完整信息来源行。")
                return
            source_info = {
                "title":  extracted.get("title", url[:60]),
                "source": extracted.get("source", ""),
                "url":    url,
            }
            source_context = _brief_source_context(
                material_text="\n".join(filter(None, [extracted.get("title"), body])),
                source_name=extracted.get("source", ""),
                source_title=extracted.get("title", ""),
                publication_date=extracted.get("pub_date", ""),
                publication_date_verified=True,
                url=url,
            )
            prompt = _build_brief_user_prompt_imported(
                title=extracted.get("title", ""),
                body=body,
                source=extracted.get("source", ""),
                url=url,
                pub_date=extracted.get("pub_date", ""),
            )

        elif len(text) >= 30:
            # ── 纯文本模式 ────────────────────────────────
            send_text(chat_id, "⏳ 正在根据文本内容生成要讯，请稍候…")
            source_info = {"title": text[:40], "source": "用户导入", "url": ""}
            source_context = _brief_source_context(material_text=text)
            prompt = _build_brief_user_prompt_imported(
                title=text[:40], body=text,
                source="", url="", pub_date="")

        else:
            send_text(chat_id,
                "💡 请发送以下任一内容：\n"
                "① 文章链接（自动抓取正文）\n"
                "② 文章正文（30字以上）\n\n"
                "发送「帮助」查看使用说明")
            return

        # ── 调用 AI ───────────────────────────────────────
        messages = [
            {"role": "system", "content": SYSTEM_PROMPT_BRIEF_WRITE},
            {"role": "user",   "content": prompt},
        ]
        result = _call_ai(messages, temperature=0.4)
        validation = _validate_brief_text(result, source_context=source_context)
        if validation.get("valid") is not True:
            errors = validation.get("errors") or ["未知校验错误"]
            send_card(chat_id, _build_error_card(
                "要讯未通过写作规范，已停止发送和导出：\n" +
                "\n".join(f"• {error}" for error in errors[:8])
            ))
            return

        # ── 回复卡片预览 ─────────────────────────────────
        send_card(chat_id, _build_brief_card(result, source_info))

        # ── 生成 DOCX 并发送文件 ─────────────────────────
        try:
            docx_bytes = _generate_brief_docx(result)
            title_short = (source_info.get("title") or "防务要讯")[:30]
            ts = time.strftime("%Y%m%d_%H%M")
            file_name = f"要讯_{title_short}_{ts}.docx"
            file_key = upload_file(file_name, docx_bytes)
            send_file(chat_id, file_key)
        except Exception as docx_err:
            logger.error("DOCX generation or upload failed (%s)", type(docx_err).__name__)
            send_text(chat_id, "⚠️ 要讯文本已生成，但DOCX文件发送失败")

    except Exception as exc:
        logger.error("feishu_bot processing failed (%s)", type(exc).__name__)
        send_card(chat_id, _build_error_card("处理失败"))


# ════════════════════════════════════════════════════════════
# Flask 路由
# ════════════════════════════════════════════════════════════

HELP_TEXT = (
    "🛡️ **防务要讯机器人** 使用说明\n\n"
    "📌 **支持的消息类型**\n"
    "• 发送文章链接 → 自动抓取正文并生成要讯\n"
    "• 发送文章正文（30字+）→ 直接生成要讯\n\n"
    "📝 **示例**\n"
    "```\nhttps://www.scmp.com/news/military/...\n```\n\n"
    "⏱️ 生成约需 10-30 秒，生成中会先收到「⏳」提示。"
)


@feishu_bp.route("/api/feishu/webhook", methods=["POST"])
def feishu_webhook():
    """飞书事件订阅回调入口（需在飞书开放平台注册为事件订阅URL）"""
    if request.content_length is not None and request.content_length > MAX_WEBHOOK_BODY_BYTES:
        return jsonify({"code": 1, "msg": "webhook body too large"}), 413
    raw_body = request.get_data()
    if len(raw_body) > MAX_WEBHOOK_BODY_BYTES:
        return jsonify({"code": 1, "msg": "webhook body too large"}), 413
    data = request.get_json(silent=True) or {}
    if not isinstance(data, dict):
        return jsonify({"code": 1, "msg": "invalid webhook payload"}), 400
    verify_token = str(FEISHU_CONFIG.get("verify_token") or "").strip()
    if not verify_token:
        logger.error("feishu_bot: webhook verification token is not configured")
        return jsonify({"code": 1, "msg": "webhook verification not configured"}), 503

    allow_legacy = token_only_development_enabled()
    signing_key = str(FEISHU_CONFIG.get("encrypt_key") or "").strip()
    try:
        verify_signed_request(
            request.headers,
            raw_body,
            signing_key=signing_key,
            allow_token_only=allow_legacy,
        )
        data = decrypt_event_payload(data, encrypt_key=signing_key)
    except WebhookRejected as exc:
        logger.warning("feishu_bot: webhook signature rejected (%s)", exc.code)
        return jsonify({"code": 1, "msg": "invalid webhook signature"}), 403
    except WebhookMisconfigured as exc:
        logger.error("feishu_bot: webhook security unavailable (%s)", exc.code)
        return jsonify({"code": 1, "msg": "webhook security not configured"}), 503

    # ── URL 验证（首次注册 webhook 时飞书发来）──────────────
    if data.get("type") == "url_verification":
        incoming_token = str(data.get("token") or "")
        if not hmac.compare_digest(verify_token, incoming_token):
            logger.warning("feishu_bot: URL verification token invalid")
            return jsonify({"code": 1, "msg": "invalid token"}), 403
        challenge = data.get("challenge", "")
        if not isinstance(challenge, str) or len(challenge) > 1024:
            return jsonify({"code": 1, "msg": "invalid challenge"}), 400
        logger.info("feishu_bot: URL verification OK")
        return jsonify({"challenge": challenge})

    # ── 事件处理（schema 2.0）────────────────────────────────
    header = data.get("header")
    if not isinstance(header, dict):
        return jsonify({"code": 1, "msg": "invalid webhook event"}), 400
    incoming_token = str(header.get("token") or "")
    if not hmac.compare_digest(verify_token, incoming_token):
        logger.warning("feishu_bot: Verification Token invalid")
        return jsonify({"code": 1, "msg": "invalid token"}), 403

    try:
        validate_event_identity(
            data,
            expected_app_id=str(FEISHU_CONFIG.get("app_id") or "").strip(),
            expected_tenant_key=str(FEISHU_CONFIG.get("tenant_key") or "").strip(),
            allow_legacy=allow_legacy,
        )
    except WebhookRejected as exc:
        logger.warning("feishu_bot: webhook identity rejected (%s)", exc.code)
        return jsonify({"code": 1, "msg": "invalid webhook identity"}), 403
    except WebhookMisconfigured as exc:
        logger.error("feishu_bot: webhook identity unavailable (%s)", exc.code)
        return jsonify({"code": 1, "msg": "webhook identity not configured"}), 503

    event_type = header.get("event_type", "")
    event = data.get("event")
    if not isinstance(event, dict):
        return jsonify({"code": 1, "msg": "invalid webhook event"}), 400

    if event_type != "im.message.receive_v1":
        return jsonify({"code": 0})

    message = event.get("message")
    if not isinstance(message, dict):
        return jsonify({"code": 1, "msg": "invalid webhook event"}), 400
    msg_type = message.get("message_type", "")
    chat_id  = message.get("chat_id", "")

    if not chat_id:
        return jsonify({"code": 0})

    # ── 持久租约：仅在线程完成后确认，崩溃后允许重试接管 ────────
    try:
        event_lease = acquire_event_lease(data, allow_legacy=allow_legacy)
    except WebhookRejected as exc:
        logger.warning("feishu_bot: webhook event rejected (%s)", exc.code)
        return jsonify({"code": 1, "msg": "invalid webhook event"}), 403
    except WebhookMisconfigured as exc:
        logger.error("feishu_bot: webhook deduplication unavailable (%s)", exc.code)
        return jsonify({"code": 1, "msg": "webhook deduplication unavailable"}), 503
    if event_lease is None:
        logger.info("feishu_bot: in-flight or completed event skipped")
        return jsonify({"code": 0})

    # 只处理文本消息
    if msg_type != "text":
        send_text(chat_id,
            "💡 目前支持文字消息（链接 或 正文）。\n发送「帮助」查看使用说明。")
        event_lease.complete()
        return jsonify({"code": 0})

    try:
        text = json.loads(message.get("content", "{}")).get("text", "").strip()
    except Exception:
        event_lease.complete()
        return jsonify({"code": 0})

    if not text:
        event_lease.complete()
        return jsonify({"code": 0})

    # 帮助指令
    if text in ["帮助", "help", "?", "？", "/help"]:
        send_text(chat_id, HELP_TEXT)
        event_lease.complete()
        return jsonify({"code": 0})

    # 异步线程池生成（立即返回 200，避免飞书重试）
    try:
        submit_leased_event(_worker_pool, event_lease, _process_async, chat_id, text)
    except WebhookMisconfigured as exc:
        logger.error("feishu_bot: webhook dispatch unavailable (%s)", exc.code)
        return jsonify({"code": 1, "msg": "webhook dispatch unavailable"}), 503

    return jsonify({"code": 0})


@feishu_bp.route("/api/feishu/config", methods=["GET", "POST"])
def feishu_config():
    """读写飞书机器人配置（供前端设置页调用，需登录）"""
    global FEISHU_CONFIG, _FEISHU_MIGRATED_SECRET_FINGERPRINT
    global _FEISHU_ROTATION_REQUIRED
    # Blueprint 为避免循环 import 在请求期复用主应用的统一鉴权；浏览器
    # cookie 是短期进程内 session，长期 master/device token 不进入 cookie。
    from app import _workspace_auth_error_response
    rejection = _workspace_auth_error_response()
    if rejection is not None:
        return rejection
    if request.method == "POST":
        data = request.get_json() or {}
        if not isinstance(data, dict):
            return jsonify({"error": "飞书配置格式无效"}), 400
        for field in FEISHU_CONFIG_FIELDS:
            value = data.get(field)
            if value:
                if not isinstance(value, str):
                    return jsonify({"error": "飞书配置字段必须是字符串"}), 400
        if (
            "old_credentials_revoked" in data
            and not isinstance(data["old_credentials_revoked"], bool)
        ):
            return jsonify({
                "error": "旧凭据撤销确认必须是布尔值",
                "code": "FEISHU_REVOCATION_CONFIRMATION_INVALID",
            }), 400
        with _feishu_config_lock:
            if (
                _FEISHU_ROTATION_REQUIRED
                and _environment_manages_feishu_secrets()
            ):
                return jsonify({
                    "error": (
                        "飞书秘密凭据由环境变量托管，本应用不会写入本地或清除"
                        "轮换状态；请在飞书开发者后台撤销旧值，更新部署环境中的"
                        "完整凭据并重启。环境覆盖存在期间，本地迁移提醒不会自动"
                        "清除；如需清除，请先移除秘密凭据环境覆盖并重启，再通过"
                        "本接口完整提交新凭据和撤销确认"
                    ),
                    "code": "FEISHU_ROTATION_ENVIRONMENT_MANAGED",
                }), 409
            candidate = dict(FEISHU_CONFIG)
            candidate.update({
                field: data[field]
                for field in FEISHU_CONFIG_FIELDS
                if data.get(field)
            })
            rotation_required = _FEISHU_ROTATION_REQUIRED
            submitted_secret_fields = {
                field for field in FEISHU_SECRET_FIELDS if data.get(field)
            }
            if _FEISHU_ROTATION_REQUIRED:
                if submitted_secret_fields != set(FEISHU_SECRET_FIELDS):
                    return jsonify({
                        "error": "轮换时必须一次提交完整的新飞书凭据",
                        "code": "FEISHU_ROTATION_INCOMPLETE",
                    }), 409
                if data.get("old_credentials_revoked") is not True:
                    return jsonify({
                        "error": (
                            "请确认已在飞书开发者后台撤销旧凭据；"
                            "本应用无法自动核验远程撤销状态"
                        ),
                        "code": "FEISHU_REVOCATION_CONFIRMATION_REQUIRED",
                    }), 409
                if _matches_migrated_feishu_secrets(candidate):
                    return jsonify({
                        "error": "新凭据不能与迁移前的旧凭据完全相同",
                        "code": "FEISHU_CREDENTIALS_UNCHANGED",
                    }), 409
                if any(
                    data[field] != data[field].strip()
                    for field in FEISHU_SECRET_FIELDS
                ):
                    return jsonify({
                        "error": "飞书秘密凭据不得为空或包含首尾空白",
                        "code": "FEISHU_ROTATION_SECRET_WHITESPACE",
                    }), 400
                rotation_required = False
            try:
                _save_feishu_config(
                    candidate,
                    rotation_required=rotation_required,
                )
            except ProtectedSecretError as exc:
                logger.error(
                    "feishu_bot: protected configuration save failed (%s)",
                    exc.code,
                )
                status = 413 if exc.code == "CONFIG_TOO_LARGE" else 503
                return jsonify({
                    "error": (
                        "飞书配置超过安全存储大小限制"
                        if status == 413
                        else "当前系统无法安全保存飞书凭据，请改用环境变量配置"
                    ),
                    "code": (
                        "FEISHU_CONFIG_TOO_LARGE"
                        if status == 413
                        else "FEISHU_SECRET_STORAGE_UNAVAILABLE"
                    ),
                }), status
            # Assignment is atomic for readers.  The token lock orders the
            # configuration switch with cache invalidation, while the outer
            # lock linearizes concurrent configuration POST requests.
            with _token_lock:
                FEISHU_CONFIG = candidate
                _FEISHU_ROTATION_REQUIRED = rotation_required
                if not rotation_required:
                    _FEISHU_MIGRATED_SECRET_FINGERPRINT = None
                _token_cache["expire"] = 0
        return jsonify({"ok": True,
                        "app_id": FEISHU_CONFIG["app_id"],
                        "configured": bool(FEISHU_CONFIG["app_id"]),
                        "signature_configured": bool(FEISHU_CONFIG.get("encrypt_key")),
                        "tenant_configured": bool(FEISHU_CONFIG.get("tenant_key")),
                        "credential_rotation_required": _FEISHU_ROTATION_REQUIRED,
                        "credential_notice": ROTATION_NOTICE if _FEISHU_ROTATION_REQUIRED else ""})
    return jsonify({
        "app_id":     FEISHU_CONFIG["app_id"],
        "configured": bool(FEISHU_CONFIG["app_id"] and FEISHU_CONFIG["app_secret"]),
        "signature_configured": bool(FEISHU_CONFIG.get("encrypt_key")),
        "tenant_configured": bool(FEISHU_CONFIG.get("tenant_key")),
        "credential_rotation_required": _FEISHU_ROTATION_REQUIRED,
        "credential_notice": ROTATION_NOTICE if _FEISHU_ROTATION_REQUIRED else "",
        "webhook_url": "/api/feishu/webhook  （需配合 ngrok 或公网域名使用）",
    })


@feishu_bp.route("/api/feishu/test", methods=["POST"])
def feishu_test():
    """向指定 chat_id 发送一条测试消息，验证机器人配置是否正确（需登录）"""
    from app import _workspace_auth_error_response
    rejection = _workspace_auth_error_response()
    if rejection is not None:
        return rejection
    data = request.get_json() or {}
    chat_id = data.get("chat_id", "")
    if not chat_id:
        return jsonify({"error": "需要 chat_id"}), 400
    try:
        send_text(chat_id, "✅ 防务要讯机器人连接正常！发送文章链接即可生成要讯。")
        return jsonify({"ok": True})
    except Exception as exc:
        logger.warning("feishu connection test failed (%s)", type(exc).__name__)
        return jsonify({"error": "飞书连接测试失败", "code": "FEISHU_UPSTREAM_FAILED"}), 502
