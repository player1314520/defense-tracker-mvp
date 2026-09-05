"""Hard-timeout PDF/DOCX text parsing in a disposable child process."""

from __future__ import annotations

import json
import os
import re
import stat
import subprocess
import sys
import tempfile
import threading
import ctypes
import time
from io import BytesIO
from pathlib import Path

import document_safety
from upload_safety import UploadValidationError, validate_upload


DOCUMENT_WORKER_MEMORY_LIMIT_BYTES = 512 * 1024 * 1024
DOCUMENT_PARSE_MAX_WORKERS = 2
DOCUMENT_PARSE_RETRY_AFTER_SECONDS = 5
_DOCUMENT_PARSE_SLOTS = threading.BoundedSemaphore(DOCUMENT_PARSE_MAX_WORKERS)
JOB_OBJECT_LIMIT_PROCESS_MEMORY = 0x00000100
JOB_OBJECT_LIMIT_KILL_ON_JOB_CLOSE = 0x00002000
JOB_OBJECT_EXTENDED_LIMIT_INFORMATION_CLASS = 9


class _IoCounters(ctypes.Structure):
    _fields_ = [
        ("ReadOperationCount", ctypes.c_uint64),
        ("WriteOperationCount", ctypes.c_uint64),
        ("OtherOperationCount", ctypes.c_uint64),
        ("ReadTransferCount", ctypes.c_uint64),
        ("WriteTransferCount", ctypes.c_uint64),
        ("OtherTransferCount", ctypes.c_uint64),
    ]


class _JobObjectBasicLimitInformation(ctypes.Structure):
    _fields_ = [
        ("PerProcessUserTimeLimit", ctypes.c_int64),
        ("PerJobUserTimeLimit", ctypes.c_int64),
        ("LimitFlags", ctypes.c_uint32),
        ("MinimumWorkingSetSize", ctypes.c_size_t),
        ("MaximumWorkingSetSize", ctypes.c_size_t),
        ("ActiveProcessLimit", ctypes.c_uint32),
        ("Affinity", ctypes.c_size_t),
        ("PriorityClass", ctypes.c_uint32),
        ("SchedulingClass", ctypes.c_uint32),
    ]


class _JobObjectExtendedLimitInformation(ctypes.Structure):
    _fields_ = [
        ("BasicLimitInformation", _JobObjectBasicLimitInformation),
        ("IoInfo", _IoCounters),
        ("ProcessMemoryLimit", ctypes.c_size_t),
        ("JobMemoryLimit", ctypes.c_size_t),
        ("PeakProcessMemoryUsed", ctypes.c_size_t),
        ("PeakJobMemoryUsed", ctypes.c_size_t),
    ]


class _WindowsJobApi:
    """Small ctypes boundary kept injectable for non-Windows unit tests."""

    def __init__(self) -> None:
        kernel32 = ctypes.WinDLL("kernel32", use_last_error=True)
        kernel32.CreateJobObjectW.argtypes = [ctypes.c_void_p, ctypes.c_wchar_p]
        kernel32.CreateJobObjectW.restype = ctypes.c_void_p
        kernel32.SetInformationJobObject.argtypes = [
            ctypes.c_void_p,
            ctypes.c_int,
            ctypes.c_void_p,
            ctypes.c_uint32,
        ]
        kernel32.SetInformationJobObject.restype = ctypes.c_int
        kernel32.AssignProcessToJobObject.argtypes = [ctypes.c_void_p, ctypes.c_void_p]
        kernel32.AssignProcessToJobObject.restype = ctypes.c_int
        kernel32.CloseHandle.argtypes = [ctypes.c_void_p]
        kernel32.CloseHandle.restype = ctypes.c_int
        self._kernel32 = kernel32

    @staticmethod
    def _raise_last_error(action: str) -> None:
        error = ctypes.get_last_error()
        raise OSError(error, f"{action} failed")

    def create(self):
        handle = self._kernel32.CreateJobObjectW(None, None)
        if not handle:
            self._raise_last_error("CreateJobObjectW")
        return handle

    def configure(self, handle, *, limit_flags: int, process_memory_limit: int) -> None:
        limits = _JobObjectExtendedLimitInformation()
        limits.BasicLimitInformation.LimitFlags = int(limit_flags)
        limits.ProcessMemoryLimit = int(process_memory_limit)
        if not self._kernel32.SetInformationJobObject(
            handle,
            JOB_OBJECT_EXTENDED_LIMIT_INFORMATION_CLASS,
            ctypes.byref(limits),
            ctypes.sizeof(limits),
        ):
            self._raise_last_error("SetInformationJobObject")

    def assign(self, handle, process_handle: int) -> None:
        if not self._kernel32.AssignProcessToJobObject(
            handle, ctypes.c_void_p(int(process_handle))
        ):
            self._raise_last_error("AssignProcessToJobObject")

    def close(self, handle) -> None:
        if handle:
            self._kernel32.CloseHandle(handle)


class _WorkerJobHandle:
    def __init__(self, api, handle) -> None:
        self._api = api
        self._handle = handle

    def close(self) -> None:
        if self._handle is not None:
            handle, self._handle = self._handle, None
            self._api.close(handle)


def _bind_worker_job(
    process,
    *,
    platform: str | None = None,
    api=None,
    memory_limit: int = DOCUMENT_WORKER_MEMORY_LIMIT_BYTES,
):
    """Bind the direct Windows worker to a kill-on-close, memory-limited job."""

    if (os.name if platform is None else platform) != "nt":
        return None
    if int(memory_limit) <= 0:
        raise ValueError("memory_limit must be positive")
    process_handle = getattr(process, "_handle", None)
    if process_handle is None:
        raise OSError("worker process handle unavailable")
    job_api = api or _WindowsJobApi()
    handle = job_api.create()
    try:
        job_api.configure(
            handle,
            limit_flags=(
                JOB_OBJECT_LIMIT_PROCESS_MEMORY | JOB_OBJECT_LIMIT_KILL_ON_JOB_CLOSE
            ),
            process_memory_limit=int(memory_limit),
        )
        job_api.assign(handle, int(process_handle))
    except BaseException:
        job_api.close(handle)
        raise
    return _WorkerJobHandle(job_api, handle)


class IsolatedDocumentParseError(RuntimeError):
    def __init__(self, code: str, message: str, *, details=None) -> None:
        self.code = str(code or "DOCUMENT_PARSE_FAILED")
        self.message = str(message or "文档无法安全解析")
        self.details = details if isinstance(details, dict) else {}
        super().__init__(f"[{self.code}] {self.message}")


def _bounded_title(text: str, fallback: str) -> str:
    first = next((line.strip() for line in text.splitlines() if line.strip()), "")
    return (first or os.path.basename(fallback) or "文档")[:100]


def _parse_document(kind: str, content: bytes, filename: str, limits: dict) -> dict:
    max_chars = max(1, int(limits.get("max_chars") or 60_000))
    max_file_size = max(1, int(limits.get("max_file_size") or 25 * 1024 * 1024))
    parts: list[str] = []
    extracted = 0

    def append_text(value: str) -> bool:
        nonlocal extracted
        value = str(value or "").strip()
        if not value:
            return True
        remaining = max_chars - extracted
        if remaining <= 0:
            return False
        parts.append(value[:remaining])
        extracted += min(len(value), remaining) + 1
        return extracted < max_chars

    if kind == "docx":
        safe_filename = filename if Path(filename).suffix else f"{filename}.docx"
        try:
            document_safety.validate_docx(content, max_input_bytes=max_file_size)
        except document_safety.DocumentSafetyError as exc:
            raise UploadValidationError(exc.code, str(exc)) from None
        validate_upload(safe_filename, content, max_file_size=max_file_size)
        from docx import Document as DocxDocument

        document = DocxDocument(BytesIO(content))
        keep_reading = True
        for paragraph in document.paragraphs:
            keep_reading = append_text(paragraph.text)
            if not keep_reading:
                break
        if keep_reading:
            for table in document.tables:
                for row in table.rows:
                    keep_reading = append_text(
                        " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                    )
                    if not keep_reading:
                        break
                if not keep_reading:
                    break
    elif kind == "pdf":
        max_pages = max(1, int(limits.get("max_pages") or 30))
        safe_filename = filename if Path(filename).suffix else f"{filename}.pdf"
        try:
            document_safety.validate_pdf(content, max_input_bytes=max_file_size)
        except document_safety.DocumentSafetyError as exc:
            raise UploadValidationError(exc.code, str(exc)) from None
        validate_upload(
            safe_filename,
            content,
            max_file_size=max_file_size,
            max_pages=max_pages,
        )
        import pdfplumber

        with pdfplumber.open(BytesIO(content)) as document:
            for page in document.pages[:max_pages]:
                if not append_text(page.extract_text() or ""):
                    break
    else:
        raise UploadValidationError("UNSUPPORTED_UPLOAD_TYPE", "仅支持 DOCX 和 PDF 文件")

    text = re.sub(r"\n{3,}", "\n\n", "\n".join(parts)).strip()[:max_chars]
    return {"title": _bounded_title(text, filename), "text": text}


def _worker_directory_valid() -> bool:
    root = Path.cwd()
    root_stat = root.lstat()
    if (
        # Windows can report an 8.3 cwd alias for the same physical temp folder.
        not root.parent.samefile(Path(tempfile.gettempdir()))
        or not re.fullmatch(r"defensetracker-parser-[A-Za-z0-9_-]+", root.name)
        or not stat.S_ISDIR(root_stat.st_mode)
        or getattr(root_stat, "st_file_attributes", 0) & 0x400
    ):
        return False
    if os.name != "nt" and (
        root_stat.st_uid != os.getuid() or root_stat.st_mode & 0o077
    ):
        return False
    return True


def _read_worker_file(name: str, limit: int) -> bytes:
    before = os.lstat(name)
    if (
        not stat.S_ISREG(before.st_mode)
        or before.st_nlink != 1
        or getattr(before, "st_file_attributes", 0) & 0x400
    ):
        raise ValueError("parser request must be a regular private file")
    flags = os.O_RDONLY | getattr(os, "O_BINARY", 0) | getattr(os, "O_NOFOLLOW", 0)
    with os.fdopen(os.open(name, flags), "rb") as stream:
        opened = os.fstat(stream.fileno())
        if (before.st_dev, before.st_ino) != (opened.st_dev, opened.st_ino):
            raise ValueError("parser request changed while opening")
        if os.name != "nt" and opened.st_uid != os.getuid():
            raise ValueError("parser request belongs to another user")
        return stream.read(limit + 1)


def _write_result(payload: dict) -> None:
    # Exclusive creation also rejects pre-existing files, links and reparse points.
    with open("result.json", "x", encoding="utf-8") as stream:
        json.dump(payload, stream, ensure_ascii=False, separators=(",", ":"))


def worker_file_entry(meta_path: str, input_path: str, output_path: str) -> int:
    """Isolated worker entry used by source Python and the frozen launcher."""
    # HTTP filenames only occur inside metadata. The worker protocol accepts no
    # paths: its cwd is the private directory created by the parent process.
    if (meta_path, input_path, output_path) != (
        "request.json", "input.bin", "result.json"
    ) or not _worker_directory_valid() or os.path.lexists("result.json"):
        return 64
    try:
        raw_meta = _read_worker_file("request.json", 65_536)
        if len(raw_meta) > 65_536:
            raise ValueError("parser metadata too large")
        meta = json.loads(raw_meta.decode("utf-8"))
        limits = meta.get("limits") if isinstance(meta, dict) else None
        if not isinstance(limits, dict):
            raise ValueError("parser metadata invalid")
        if meta.get("wait_for_start") is True:
            gate_deadline = time.monotonic() + 30.0
            while not Path("worker.start").exists():
                if time.monotonic() >= gate_deadline:
                    raise TimeoutError("parser start gate timed out")
                time.sleep(0.01)
            if _read_worker_file("worker.start", 2) != b"go":
                raise ValueError("parser start gate invalid")
        max_file_size = max(1, int(limits.get("max_file_size") or 25 * 1024 * 1024))
        content = _read_worker_file("input.bin", max_file_size)
        if len(content) > max_file_size:
            raise UploadValidationError("UPLOAD_TOO_LARGE", "文件超过隔离解析上限")
        result = _parse_document(
            str(meta.get("kind") or "").lower(),
            content,
            os.path.basename(str(meta.get("filename") or "document")),
            limits,
        )
        _write_result({"status": "ok", "result": result})
        return 0
    except UploadValidationError as exc:
        _write_result(
            {
                "status": "validation_error",
                "code": exc.code,
                "message": exc.message,
                "details": exc.details,
            },
        )
        return 2
    except BaseException:
        _write_result(
            {
                "status": "error",
                "code": "DOCUMENT_PARSE_FAILED",
                "message": "文档无法安全解析",
            },
        )
        return 3


def _terminate_process(process) -> None:
    if process.poll() is None:
        process.terminate()
        try:
            process.wait(timeout=2)
        except subprocess.TimeoutExpired:
            process.kill()
            process.wait(timeout=2)


def _wait_for_worker(process, timeout: float) -> int:
    try:
        return int(process.wait(timeout=max(0.01, float(timeout))))
    except subprocess.TimeoutExpired as exc:
        _terminate_process(process)
        raise IsolatedDocumentParseError(
            "DOCUMENT_PARSE_TIMEOUT", "文档解析超时，已终止隔离解析进程"
        ) from exc


def _worker_command() -> list[str]:
    if getattr(sys, "frozen", False):
        return [
            sys.executable,
            "--document-parser-worker",
            "request.json",
            "input.bin",
            "result.json",
        ]
    return [
        sys.executable,
        str(Path(__file__).resolve()),
        "--worker",
        "request.json",
        "input.bin",
        "result.json",
    ]


def _parse_document_isolated_once(
    kind: str,
    content: bytes,
    filename: str,
    *,
    timeout: float = 12.0,
    max_chars: int = 60_000,
    max_pages: int = 30,
    max_file_size: int = 25 * 1024 * 1024,
) -> dict:
    with tempfile.TemporaryDirectory(prefix="defensetracker-parser-") as temp_dir:
        root = Path(temp_dir).resolve()
        meta_path = root / "request.json"
        input_path = root / "input.bin"
        output_path = root / "result.json"
        start_gate_path = root / "worker.start"
        meta_path.write_text(
            json.dumps(
                {
                    "kind": str(kind or "").lower(),
                    "filename": os.path.basename(str(filename or "document")),
                    "wait_for_start": True,
                    "limits": {
                        "max_chars": int(max_chars),
                        "max_pages": int(max_pages),
                        "max_file_size": int(max_file_size),
                    },
                },
                ensure_ascii=False,
            ),
            encoding="utf-8",
        )
        input_path.write_bytes(bytes(content))
        creation_flags = 0x08000000 if os.name == "nt" else 0
        process = None
        worker_job = None
        try:
            try:
                process = subprocess.Popen(
                    _worker_command(),
                    cwd=root,
                    stdin=subprocess.DEVNULL,
                    stdout=subprocess.DEVNULL,
                    stderr=subprocess.DEVNULL,
                    close_fds=True,
                    creationflags=creation_flags,
                )
            except (OSError, subprocess.SubprocessError) as exc:
                raise IsolatedDocumentParseError(
                    "DOCUMENT_PARSER_ISOLATION_FAILED",
                    "文档隔离解析进程无法启动",
                ) from exc
            try:
                worker_job = _bind_worker_job(process)
            except BaseException as exc:
                _terminate_process(process)
                raise IsolatedDocumentParseError(
                    "DOCUMENT_PARSER_ISOLATION_FAILED",
                    "文档隔离解析环境无法建立",
                ) from exc
            try:
                start_gate_path.write_bytes(b"go")
            except OSError as exc:
                _terminate_process(process)
                raise IsolatedDocumentParseError(
                    "DOCUMENT_PARSER_ISOLATION_FAILED",
                    "文档隔离解析环境无法启动",
                ) from exc
            _wait_for_worker(process, timeout)
            if not output_path.is_file() or output_path.stat().st_size > 1_000_000:
                raise IsolatedDocumentParseError(
                    "DOCUMENT_PARSE_FAILED", "隔离解析进程未返回有效结果"
                )
            try:
                message = json.loads(output_path.read_text(encoding="utf-8"))
            except (OSError, UnicodeError, ValueError) as exc:
                raise IsolatedDocumentParseError(
                    "DOCUMENT_PARSE_FAILED", "隔离解析结果无效"
                ) from exc
        finally:
            if process is not None and process.poll() is None:
                _terminate_process(process)
            if worker_job is not None:
                worker_job.close()

    if not isinstance(message, dict) or message.get("status") != "ok":
        raise IsolatedDocumentParseError(
            str(message.get("code") or "DOCUMENT_PARSE_FAILED")
            if isinstance(message, dict)
            else "DOCUMENT_PARSE_FAILED",
            str(message.get("message") or "文档无法安全解析")
            if isinstance(message, dict)
            else "文档无法安全解析",
            details=message.get("details") if isinstance(message, dict) else None,
        )
    result = message.get("result")
    if not isinstance(result, dict):
        raise IsolatedDocumentParseError("DOCUMENT_PARSE_FAILED", "隔离解析结果无效")
    return result


def parse_document_isolated(
    kind: str,
    content: bytes,
    filename: str,
    *,
    timeout: float = 12.0,
    max_chars: int = 60_000,
    max_pages: int = 30,
    max_file_size: int = 25 * 1024 * 1024,
) -> dict:
    slots = _DOCUMENT_PARSE_SLOTS
    if not slots.acquire(blocking=False):
        raise IsolatedDocumentParseError(
            "DOCUMENT_PARSE_QUEUE_FULL",
            "文档解析队列已满，请稍后重试",
            details={"retry_after": DOCUMENT_PARSE_RETRY_AFTER_SECONDS},
        )
    try:
        return _parse_document_isolated_once(
            kind,
            content,
            filename,
            timeout=timeout,
            max_chars=max_chars,
            max_pages=max_pages,
            max_file_size=max_file_size,
        )
    finally:
        slots.release()


def _module_main(argv: list[str]) -> int:
    if len(argv) == 4 and argv[0] == "--worker":
        return worker_file_entry(argv[1], argv[2], argv[3])
    return 64


if __name__ == "__main__":
    raise SystemExit(_module_main(sys.argv[1:]))


__all__ = [
    "DOCUMENT_PARSE_MAX_WORKERS",
    "DOCUMENT_PARSE_RETRY_AFTER_SECONDS",
    "DOCUMENT_WORKER_MEMORY_LIMIT_BYTES",
    "IsolatedDocumentParseError",
    "parse_document_isolated",
    "worker_file_entry",
]
