"""Evidence-bound V9 document, layout and immutable publication helpers."""
from __future__ import annotations

import copy
import html
import re
import uuid
from datetime import datetime, timezone
from io import BytesIO
from typing import Any


DOCUMENT_KINDS = {"report", "brief"}
DOCUMENT_STAGES = {"outline", "draft", "review", "ready"}
SOURCE_STATUSES = {
    "verified",
    "source_claim",
    "inference",
    "scenario_assumption",
}
FACT_CHECK_STATUSES = {"pending", "passed", "failed"}
BOARD_STATUSES = {
    "evidence_needed",
    "editing",
    "pending_approval",
    "signed",
    "recalled",
}
LOCAL_SIGNING_NOTICE = (
    "签发仅表示本地批准并冻结快照，"
    "不代表已向外部渠道发布"
)


class PublicationRecalled(ValueError):
    code = "PUBLICATION_RECALLED"


def _now() -> str:
    return datetime.now(timezone.utc).isoformat()


def _text(value: Any, *, limit: int = 20000) -> str:
    return str(value or "").strip()[:limit]


def _ids(value: Any) -> list[str]:
    if isinstance(value, str):
        values = re.split(r"[,，\n]+", value)
    elif isinstance(value, list):
        values = value
    else:
        values = []
    return list(dict.fromkeys(_text(item, limit=100) for item in values if _text(item, limit=100)))


def normalize_paragraph(value: dict, *, existing_id: str = "") -> dict:
    value = value if isinstance(value, dict) else {}
    source_status = _text(value.get("source_status"), limit=30) or "source_claim"
    fact_check = _text(value.get("fact_check"), limit=30) or "pending"
    if source_status not in SOURCE_STATUSES:
        raise ValueError("无效来源状态")
    if fact_check not in FACT_CHECK_STATUSES:
        raise ValueError("无效事实核查状态")
    return {
        "paragraph_id": _text(value.get("paragraph_id"), limit=100)
        or existing_id
        or str(uuid.uuid4()),
        "heading": _text(value.get("heading"), limit=300),
        "text": _text(value.get("text")),
        "evidence_ids": _ids(value.get("evidence_ids")),
        "claim_ids": _ids(value.get("claim_ids")),
        "source_status": source_status,
        "fact_check": fact_check,
        "fact_check_note": _text(value.get("fact_check_note"), limit=2000),
    }


def validate_document(content: dict) -> dict:
    errors: list[str] = []
    paragraphs = content.get("paragraphs")
    if not isinstance(paragraphs, list) or not paragraphs:
        errors.append("稿件至少需要一个段落")
        paragraphs = []
    for index, paragraph in enumerate(paragraphs, start=1):
        label = paragraph.get("heading") or f"第 {index} 段"
        if not paragraph.get("text"):
            errors.append(f"{label}：正文为空")
        if not paragraph.get("evidence_ids"):
            errors.append(f"{label}：缺少引用证据")
        if paragraph.get("fact_check") != "passed":
            errors.append(f"{label}：事实核查未通过")
        if paragraph.get("source_status") not in SOURCE_STATUSES:
            errors.append(f"{label}：来源状态无效")
    if content.get("kind") == "brief":
        errors.append(
            "V9通用稿件不承载要讯签发；请使用写作室中的要讯专用生成与来源校验流程"
        )
    return {
        "ready": not errors,
        "errors": errors,
        "checked_at": _now(),
        "paragraph_count": len(paragraphs),
        "evidence_count": len(
            {
                evidence_id
                for paragraph in paragraphs
                for evidence_id in paragraph.get("evidence_ids", [])
            }
        ),
    }


def new_document(value: dict) -> dict:
    value = value if isinstance(value, dict) else {}
    kind = _text(value.get("kind"), limit=30) or "report"
    stage = _text(value.get("stage"), limit=30) or "outline"
    title = _text(value.get("title"), limit=300)
    if kind not in DOCUMENT_KINDS:
        raise ValueError("稿件类型必须为 report 或 brief")
    if stage not in DOCUMENT_STAGES:
        raise ValueError("无效稿件阶段")
    if not title:
        raise ValueError("稿件标题必填")
    paragraphs = [
        normalize_paragraph(paragraph)
        for paragraph in value.get("paragraphs", [])
        if isinstance(paragraph, dict)
    ]
    now = _now()
    content = {
        "kind": kind,
        "title": title,
        "stage": stage,
        "outline": _text(value.get("outline")),
        "paragraphs": paragraphs,
        "revision": 1,
        "revisions": [],
        "created_at": now,
        "updated_at": now,
    }
    content["validation"] = validate_document(content)
    return content


def apply_document_changes(current: dict, changes: dict) -> dict:
    changes = changes if isinstance(changes, dict) else {}
    result = copy.deepcopy(current)
    snapshot = {
        "revision": int(current.get("revision") or 1),
        "title": current.get("title", ""),
        "stage": current.get("stage", "outline"),
        "outline": current.get("outline", ""),
        "paragraphs": copy.deepcopy(current.get("paragraphs", [])),
        "saved_at": current.get("updated_at") or _now(),
        "note": _text(changes.get("revision_note"), limit=1000),
    }
    result.setdefault("revisions", []).append(snapshot)
    if "title" in changes:
        result["title"] = _text(changes.get("title"), limit=300)
        if not result["title"]:
            raise ValueError("稿件标题必填")
    if "kind" in changes:
        kind = _text(changes.get("kind"), limit=30)
        if kind not in DOCUMENT_KINDS:
            raise ValueError("稿件类型必须为 report 或 brief")
        result["kind"] = kind
    if "stage" in changes:
        stage = _text(changes.get("stage"), limit=30)
        if stage not in DOCUMENT_STAGES:
            raise ValueError("无效稿件阶段")
        result["stage"] = stage
    if "outline" in changes:
        result["outline"] = _text(changes.get("outline"))
    if "paragraphs" in changes:
        if not isinstance(changes["paragraphs"], list):
            raise ValueError("paragraphs 必须为数组")
        existing = {
            item.get("paragraph_id"): item
            for item in current.get("paragraphs", [])
            if item.get("paragraph_id")
        }
        result["paragraphs"] = [
            normalize_paragraph(
                paragraph,
                existing_id=existing.get(
                    _text(paragraph.get("paragraph_id"), limit=100), {}
                ).get("paragraph_id", ""),
            )
            for paragraph in changes["paragraphs"]
            if isinstance(paragraph, dict)
        ]
    result["revision"] = int(current.get("revision") or 1) + 1
    result["updated_at"] = _now()
    result["validation"] = validate_document(result)
    return result


def evidence_ids_for_document(content: dict) -> list[str]:
    return list(
        dict.fromkeys(
            evidence_id
            for paragraph in content.get("paragraphs", [])
            for evidence_id in paragraph.get("evidence_ids", [])
        )
    )


def new_publication_item(document: dict, actor_user_id: str) -> dict:
    validation = validate_document(document["content"])
    now = _now()
    return {
        "document_id": document["record_id"],
        "title": document["content"].get("title", ""),
        "kind": document["content"].get("kind", "report"),
        "status": "editing" if validation.get("ready") else "evidence_needed",
        "position": 0,
        "document_version": document["version"],
        "document_content_hash": document["content_hash"],
        "signed_snapshot": None,
        "publication_semantics": {
            "scope": "local_approval_snapshot",
            "external_published": False,
            "notice": LOCAL_SIGNING_NOTICE,
        },
        "created_by": actor_user_id,
        "created_at": now,
        "updated_at": now,
    }


def build_source_index(document_content: dict, evidence_records: list[dict]) -> list[dict]:
    required = set(evidence_ids_for_document(document_content))
    index = []
    for evidence in evidence_records:
        if evidence["record_id"] not in required:
            continue
        content = evidence.get("content") or {}
        provenance = content.get("provenance")
        if not isinstance(provenance, dict):
            provenance = {}
        index.append(
            {
                "record_id": evidence["record_id"],
                "title": _text(content.get("title"), limit=500),
                "source": _text(content.get("source"), limit=300),
                "url": _text(
                    content.get("url")
                    or content.get("link")
                    or provenance.get("url"),
                    limit=2000,
                ),
                "source_hash": _text(
                    content.get("source_hash")
                    or content.get("content_hash")
                    or evidence.get("content_hash"),
                    limit=200,
                ),
                "archived_at": _text(
                    content.get("archived_at")
                    or provenance.get("archived_at")
                    or evidence.get("updated_at"),
                    limit=100,
                ),
            }
        )
    return index


def signed_publication_content(
    publication: dict,
    document: dict,
    source_index: list[dict],
    actor_user_id: str,
) -> dict:
    if publication.get("status") != "pending_approval":
        raise ValueError("只有待签发稿件可以签发")
    validation = validate_document(document["content"])
    if not validation.get("ready"):
        raise ValueError("稿件校验未通过，禁止签发")
    result = copy.deepcopy(publication)
    signed_at = _now()
    result["status"] = "signed"
    result["updated_at"] = signed_at
    result["signed_snapshot"] = {
        "document": copy.deepcopy(document["content"]),
        "source_index": copy.deepcopy(source_index),
        "publication_semantics": {
            "scope": "local_approval_snapshot",
            "external_published": False,
            "notice": LOCAL_SIGNING_NOTICE,
        },
        "receipt": {
            "document_id": document["record_id"],
            "document_version": document["version"],
            "document_content_hash": document["content_hash"],
            "signed_at": signed_at,
            "signed_by": actor_user_id,
        },
    }
    return result


def build_recall_audit_notice(publication: dict) -> dict:
    if publication.get("status") != "recalled":
        raise ValueError("只有已撤回材料可生成撤回审计件")
    snapshot = publication.get("signed_snapshot")
    if not isinstance(snapshot, dict) or not isinstance(
        snapshot.get("receipt"), dict
    ):
        raise ValueError("已撤回材料缺少原签发回执")
    return {
        "status": "recalled",
        "recalled_at": _text(publication.get("recalled_at"), limit=100),
        "recalled_by": _text(publication.get("recalled_by"), limit=200),
        "reason": _text(publication.get("recall_reason"), limit=2000),
        "signed_receipt": copy.deepcopy(snapshot["receipt"]),
        "publication_semantics": copy.deepcopy(
            snapshot.get("publication_semantics")
            or publication.get("publication_semantics")
            or {
                "scope": "local_approval_snapshot",
                "external_published": False,
                "notice": LOCAL_SIGNING_NOTICE,
            }
        ),
    }


def safe_filename(title: str, suffix: str) -> str:
    base = re.sub(r'[<>:"/\\|?*\x00-\x1f]', "_", _text(title, limit=120))
    return f"{base or 'V9稿件'}.{suffix}"


def build_document_docx(
    document_content: dict,
    source_index: list[dict],
    *,
    recall_audit: dict | None = None,
) -> bytes:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt, RGBColor

    if not validate_document(document_content).get("ready"):
        raise ValueError("稿件校验未通过，禁止生成DOCX")
    document = Document()
    if recall_audit:
        for section in document.sections:
            watermark = section.header.paragraphs[0]
            watermark.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = watermark.add_run("已撤回 · 审计件")
            run.bold = True
            run.font.size = Pt(24)
            run.font.color.rgb = RGBColor(0xB0, 0x20, 0x20)
    document.add_heading(document_content.get("title") or "V9 稿件", 0)
    document.add_paragraph(
        f"类型：{document_content.get('kind', 'report')}　"
        f"修订：V{document_content.get('revision', 1)}"
    )
    if recall_audit:
        receipt = recall_audit.get("signed_receipt") or {}
        document.add_heading("已撤回 · 仅供审计", level=1)
        document.add_paragraph(LOCAL_SIGNING_NOTICE)
        document.add_heading("撤回回执", level=2)
        document.add_paragraph(
            f"撤回时间：{recall_audit.get('recalled_at') or '未记录'}；"
            f"撤回操作人：{recall_audit.get('recalled_by') or '未记录'}；"
            f"撤回原因：{recall_audit.get('reason') or '未记录'}"
        )
        document.add_paragraph(
            f"原签发回执：稿件ID={receipt.get('document_id') or '未记录'}；"
            f"版本={receipt.get('document_version') or '未记录'}；"
            f"内容哈希={receipt.get('document_content_hash') or '未记录'}；"
            f"签发时间={receipt.get('signed_at') or '未记录'}；"
            f"签发人={receipt.get('signed_by') or '未记录'}"
        )
    if document_content.get("outline"):
        document.add_heading("大纲", level=1)
        document.add_paragraph(document_content["outline"])
    evidence_numbers = {
        row["record_id"]: index for index, row in enumerate(source_index, start=1)
    }
    for paragraph in document_content.get("paragraphs", []):
        if paragraph.get("heading"):
            document.add_heading(paragraph["heading"], level=1)
        citations = [
            f"[E{evidence_numbers[evidence_id]}]"
            for evidence_id in paragraph.get("evidence_ids", [])
            if evidence_id in evidence_numbers
        ]
        document.add_paragraph(
            f"{paragraph.get('text', '')} {' '.join(citations)}".strip()
        )
        document.add_paragraph(
            f"来源状态：{paragraph.get('source_status', '')}；"
            f"事实核查：{paragraph.get('fact_check', '')}"
        )
    document.add_heading("来源索引", level=1)
    for index, source in enumerate(source_index, start=1):
        document.add_paragraph(
            f"[E{index}] {source.get('title') or '未命名来源'}；"
            f"来源：{source.get('source') or '未标注'}；"
            f"记录ID：{source['record_id']}；"
            f"内容哈希：{source.get('source_hash') or '无'}；"
            f"归档时间：{source.get('archived_at') or '无'}；"
            f"链接：{source.get('url') or '无'}"
        )
    output = BytesIO()
    document.save(output)
    return output.getvalue()


def build_document_pdf(
    document_content: dict,
    source_index: list[dict],
    *,
    recall_audit: dict | None = None,
) -> bytes:
    if not validate_document(document_content).get("ready"):
        raise ValueError("稿件校验未通过，禁止生成PDF")
    from reportlab.lib.enums import TA_LEFT
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.cidfonts import UnicodeCIDFont
    from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer

    try:
        pdfmetrics.getFont("STSong-Light")
    except KeyError:
        pdfmetrics.registerFont(UnicodeCIDFont("STSong-Light"))
    styles = getSampleStyleSheet()
    title_style = ParagraphStyle(
        "V9Title",
        parent=styles["Title"],
        fontName="STSong-Light",
        fontSize=20,
        leading=28,
    )
    body_style = ParagraphStyle(
        "V9Body",
        parent=styles["BodyText"],
        fontName="STSong-Light",
        fontSize=10.5,
        leading=18,
        alignment=TA_LEFT,
    )
    heading_style = ParagraphStyle(
        "V9Heading",
        parent=body_style,
        fontSize=14,
        leading=22,
        spaceBefore=10,
    )
    evidence_numbers = {
        row["record_id"]: index for index, row in enumerate(source_index, start=1)
    }
    story = [
        Paragraph(
            html.escape(document_content.get("title") or "V9 稿件"),
            title_style,
        ),
        Spacer(1, 12),
    ]
    if recall_audit:
        receipt = recall_audit.get("signed_receipt") or {}
        story.extend(
            [
                Paragraph("已撤回 · 仅供审计", heading_style),
                Paragraph(html.escape(LOCAL_SIGNING_NOTICE), body_style),
                Paragraph("撤回回执", heading_style),
                Paragraph(
                    html.escape(f"撤回时间：{recall_audit.get('recalled_at') or '未记录'}"),
                    body_style,
                ),
                Paragraph(
                    html.escape(f"撤回操作人：{recall_audit.get('recalled_by') or '未记录'}"),
                    body_style,
                ),
                Paragraph(
                    html.escape(f"撤回原因：{recall_audit.get('reason') or '未记录'}"),
                    body_style,
                ),
                Paragraph(
                    html.escape(f"原签发稿件ID：{receipt.get('document_id') or '未记录'}"),
                    body_style,
                ),
                Paragraph(
                    html.escape(f"原签发版本：{receipt.get('document_version') or '未记录'}"),
                    body_style,
                ),
                Paragraph(
                    html.escape(f"原签发内容哈希：{receipt.get('document_content_hash') or '未记录'}"),
                    body_style,
                ),
                Paragraph(
                    html.escape(f"原签发时间：{receipt.get('signed_at') or '未记录'}"),
                    body_style,
                ),
                Paragraph(
                    html.escape(f"原签发人：{receipt.get('signed_by') or '未记录'}"),
                    body_style,
                ),
            ]
        )
    for paragraph in document_content.get("paragraphs", []):
        if paragraph.get("heading"):
            story.append(
                Paragraph(html.escape(paragraph["heading"]), heading_style)
            )
        citations = " ".join(
            f"[E{evidence_numbers[evidence_id]}]"
            for evidence_id in paragraph.get("evidence_ids", [])
            if evidence_id in evidence_numbers
        )
        story.append(
            Paragraph(
                html.escape(
                    f"{paragraph.get('text', '')} {citations}".strip()
                ),
                body_style,
            )
        )
    story.append(Paragraph("来源索引", heading_style))
    for index, source in enumerate(source_index, start=1):
        story.append(
            Paragraph(
                html.escape(
                    f"[E{index}] {source.get('title') or '未命名来源'}；"
                    f"来源：{source.get('source') or '未标注'}；"
                    f"记录ID：{source['record_id']}；"
                    f"内容哈希：{source.get('source_hash') or '无'}；"
                    f"归档时间：{source.get('archived_at') or '无'}；"
                    f"链接：{source.get('url') or '无'}"
                ),
                body_style,
            )
        )
    output = BytesIO()
    pdf_title = document_content.get("title") or "V9 稿件"
    if recall_audit:
        pdf_title = f"{pdf_title}-已撤回-审计件"
    pdf = SimpleDocTemplate(
        output,
        pagesize=A4,
        title=pdf_title,
        author="DefenseTracker V9",
    )

    def draw_recall_watermark(canvas, _document):
        if not recall_audit:
            return
        canvas.saveState()
        try:
            canvas.setFillAlpha(0.18)
        except AttributeError:
            pass
        canvas.setFillColorRGB(0.7, 0.12, 0.12)
        canvas.setFont("STSong-Light", 38)
        canvas.translate(A4[0] / 2, A4[1] / 2)
        canvas.rotate(35)
        canvas.drawCentredString(0, 0, "已撤回 · 审计件")
        canvas.restoreState()

    pdf.build(
        story,
        onFirstPage=draw_recall_watermark,
        onLaterPages=draw_recall_watermark,
    )
    return output.getvalue()
