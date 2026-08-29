from datetime import datetime, timezone
from io import BytesIO
from pathlib import Path
import re

import pytest

import report_agent

_RETIRED_USER_MARKER = "131" + "4520"

if report_agent.DOCX_AVAILABLE:
    from docx import Document
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
    from docx.oxml.ns import qn


def _candidate(title="台海方向联合演训值得警惕"):
    return {
        "article_id": "article-001",
        "title": title,
        "summary": "公开报道显示，相关部队围绕台海方向组织联合演训，检验远程火力和体系支撑能力。",
        "source": "Jamestown China Brief",
        "source_cn": "詹姆斯敦中国简报",
        "link": "https://example.test/article-001",
        "date": datetime.now(timezone.utc).isoformat(),
        "quality_score": 91,
        "quality_level": "S",
        "quality_reasons": ["高权威信源", "贴近对华/周边防务重点"],
        "brief_hits": ["PLA备战", "对华威胁"],
    }


def _institution_pack_evidence(count=7):
    return [
        {
            "evidence_id": f"evidence-{idx}",
            "title": f"机构原文{idx}",
            "source": f"Institution {idx}",
            "link": f"https://sources.example.test/report-{idx}",
            "source_type": "已抓取公开报告/原文",
            "payload": {
                "asset_status": "archived",
                "asset_id": f"asset-{idx}",
                "text": f"这是第{idx}份可引用公开原文正文。",
            },
        }
        for idx in range(1, count + 1)
    ]

def _institution_pack_draft(citation_count=7):
    citations = "\n".join(
        f"{idx}. 机构信息要点{idx} [{idx}]" for idx in range(1, citation_count + 1)
    )
    return {
        "payload": {},
        "content": (
            "# 机构开源情报整编包\n\n"
            "## 信息清单\n"
            f"{citations}\n\n"
            "## 短消息\n"
            "### 短消息一\n已核实事实：第一项动态已由原文支持。[1]\n\n"
            "### 短消息二\n来源陈述：第二项动态仍需持续观察。[2]\n\n"
            "### 短消息三\n分析推断：多项公开信号具有联动性。[3]\n\n"
            "## 专题报告\n围绕能力、部署与规则影响展开专题分析。[4][5]\n\n"
            "## 事实来源追溯表\n"
            "| 事实 | 事实层级 | 来源 |\n|---|---|---|\n"
            "| 动态一 | 已核实事实 | [6] |\n\n"
            "## 不确定性与观察指标\n证据边界是公开材料时效有限；观察指标为后续正式披露。[7]\n\n"
            "## 诚实边界\n"
            "1. 本包不能验证未公开的内部意图。\n"
            "2. 本包不能把单一机构判断当作既成事实。\n"
            "3. 本包不能替代后续原文复核与动态更新。\n"
        ),
    }

@pytest.fixture()
def agent_db(monkeypatch, tmp_path):
    db_file = tmp_path / "report_agent.sqlite3"
    monkeypatch.setattr(report_agent, "REPORT_AGENT_DB_FILE", str(db_file))
    return db_file


def test_chinese_date_format_is_locale_independent():
    class DateLike:
        year = 2026
        month = 8
        day = 9

        def strftime(self, _format):
            raise AssertionError("locale-sensitive strftime must not be used")

    assert report_agent._format_chinese_date(DateLike()) == "2026年08月09日"


def test_create_project_uses_report_type_defaults_and_records_event(agent_db):
    project = report_agent.create_project(
        title="台海态势日报",
        report_type="strategic",
        topic="台海",
    )

    assert project["title"] == "台海态势日报"
    assert project["report_type"] == "strategic"
    assert project["target_count"] == 12
    assert project["time_window_days"] == 14
    assert project["voice"] == "strategic_analysis"

    bundle = report_agent.get_project_bundle(project["project_id"])
    assert bundle["project"]["project_id"] == project["project_id"]
    assert bundle["events"][0]["event_type"] == "project_created"


def test_institution_pack_defaults_and_public_sod_path(agent_db):
    project = report_agent.create_project(
        title="机构开源情报整编包",
        report_type="institution_pack",
        topic="印太防务态势",
    )

    assert project["target_count"] == 8
    assert project["time_window_days"] == 30
    assert report_agent.REPORT_TYPE_DEFAULTS["institution_pack"]["label"] == "机构开源情报整编包"

    candidates = [Path(path) for path in report_agent._writing_spec_candidates()]
    assert candidates[0].name == "defensetracker_sod_writing.md"
    assert candidates[0].parent.name == "docs"
    assert any(path.name == "report_agent_sod_writing.md" for path in candidates)
    assert _RETIRED_USER_MARKER not in "\n".join(str(path) for path in candidates)


def test_evidence_lines_use_stable_bracketed_citation_numbers():
    evidence = _institution_pack_evidence(3)

    rendered = report_agent._evidence_lines(evidence)

    assert rendered.count("[1]") == 1
    assert rendered.count("[2]") == 1
    assert rendered.count("[3]") == 1
    assert "1. 【" not in rendered

@pytest.mark.parametrize("builder", [report_agent.build_outline_messages, report_agent.build_draft_messages])
def test_institution_pack_prompts_require_closed_loop_delivery(builder):
    project = {
        "report_type": "institution_pack",
        "title": "机构开源情报整编包",
        "topic": "印太防务态势",
        "client_request": "形成今日可交付闭环",
        "target_count": 8,
    }

    prompt = "\n".join(message["content"] for message in builder(project, _institution_pack_evidence()))

    for requirement in (
        "7—10条",
        "恰好3篇短消息",
        "1篇专题",
        "事实来源追溯表",
        "事实分级",
        "不确定性",
        "观察指标",
        "来源索引",
        "至少3条诚实边界",
        "不得编造",
        "不得越界引用",
    ):
        assert requirement in prompt

def test_institution_pack_delivery_preflight_accepts_closed_loop_package():
    project = {"report_type": "institution_pack"}

    result = report_agent.build_delivery_preflight(
        project,
        _institution_pack_draft(),
        _institution_pack_evidence(),
    )

    assert result["ok"] is True
    assert result["status"] == "ready"
    assert all(check["ok"] for check in result["checks"])
    assert all({"id", "ok", "label", "detail"} <= set(check) for check in result["checks"])
    assert result["counts"]["evidence"] == 7
    assert result["counts"]["short_messages"] == 3
    assert result["counts"]["honest_boundaries"] == 3
    assert result["blocked_domains"] == []
    assert result["missing_citations"] == []
    assert result["out_of_range_citations"] == []

@pytest.mark.parametrize(
    ("mutation", "failed_check", "detail_fragment"),
    [
        ("too_few_sources", "sources", "7—10"),
        ("non_http_link", "source_links", "HTTP"),
        ("malformed_link", "source_links", "HTTP"),
        ("backslash_authority", "source_links", "HTTP"),
        ("userinfo_authority", "source_links", "HTTP"),
        ("blocked_subdomain", "source_policy", "zhihu.com"),
        ("missing_original", "content", "原文"),
        ("directory_seed_with_text", "content", "原文"),
        ("missing_section", "sections", "专题报告"),
        ("too_few_short_messages", "short_messages", "3篇"),
        ("too_many_short_messages", "short_messages", "恰好3篇"),
        ("generic_short_message_children", "short_messages", "恰好3篇"),
        ("missing_fact_level", "fact_levels", "事实分级"),
        ("missing_citation", "citations", "[7]"),
        ("out_of_range_citation", "citation_range", "[8]"),
        ("too_few_boundaries", "honest_boundaries", "3条"),
    ],
)
def test_institution_pack_delivery_preflight_fails_closed_with_specific_gap(
    mutation, failed_check, detail_fragment
):
    project = {"report_type": "institution_pack"}
    evidence = _institution_pack_evidence()
    draft = _institution_pack_draft()

    if mutation == "too_few_sources":
        evidence = evidence[:6]
        draft = _institution_pack_draft(6)
    elif mutation == "non_http_link":
        evidence[0]["link"] = "file:///private/source.pdf"
    elif mutation == "malformed_link":
        evidence[0]["link"] = "https://[invalid"
    elif mutation == "backslash_authority":
        evidence[0]["link"] = r"https://zhihu.com\@evil.com/report"
    elif mutation == "userinfo_authority":
        evidence[0]["link"] = "https://reader@example.test/report"
    elif mutation == "blocked_subdomain":
        evidence[0]["link"] = "https://column.zhihu.com/p/123"
    elif mutation == "missing_original":
        evidence[0]["payload"] = {"asset_status": "partial"}
    elif mutation == "directory_seed_with_text":
        evidence[0]["source_type"] = "智库/报告源"
        evidence[0]["payload"] = {"text": "这只是目录 seed 的描述文本，不是已抓取原文。"}
    elif mutation == "missing_section":
        draft["content"] = draft["content"].replace("## 专题报告", "## 综合研判")
    elif mutation == "too_few_short_messages":
        draft["content"] = draft["content"].replace(
            "### 短消息三\n分析推断：多项公开信号具有联动性。[3]\n\n", ""
        )
    elif mutation == "too_many_short_messages":
        draft["content"] = draft["content"].replace(
            "## 专题报告",
            "### 短消息四\n已核实事实：补充消息仍引用既有来源。[4]\n\n## 专题报告",
        )
    elif mutation == "generic_short_message_children":
        for source, replacement in (
            ("### 短消息一", "### 背景"),
            ("### 短消息二", "### 态势"),
            ("### 短消息三", "### 结论"),
        ):
            draft["content"] = draft["content"].replace(source, replacement)
    elif mutation == "missing_fact_level":
        for term in ("已核实事实", "来源陈述", "分析推断", "事实层级"):
            draft["content"] = draft["content"].replace(term, "研判")
    elif mutation == "missing_citation":
        draft["content"] = draft["content"].replace("[7]", "")
    elif mutation == "out_of_range_citation":
        draft["content"] += "\n越界来源[8]。"
    elif mutation == "too_few_boundaries":
        draft["content"] = draft["content"].replace(
            "3. 本包不能替代后续原文复核与动态更新。\n", ""
        )

    result = report_agent.build_delivery_preflight(project, draft, evidence)
    check = next(item for item in result["checks"] if item["id"] == failed_check)

    assert result["ok"] is False
    assert result["status"] == "blocked"
    assert check["ok"] is False
    assert detail_fragment in check["detail"]

def test_institution_pack_preflight_accepts_fetched_original_body_without_archive_asset():
    evidence = _institution_pack_evidence()
    evidence[0].pop("source_type")
    evidence[0]["payload"] = {
        "source_type": "已抓取公开报告/原文",
        "text": "这是已抓取并转入的公开报告正文。",
    }

    result = report_agent.build_delivery_preflight(
        {"report_type": "institution_pack"},
        _institution_pack_draft(),
        evidence,
    )

    assert result["ok"] is True

def test_institution_pack_preflight_blocks_citations_only_in_information_list():
    draft = _institution_pack_draft()
    info_list, remainder = draft["content"].split("## 短消息", 1)
    draft["content"] = info_list + "## 短消息" + re.sub(r"\[(\d+)\]", "", remainder)

    result = report_agent.build_delivery_preflight(
        {"report_type": "institution_pack"},
        draft,
        _institution_pack_evidence(),
    )
    check = next(item for item in result["checks"] if item["id"] == "section_citations")

    assert result["missing_citations"] == []
    assert result["ok"] is False
    assert check["ok"] is False
    assert "短消息一" in check["detail"]
    assert "专题报告" in check["detail"]
    assert "事实来源追溯表" in check["detail"]

@pytest.mark.parametrize(
    ("trace", "detail_fragment"),
    [
        (f"制作说明：遵循{_RETIRED_USER_MARKER} SOD/SOP。", "SOD/SOP"),
        (r"工作底稿：F:\private\draft.md", "本地绝对路径"),
        ("本文由AI生成并整理。", "AI制作痕迹"),
        ("作为AI语言模型，我依据提示词完成本稿。", "AI制作痕迹"),
        ("本报告由生成式人工智能协助起草。", "AI制作痕迹"),
        ("本稿在ChatGPT帮助下形成。", "AI制作痕迹"),
        ("本文采用大模型协助起草。", "AI制作痕迹"),
    ],
)
def test_institution_pack_delivery_preflight_blocks_production_traces(trace, detail_fragment):
    draft = _institution_pack_draft()
    draft["content"] += f"\n{trace}\n"

    result = report_agent.build_delivery_preflight(
        {"report_type": "institution_pack"},
        draft,
        _institution_pack_evidence(),
    )
    check = next(item for item in result["checks"] if item["id"] == "trace_hygiene")

    assert result["ok"] is False
    assert check["ok"] is False
    assert detail_fragment in check["detail"]

def test_institution_pack_trace_hygiene_allows_source_backed_ai_subject_matter():
    draft = _institution_pack_draft()
    draft["content"] += "\n本文分析AI生成图像对舆论环境的影响，并以来源[1]为依据。\n"

    result = report_agent.build_delivery_preflight(
        {"report_type": "institution_pack"},
        draft,
        _institution_pack_evidence(),
    )
    check = next(item for item in result["checks"] if item["id"] == "trace_hygiene")

    assert check["ok"] is True

def test_institution_pack_trace_hygiene_allows_operational_sop_subject_matter():
    draft = _institution_pack_draft()
    draft["content"] += "\n已核实事实：部队修订 SOP 后开展训练。[1]\n"

    result = report_agent.build_delivery_preflight(
        {"report_type": "institution_pack"},
        draft,
        _institution_pack_evidence(),
    )
    check = next(item for item in result["checks"] if item["id"] == "trace_hygiene")

    assert check["ok"] is True

def test_delivery_preflight_and_export_remain_backward_compatible_for_other_types():
    project = {"report_type": "strategic", "client_request": ""}
    draft = {"payload": {}, "content": "普通战略报告正文。"}

    result = report_agent.build_delivery_preflight(project, draft, [])
    quality = report_agent.assert_report_exportable(draft, project, evidence=[])

    assert result["ok"] is True
    assert result["status"] == "not_required"
    assert quality["target_word_count"] == 0

def test_institution_pack_export_error_names_first_concrete_gap():
    project = {"report_type": "institution_pack", "client_request": ""}

    with pytest.raises(ValueError, match="证据数量.*7—10"):
        report_agent.assert_report_exportable(
            _institution_pack_draft(),
            project,
            evidence=[],
        )

@pytest.mark.skipif(not report_agent.DOCX_AVAILABLE, reason="python-docx 未安装")
def test_build_report_docx_passes_institution_pack_evidence_to_preflight():
    buf = report_agent.build_report_docx(
        {"report_type": "institution_pack", "title": "机构开源情报整编包", "client_request": ""},
        _institution_pack_draft(),
        _institution_pack_evidence(),
    )

    assert buf.getvalue().startswith(b"PK")
    doc = Document(BytesIO(buf.getvalue()))
    visible_text = [p.text for p in doc.paragraphs]
    visible_text.extend(
        cell.text for table in doc.tables for row in table.rows for cell in row.cells
    )
    visible_text.extend(
        p.text for section in doc.sections for p in section.header.paragraphs
    )
    combined = "\n".join(visible_text)
    assert _RETIRED_USER_MARKER not in combined
    assert "SOD/SOP" not in combined
    assert "FACT-DATA-CITE" not in combined

@pytest.mark.skipif(not report_agent.DOCX_AVAILABLE, reason="python-docx 未安装")
def test_institution_pack_docx_tables_use_fixed_geometry_and_real_headers():
    buf = report_agent.build_report_docx(
        {"report_type": "institution_pack", "title": "机构开源情报整编包", "client_request": ""},
        _institution_pack_draft(),
        _institution_pack_evidence(),
    )
    doc = Document(BytesIO(buf.getvalue()))
    section = doc.sections[-1]
    printable_width = int(
        section.page_width.twips - section.left_margin.twips - section.right_margin.twips
    )

    assert len(doc.tables) >= 3
    for table in doc.tables:
        tbl_pr = table._tbl.tblPr
        tbl_width = tbl_pr.find(qn("w:tblW"))
        tbl_indent = tbl_pr.find(qn("w:tblInd"))
        tbl_layout = tbl_pr.find(qn("w:tblLayout"))
        assert tbl_width is not None and tbl_width.get(qn("w:type")) == "dxa"
        assert tbl_indent is not None and tbl_indent.get(qn("w:type")) == "dxa"
        assert tbl_layout is not None and tbl_layout.get(qn("w:type")) == "fixed"
        width = int(tbl_width.get(qn("w:w")))
        indent = int(tbl_indent.get(qn("w:w")))
        assert indent + width <= printable_width
        grid = [int(col.get(qn("w:w"))) for col in table._tbl.tblGrid]
        assert sum(grid) == width

        start_margins = set()
        for row in table.rows:
            cell_widths = []
            for cell in row.cells:
                tc_pr = cell._tc.get_or_add_tcPr()
                tc_width = tc_pr.find(qn("w:tcW"))
                cell_widths.append(int(tc_width.get(qn("w:w"))))
                start_margin = tc_pr.find(qn("w:tcMar"))
                assert start_margin is not None
                start_node = start_margin.find(qn("w:start"))
                assert start_node is not None
                start_margins.add(int(start_node.get(qn("w:w"))))
            assert cell_widths == grid
        assert len(start_margins) == 1
        assert indent == next(iter(start_margins))

    for table in doc.tables[1:]:
        header = table.rows[0]._tr.get_or_add_trPr().find(qn("w:tblHeader"))
        assert header is not None and header.get(qn("w:val")) in ("1", "true")

@pytest.mark.skipif(not report_agent.DOCX_AVAILABLE, reason="python-docx 未安装")
def test_fixed_table_geometry_rejects_merged_cells():
    doc = Document()
    table = doc.add_table(rows=1, cols=3)
    table.cell(0, 0).merge(table.cell(0, 1))

    with pytest.raises(ValueError, match="合并单元格"):
        report_agent._apply_table_geometry(doc, table, [1, 1, 1], cell_margin=140)

    doc = Document()
    table = doc.add_table(rows=2, cols=3)
    table.cell(0, 0).merge(table.cell(1, 0))

    with pytest.raises(ValueError, match="合并单元格"):
        report_agent._apply_table_geometry(doc, table, [1, 1, 1], cell_margin=140)

@pytest.mark.skipif(not report_agent.DOCX_AVAILABLE, reason="python-docx 未安装")
def test_fixed_table_geometry_reduces_margin_for_many_columns():
    doc = Document()
    table = doc.add_table(rows=1, cols=40)
    report_agent._apply_table_geometry(doc, table, [1] * 40, cell_margin=140)

    for cell in table.rows[0].cells:
        tc_pr = cell._tc.get_or_add_tcPr()
        width = int(tc_pr.find(qn("w:tcW")).get(qn("w:w")))
        margins = tc_pr.find(qn("w:tcMar"))
        start = int(margins.find(qn("w:start")).get(qn("w:w")))
        end = int(margins.find(qn("w:end")).get(qn("w:w")))
        assert start + end < width

def test_create_project_from_client_request_derives_strategy_title(agent_db):
    project = report_agent.create_project(
        title="",
        report_type="strategic",
        client_request="帮我做一个台海军力平衡报告",
    )

    assert project["title"] == "台海军力平衡战略分析报告"
    assert project["topic"] == "台海军力平衡"
    assert project["client_request"] == "帮我做一个台海军力平衡报告"


def test_create_project_extracts_requested_source_count_without_upper_cap(agent_db):
    project = report_agent.create_project(
        title="",
        report_type="strategic",
        client_request="帮我做一个台海军力平衡报告，搜集120份信息源",
    )

    assert project["target_count"] == 120
    assert project["title"] == "台海军力平衡战略分析报告"


def test_upsert_evidence_preserves_quality_source_link_and_dedup_id(agent_db):
    project = report_agent.create_project("专题短报", "short_topic", topic="六代机")

    first = report_agent.upsert_project_evidence(project["project_id"], [_candidate()])
    second = report_agent.upsert_project_evidence(project["project_id"], [_candidate(title="标题更新")])

    assert len(first) == 1
    assert second[0]["evidence_id"] == first[0]["evidence_id"]
    assert second[0]["title"] == "标题更新"
    assert second[0]["source"] == "Jamestown China Brief"
    assert second[0]["link"] == "https://example.test/article-001"
    assert second[0]["quality_score"] == 91
    assert second[0]["quality_level"] == "S"
    assert "高权威信源" in second[0]["quality_reasons"]


def test_build_messages_and_save_draft(agent_db):
    project = report_agent.create_project("台海战略分析报告", "strategic", topic="台海")
    evidence = report_agent.upsert_project_evidence(project["project_id"], [_candidate()])

    outline_messages = report_agent.build_outline_messages(project, evidence)
    assert outline_messages[0]["role"] == "system"
    assert "防务战略分析报告" in outline_messages[0]["content"]
    assert "不是要讯" in outline_messages[0]["content"]
    assert "目录提纲" in outline_messages[1]["content"]
    assert "台海方向联合演训值得警惕" in outline_messages[1]["content"]

    draft = report_agent.save_draft(
        project["project_id"],
        kind="draft",
        content="## 核心判断\n台海方向联合演训值得持续跟踪。",
        model="unit-test-model",
    )

    assert draft["kind"] == "draft"
    assert draft["model"] == "unit-test-model"
    bundle = report_agent.get_project_bundle(project["project_id"])
    assert bundle["drafts"][0]["draft_id"] == draft["draft_id"]
    assert any(e["event_type"] == "draft_saved" for e in bundle["events"])


def test_draft_messages_include_sod_writing_requirements(agent_db, monkeypatch, tmp_path):
    spec_file = tmp_path / "defensetracker_sod_writing.md"
    spec_file.write_text(
        "# DefenseTracker SOD写作要求\n"
        "必须采用技术—能力—规则三维分析框架。\n"
        "每段遵循PARA结构，并建立FACT-DATA-CITE证据链。",
        encoding="utf-8",
    )
    monkeypatch.setattr(report_agent, "_WRITING_SPEC_ROOT", tmp_path)
    monkeypatch.setattr(report_agent, "REPORT_AGENT_WRITING_SPEC_FILE", str(spec_file), raising=False)
    project = report_agent.create_project("台海战略分析报告", "strategic", topic="台海")
    evidence = report_agent.upsert_project_evidence(project["project_id"], [_candidate()])

    messages = report_agent.build_draft_messages(project, evidence)
    prompt = "\n".join(m["content"] for m in messages)

    assert "DefenseTracker" in prompt
    assert "技术—能力—规则" in prompt
    assert "PARA" in prompt
    assert "FACT-DATA-CITE" in prompt
    assert str(spec_file) not in prompt


def test_writing_spec_override_cannot_read_outside_public_docs(
    monkeypatch, tmp_path
):
    docs = tmp_path / "docs"
    docs.mkdir()
    secret = tmp_path / "private.txt"
    secret.write_text("must-not-reach-an-ai-provider", encoding="utf-8")
    monkeypatch.setattr(report_agent, "_WRITING_SPEC_ROOT", docs)
    monkeypatch.setattr(
        report_agent,
        "REPORT_AGENT_WRITING_SPEC_FILE",
        str(secret),
    )

    with pytest.raises(ValueError, match="writing specification file is unsafe"):
        report_agent.load_report_writing_requirements()


def test_writing_spec_rejects_oversized_file(monkeypatch, tmp_path):
    docs = tmp_path / "docs"
    docs.mkdir()
    oversized = docs / "oversized.md"
    oversized.write_bytes(b"x" * (report_agent._WRITING_SPEC_MAX_BYTES + 1))
    monkeypatch.setattr(report_agent, "_WRITING_SPEC_ROOT", docs)
    monkeypatch.setattr(
        report_agent,
        "REPORT_AGENT_WRITING_SPEC_FILE",
        str(oversized),
    )

    with pytest.raises(ValueError, match="writing specification file is unsafe"):
        report_agent.load_report_writing_requirements()


def test_writing_spec_rejects_file_swap_between_check_and_open(
    monkeypatch, tmp_path
):
    docs = tmp_path / "docs"
    docs.mkdir()
    spec = docs / "spec.md"
    replacement = docs / "replacement.md"
    spec.write_text("approved public requirements", encoding="utf-8")
    replacement.write_text("must-not-reach-an-ai-provider", encoding="utf-8")
    real_open = report_agent.os.open
    swapped = False

    def racing_open(path, flags, *args, **kwargs):
        nonlocal swapped
        if Path(path) == spec and not swapped:
            swapped = True
            spec.unlink()
            replacement.replace(spec)
        return real_open(path, flags, *args, **kwargs)

    monkeypatch.setattr(report_agent, "_WRITING_SPEC_ROOT", docs)
    monkeypatch.setattr(
        report_agent,
        "REPORT_AGENT_WRITING_SPEC_FILE",
        str(spec),
    )
    monkeypatch.setattr(report_agent.os, "open", racing_open)

    with pytest.raises(ValueError, match="writing specification file is unsafe"):
        report_agent.load_report_writing_requirements()
    assert swapped


def test_agent_ui_uses_complete_report_button():
    html = (Path(__file__).resolve().parents[1] / "templates" / "index.html").read_text(encoding="utf-8")

    assert 'onclick="agentGenerateDraft()">生成完整报告</button>' in html


def test_report_rules_markdown_records_forbidden_classification_terms():
    path = Path(__file__).resolve().parents[1] / "docs" / "报告Agent写作导出规范.md"

    assert path.exists()
    text = path.read_text(encoding="utf-8")
    assert "秘密" in text
    assert "机密" in text
    assert "绝密" in text
    assert "导出前必须阻断" in text


def test_extract_target_word_count_and_sanitize_forbidden_terms():
    assert report_agent.extract_target_word_count("报告要1万字") == 10000
    assert report_agent.extract_target_word_count("生成一万字") == 10000
    assert report_agent.extract_target_word_count("正文不少于12000字") == 12000

    cleaned = report_agent.sanitize_report_text("绝密报告：这是机密材料和秘密来源。")

    assert "绝密" not in cleaned
    assert "机密" not in cleaned
    assert "秘密" not in cleaned
    assert "报告" in cleaned


def test_report_text_rejects_oversized_payload_and_line_before_regexes(monkeypatch):
    original_sub = report_agent.re.sub

    def reject_regex(*_args, **_kwargs):
        pytest.fail("overlong report text must not reach regexes")

    monkeypatch.setattr(report_agent.re, "sub", reject_regex)
    with pytest.raises(ValueError, match="2 MiB"):
        report_agent.sanitize_report_text("x" * (report_agent.MAX_REPORT_TEXT_CHARS + 1))
    with pytest.raises(ValueError, match="32 KiB"):
        report_agent.sanitize_report_text("x" * (report_agent.MAX_REPORT_LINE_CHARS + 1))
    monkeypatch.setattr(report_agent.re, "sub", original_sub)


def test_create_project_rejects_overlong_request_before_topic_regexes(agent_db, monkeypatch):
    monkeypatch.setattr(
        report_agent,
        "_derive_topic_from_request",
        lambda _value: pytest.fail("overlong requests must not reach topic regexes"),
    )

    with pytest.raises(ValueError, match="4096"):
        report_agent.create_project(client_request="x" * 4097)


def test_assert_report_exportable_ignores_word_count_in_body():
    """回归：报告正文里的'可扩写至8000字'不应被当成目标字数而阻断导出。
    字数目标只应来自用户意图（client_request / payload.target_word_count），不从生成的正文解析。"""
    draft = {
        "payload": {},  # 用户未设任何目标字数
        "content": "本报告分析中国高超声速武器发展态势。可进一步扩写至8000字的完整版本。",
    }
    project = {"client_request": "分析中国高超声速武器发展态势"}  # 无任何字数要求
    quality = report_agent.assert_report_exportable(draft, project)  # 不应抛"低于目标字数"
    assert quality["target_word_count"] == 0

    # 反向保证：用户在 payload 显式设了目标字数时，短稿仍被正确阻断导出
    blocked = {"payload": {"target_word_count": 8000}, "content": "太短。"}
    with pytest.raises(ValueError):
        report_agent.assert_report_exportable(blocked, {"client_request": ""})


def test_markdown_hr_not_treated_as_table():
    """回归：含 | 的正文行后跟纯 --- 水平线，不应被误判为 markdown 表格而吞掉正文。"""
    blocks = report_agent._markdown_blocks("某系统 A|B 说明如下。\n---\n下一段正文。")
    assert all(b.get("type") != "table" for b in blocks)   # --- 是水平线不是表分隔行
    assert any("A|B" in str(b) for b in blocks)             # 含 | 的正文未被表格吞掉
    # 正向保证：带 | 分隔行的真表格仍被正确解析为表格
    real = report_agent._markdown_blocks("| 维度 | 研判 |\n|---|---|\n| 技术 | 领先 |")
    assert any(b.get("type") == "table" for b in real)


@pytest.mark.skipif(not report_agent.DOCX_AVAILABLE, reason="python-docx 未安装")
def test_build_report_docx_uses_defensetracker_paper_format(agent_db):
    project = report_agent.create_project("台海军力平衡战略分析报告", "strategic", topic="台海")
    draft = report_agent.save_draft(
        project["project_id"],
        "draft",
        "\n".join([
            "# 台海军力平衡战略分析报告",
            "## 摘要",
            "本文基于公开源资料，围绕技术—能力—规则三维框架评估台海军力平衡，严禁使用绝密、机密、秘密等等级标识。",
            "## 第一章  技术—能力—规则三维分析框架",
            "### 1.1 技术维度",
            "关键判断必须具备FACT-DATA-CITE证据链，并避免单源孤证。",
            "| 维度 | 研判 |",
            "|---|---|",
            "| 技术 | 关键装备迭代加快 |",
            "| 能力 | 跨域杀伤链更完整 |",
        ]),
        model="unit-test-model",
    )
    evidence = report_agent.upsert_project_evidence(project["project_id"], [_candidate()])

    buf = report_agent.build_report_docx(project, draft, evidence)
    doc = Document(BytesIO(buf.getvalue()))
    section = doc.sections[0]
    texts = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    table_texts = [cell.text.strip() for table in doc.tables for row in table.rows for cell in row.cells if cell.text.strip()]
    header_texts = [p.text.strip() for sec in doc.sections for p in sec.header.paragraphs if p.text.strip()]

    assert round(section.top_margin.cm, 1) == 3.7
    assert round(section.bottom_margin.cm, 1) == 3.5
    assert round(section.left_margin.cm, 1) == 2.8
    assert round(section.right_margin.cm, 1) == 2.6
    assert texts[0] == "台海军力平衡战略分析报告"
    assert not any("绝密" in text or "机密" in text or "秘密" in text for text in texts)
    assert any("研究报告" in text for text in texts)
    assert "目          录" in texts
    assert any("DefenseTracker SOD/SOP" in text for text in texts + table_texts + header_texts)
    assert any("第一章  技术—能力—规则三维分析框架" in text for text in texts)
    assert any("附录A：来源索引" in text for text in texts)
    assert len(doc.tables) >= 3
    matrix = next(table for table in doc.tables if table.cell(0, 0).text == "维度")
    assert matrix.cell(1, 1).text == "关键装备迭代加快"
    source_index = next(table for table in doc.tables if table.cell(0, 0).text == "序号")
    assert source_index.cell(0, 2).text == "资料标题"

    title_run = doc.paragraphs[0].runs[0]
    assert title_run._element.rPr.rFonts.get(qn("w:eastAsia")) == "方正小标宋简体"
    assert title_run.font.size.pt == 26
    assert doc.paragraphs[0].alignment == WD_PARAGRAPH_ALIGNMENT.CENTER

    body_para = next(p for p in doc.paragraphs if p.text.startswith("本文基于公开源资料"))
    assert body_para.alignment == WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    assert body_para.paragraph_format.first_line_indent is not None


def test_build_source_candidates_from_thinktank_directory(agent_db):
    directory = [
        {
            "id": "pla_research",
            "category": "PLA专项研究机构",
            "sites": [
                {
                    "name": "RAND China Research",
                    "name_cn": "兰德中国研究",
                    "url": "https://www.rand.org/topics/china.html",
                    "desc_cn": "军事现代化深度报告",
                    "desc_en": "China military modernization reports",
                }
            ],
        }
    ]
    project = report_agent.create_project("军事现代化战略分析报告", "strategic", topic="军事现代化")

    candidates = report_agent.build_source_candidates(project, directory, limit=5)

    assert candidates
    assert candidates[0]["source_type"] == "智库/报告源"
    assert candidates[0]["source"] == "RAND China Research"
    assert candidates[0]["quality_level"] == "A"
    assert "军事现代化" in candidates[0]["summary"]


def test_build_source_candidates_filters_china_sources_for_iran_missile_topic(agent_db):
    directory = [
        {
            "id": "china_zone",
            "category": "中国防务智库",
            "sites": [
                {
                    "name": "RAND China Research",
                    "name_cn": "兰德中国研究",
                    "url": "https://www.rand.org/topics/china.html",
                    "desc_cn": "中国军事现代化与台海军力评估",
                    "desc_en": "China military modernization reports",
                }
            ],
        },
        {
            "id": "missile_mideast_research",
            "category": "导弹与中东安全研究",
            "sites": [
                {
                    "name": "CSIS Missile Threat",
                    "name_cn": "CSIS导弹威胁项目",
                    "url": "https://missilethreat.csis.org/",
                    "desc_cn": "伊朗弹道导弹、巡航导弹与导弹作战能力评估",
                    "desc_en": "Iran ballistic missile and missile force assessments",
                }
            ],
        },
    ]
    project = report_agent.create_project(
        title="伊朗导弹作战运用战例梳理战略分析报告",
        report_type="strategic",
        topic="伊朗导弹作战运用战例梳理",
    )

    candidates = report_agent.build_source_candidates(project, directory, limit=5)

    assert candidates
    assert candidates[0]["source"] == "CSIS Missile Threat"
    assert all("China" not in row["source"] for row in candidates)


def test_build_source_candidates_honors_large_requested_source_count(agent_db):
    directory = [{
        "id": "pla_research",
        "category": "PLA专项研究机构",
        "sites": [
            {
                "name": f"Source {i}",
                "name_cn": f"报告源{i}",
                "url": f"https://example.test/source-{i}",
                "desc_cn": "台海军力平衡长期研究报告",
                "desc_en": "Taiwan balance report",
            }
            for i in range(130)
        ],
    }]
    project = report_agent.create_project(
        title="",
        report_type="strategic",
        client_request="帮我做一个台海军力平衡报告，搜集120份信息源",
    )

    candidates = report_agent.build_source_candidates(project, directory, limit=project["target_count"])

    assert len(candidates) == 120


def test_compose_evidence_verdict_varies_by_signal():
    ev_topic = {
        "source_type": "公开信息",
        "quality_level": "S",
        "dims": {"source": 15, "topic": 34, "density": 8, "novelty": 4, "writability": 6},
        "brief_hits": ["PLA备战", "对华威胁"],
        "date": "",
    }
    ev_density = {
        "source_type": "公开信息",
        "quality_level": "A",
        "dims": {"source": 10, "topic": 12, "density": 19, "novelty": 2, "writability": 3},
        "brief_hits": [],
        "date": "",
    }
    v_topic = report_agent._compose_evidence_verdict(ev_topic)
    v_density = report_agent._compose_evidence_verdict(ev_density)

    assert v_topic and v_density
    assert v_topic != v_density                       # 逐条拼装，非死模板
    assert len(v_topic) <= 40 and len(v_density) <= 40
    assert "契合选题" in v_topic                        # 最高维=topic
    assert "命中" in v_topic
    assert "信息密度" in v_density                      # 最高维=density


def test_compose_evidence_verdict_thinktank_without_dims():
    ev = {
        "source_type": "智库/报告源",
        "quality_level": "A",
        "quality_reasons": ["智库/报告源", "PLA专项研究机构"],
        "brief_hits": [],
        "date": "",
    }
    verdict = report_agent._compose_evidence_verdict(ev)
    assert "权威智库源" in verdict
    assert "PLA专项研究机构" in verdict                  # 类别带入，可区分不同智库卡


def test_evidence_row_surfaces_dims_and_verdict(agent_db):
    project = report_agent.create_project("台海态势", "strategic", topic="台海")
    candidate = {
        "article_id": "art-dims-1",
        "title": "台海联合演训检验体系支撑值得关注",
        "summary": "公开报道显示相关部队组织联合演训，检验远程火力与体系支撑。",
        "source": "Jamestown China Brief",
        "date": datetime.now(timezone.utc).isoformat(),
        "quality_score": 88,
        "quality_level": "A",
        "quality_reasons": ["高权威信源"],
        "brief_hits": ["PLA备战"],
        # select_quality_candidates 会把完整 quality blob 挂进候选 → payload
        "quality": {
            "total": 88,
            "level": "A",
            "dims": {"source": 17, "topic": 30, "density": 15, "novelty": 6, "writability": 8},
            "reasons": ["高权威信源"],
            "penalties": [],
        },
    }
    report_agent.upsert_project_evidence(project["project_id"], [candidate])
    evidence = report_agent.get_project_evidence(project["project_id"])

    assert len(evidence) == 1
    ev = evidence[0]
    assert ev["dims"]["topic"] == 30                    # 5 维经 payload 透出到证据行
    assert ev["verdict_line"]
    assert "命中" in ev["verdict_line"]


def test_build_source_candidates_uses_level_fn_when_provided(agent_db):
    directory = [{
        "id": "pla_research",
        "category": "PLA专项研究机构",
        "sites": [{
            "name": "RAND China Research",
            "name_cn": "兰德中国研究",
            "url": "https://www.rand.org/topics/china.html",
            "desc_cn": "军事现代化深度报告",
            "desc_en": "China military modernization reports",
        }],
    }]
    project = report_agent.create_project("军事现代化战略分析报告", "strategic", topic="军事现代化")

    def _lvl(score):
        return "S" if score >= 90 else "A" if score >= 80 else "B"

    with_fn = report_agent.build_source_candidates(project, directory, limit=5, level_fn=_lvl)
    without_fn = report_agent.build_source_candidates(project, directory, limit=5)

    assert with_fn[0]["quality_level"] == _lvl(with_fn[0]["quality_score"])   # 分数驱动，非硬编码
    assert without_fn[0]["quality_level"] == "A"                              # 默认保持不变


def test_evidence_surfaces_multi_source_corroboration(agent_db):
    project = report_agent.create_project("南海态势", "strategic", topic="南海")
    multi = {
        "article_id": "art-corrob-multi",
        "title": "多国海军南海方向演训动态汇总",
        "summary": "多家外媒同步报道相关海军活动。",
        "source": "Reuters",
        "date": datetime.now(timezone.utc).isoformat(),
        "quality_score": 80, "quality_level": "A",
        "quality_reasons": ["高权威信源"], "brief_hits": ["对华威胁"],
        "_sources": ["Reuters", "USNI News", "Naval News"],   # _dedup_articles 跨源合并结果
    }
    single = {
        "article_id": "art-corrob-single",
        "title": "某单一来源披露的装备动态",
        "summary": "仅单一来源报道。",
        "source": "SomeBlog",
        "date": datetime.now(timezone.utc).isoformat(),
        "quality_score": 60, "quality_level": "C",
        "quality_reasons": [], "brief_hits": [],
        "_sources": ["SomeBlog"],
    }
    report_agent.upsert_project_evidence(project["project_id"], [multi, single])
    ev = {e["title"]: e for e in report_agent.get_project_evidence(project["project_id"])}

    assert ev["多国海军南海方向演训动态汇总"]["corroboration_count"] == 3
    assert set(ev["多国海军南海方向演训动态汇总"]["corroborating_sources"]) == {"Reuters", "USNI News", "Naval News"}
    assert ev["某单一来源披露的装备动态"]["corroboration_count"] == 1


def test_thinktank_source_has_no_corroboration(agent_db):
    directory = [{
        "id": "pla_research", "category": "PLA专项研究机构",
        "sites": [{"name": "RAND China Research", "name_cn": "兰德中国研究",
                   "url": "https://www.rand.org/topics/china.html",
                   "desc_cn": "军事现代化", "desc_en": "reports"}],
    }]
    project = report_agent.create_project("军力评估", "strategic", topic="军力评估")
    cands = report_agent.build_source_candidates(project, directory, limit=3)
    report_agent.upsert_project_evidence(project["project_id"], cands)
    ev = report_agent.get_project_evidence(project["project_id"])

    assert ev[0]["corroboration_count"] == 0   # 智库目录源无跨源合并信息，不适用


def test_estimate_reading_minutes():
    assert report_agent.estimate_reading_minutes(0) == 1
    assert report_agent.estimate_reading_minutes(500) == 1
    assert report_agent.estimate_reading_minutes(9000) == 18


def test_build_newspaper_front_matter_extracts_structure():
    project = {"report_type": "strategic", "topic": "台海军力平衡", "voice": "newspaper"}
    content = (
        "# 台海军力平衡战略分析报告\n\n"
        "## 执行摘要\n本报告综述当前态势……\n\n"
        "## 核心判断\n"
        "1. 判断一：区域拒止能力显著增强，值得警惕。\n"
        "2. 判断二：联合作战体系加速成型。\n\n"
        "## 战略影响研判\n……\n\n"
        "## 来源附录\n……\n"
    )
    fm = report_agent.build_newspaper_front_matter(project, content, issue_date="2026.07.03")

    assert fm["issue"] == "VOL.2026.07.03"
    assert "战略分析报告" in fm["byline"] and "台海军力平衡" in fm["byline"]
    assert fm["reading_minutes"] >= 1
    assert fm["toc"] == ["执行摘要", "核心判断", "战略影响研判", "来源附录"]
    assert fm["section_count"] == 4
    assert len(fm["cards"]) == 2
    assert "区域拒止能力显著增强" in fm["cards"][0]


def test_bundle_attaches_newspaper_when_report_draft_exists(agent_db):
    project = report_agent.create_project("台海态势", "strategic", topic="台海", voice="newspaper")
    pid = project["project_id"]
    report_agent.save_draft(
        pid, "draft",
        "# 台海报告\n\n## 执行摘要\n概述。\n\n## 核心判断\n1. 判断一：形势严峻值得关注。\n",
    )
    bundle = report_agent.get_project_bundle(pid)

    assert "newspaper" in bundle
    assert bundle["newspaper"]["active"] is True
    assert "执行摘要" in bundle["newspaper"]["toc"]
    assert bundle["newspaper"]["reading_minutes"] >= 1


@pytest.mark.skipif(not report_agent.DOCX_AVAILABLE, reason="python-docx 未安装")
def test_docx_masthead_only_for_newspaper_voice(agent_db):
    from docx import Document

    def _all_text(buf):
        d = Document(buf)
        parts = [p.text for p in d.paragraphs]
        for t in d.tables:
            for row in t.rows:
                for cell in row.cells:
                    parts.append(cell.text)
        return "\n".join(parts)

    content = (
        "# 台海报告\n\n## 执行摘要\n概述若干。\n\n## 核心判断\n"
        "1. 判断一：区域态势严峻值得警惕。\n2. 判断二：联合体系加速成型。\n\n## 来源附录\n略。\n"
    )
    # newspaper voice → title page 带 masthead + 判断卡
    np = report_agent.create_project("台海报纸式", "strategic", topic="台海", voice="newspaper")
    report_agent.save_draft(np["project_id"], "draft", content)
    np_draft = report_agent.get_project_drafts(np["project_id"])[0]
    np_text = _all_text(report_agent.build_report_docx(np, np_draft))
    assert "VOL." in np_text
    assert "预计阅读" in np_text
    assert "核心判断卡" in np_text

    # 默认 strategic voice → 无 masthead（不影响既有版式）
    st = report_agent.create_project("台海常规", "strategic", topic="台海")
    report_agent.save_draft(st["project_id"], "draft", content)
    st_draft = report_agent.get_project_drafts(st["project_id"])[0]
    st_text = _all_text(report_agent.build_report_docx(st, st_draft))
    assert "预计阅读" not in st_text
    assert "核心判断卡" not in st_text
