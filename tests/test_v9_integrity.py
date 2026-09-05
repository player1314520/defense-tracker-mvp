# -*- coding: utf-8 -*-
import base64
import hashlib
import os
import sqlite3
from io import BytesIO

import pytest
from flask import Flask


def _service_with_record(tmp_path):
    from v9.service import V9Service

    database_path = tmp_path / "v9.sqlite3"
    key_path = tmp_path / ".master"
    service = V9Service(database_path, key_path)
    context = service.get_or_create_personal_context()
    service.acknowledge_personal_recovery(context["organization_id"])
    record = service.create_record(
        context["organization_id"],
        context["user_id"],
        context["device_id"],
        "evidence",
        {"body": "existing encrypted evidence"},
    )
    return service, context, record, database_path, key_path


def _database_hash(path):
    return hashlib.sha256(path.read_bytes()).hexdigest()


def _locked_client(service):
    from v9.api import create_blueprint

    app = Flask(__name__)
    app.register_blueprint(create_blueprint(lambda: service))
    return app.test_client()


def _personal_headers(context):
    return {
        "X-V9-Context-Mode": "personal",
        "X-V9-Organization-ID": context["organization_id"],
    }


def _ready_publication(tmp_path, *, recalled=False):
    from v9.service import V9Service

    service = V9Service(tmp_path / "v9.sqlite3", tmp_path / ".master")
    context = service.get_or_create_personal_context()
    service.acknowledge_personal_recovery(context["organization_id"])
    evidence = service.archive_news_evidence(
        context,
        {
            "aid": "integrity-source",
            "title": "公开来源证据",
            "summary": "可追溯摘要",
            "source": "Source A",
            "link": "https://example.test/integrity-source",
            "date": "2026-08-31T08:00:00+00:00",
        },
    )
    document = service.create_document(
        context,
        {
            "kind": "report",
            "title": "本地签发完整性报告",
            "paragraphs": [
                {
                    "heading": "判断",
                    "text": "公开材料支持本地判断。",
                    "evidence_ids": [evidence["record_id"]],
                    "source_status": "source_claim",
                    "fact_check": "passed",
                }
            ],
        },
    )
    publication = service.create_publication_item(
        context, document["record_id"]
    )
    publication = service.move_publication_item(
        context,
        publication["record_id"],
        expected_version=publication["version"],
        status="pending_approval",
    )
    if recalled:
        publication = service.sign_publication_item(
            context,
            publication["record_id"],
            expected_version=publication["version"],
        )
        publication = service.recall_publication_item(
            context,
            publication["record_id"],
            expected_version=publication["version"],
            reason="来源状态需要重新确认",
        )
    return service, context, publication


def _workflow_rows(service):
    with sqlite3.connect(service.database_path) as conn:
        records = conn.execute(
            """
            SELECT record_id,record_type,version,content_hash
            FROM encrypted_records ORDER BY record_id
            """
        ).fetchall()
        outbox = conn.execute(
            """
            SELECT event_id,record_id,operation,payload_json,state
            FROM sync_outbox ORDER BY event_id
            """
        ).fetchall()
    return records, outbox


def _install_crash_trigger(database_path, stage):
    triggers = {
        "publication_record": """
            CREATE TRIGGER crash_publication_record
            BEFORE UPDATE ON encrypted_records
            WHEN OLD.record_type='publication_item'
            BEGIN SELECT RAISE(ABORT, 'crash publication record'); END
        """,
        "publication_outbox": """
            CREATE TRIGGER crash_publication_outbox
            BEFORE INSERT ON sync_outbox
            WHEN instr(NEW.payload_json, '\"record_type\": \"publication_item\"') > 0
            BEGIN SELECT RAISE(ABORT, 'crash publication outbox'); END
        """,
        "audit_record": """
            CREATE TRIGGER crash_audit_record
            BEFORE INSERT ON encrypted_records
            WHEN NEW.record_type='audit_event'
            BEGIN SELECT RAISE(ABORT, 'crash audit record'); END
        """,
        "audit_outbox": """
            CREATE TRIGGER crash_audit_outbox
            BEFORE INSERT ON sync_outbox
            WHEN instr(NEW.payload_json, '\"record_type\": \"audit_event\"') > 0
            BEGIN SELECT RAISE(ABORT, 'crash audit outbox'); END
        """,
    }
    with sqlite3.connect(database_path) as conn:
        conn.executescript(triggers[stage])


def test_empty_database_can_create_a_new_master_key(tmp_path):
    from v9.service import V9Service

    key_path = tmp_path / ".master"
    service = V9Service(tmp_path / "v9.sqlite3", key_path)

    assert service.is_key_locked is False
    assert key_path.is_file()


def test_missing_master_key_locks_existing_database_without_mutation(tmp_path):
    from v9.service import V9KeyLocked, V9Service

    service, context, _, database_path, key_path = _service_with_record(tmp_path)
    key_path.unlink()
    before = _database_hash(database_path)

    locked = V9Service(database_path, key_path)

    assert locked.is_key_locked is True
    assert locked.key_status()["code"] == "V9_KEY_LOCKED"
    assert not key_path.exists()
    assert _database_hash(database_path) == before
    with pytest.raises(V9KeyLocked):
        locked.create_record(
            context["organization_id"],
            context["user_id"],
            context["device_id"],
            "evidence",
            {"body": "must not be written"},
        )
    assert _database_hash(database_path) == before


def test_wrong_master_key_locks_then_correct_key_restores_old_records(tmp_path):
    from v9.service import V9KeyLocked, V9Service

    service, context, record, database_path, key_path = _service_with_record(
        tmp_path
    )
    correct_key = bytes(service._master_key)
    wrong_payload = os.urandom(32)
    key_path.write_bytes(wrong_payload)
    before = _database_hash(database_path)

    locked = V9Service(database_path, key_path)

    assert locked.is_key_locked is True
    assert key_path.read_bytes() == wrong_payload
    assert _database_hash(database_path) == before
    with pytest.raises(V9KeyLocked):
        locked.restore_local_master_key(os.urandom(32))
    assert key_path.read_bytes() == wrong_payload
    assert _database_hash(database_path) == before

    result = locked.restore_local_master_key(correct_key)
    restored = locked.read_record(
        context["organization_id"], context["user_id"], record["record_id"]
    )

    assert result["status"] == "ready"
    assert locked.is_key_locked is False
    assert restored["content"]["body"] == "existing encrypted evidence"


@pytest.mark.skipif(os.name != "nt", reason="Windows DPAPI is required")
def test_dpapi_unprotect_failure_enters_lock_without_database_mutation(
    tmp_path, monkeypatch
):
    import v9.service as service_module

    _, _, _, database_path, key_path = _service_with_record(tmp_path)
    key_payload = key_path.read_bytes()
    before = _database_hash(database_path)

    def fail_unprotect(_payload):
        raise OSError("simulated DPAPI failure")

    monkeypatch.setattr(
        service_module, "unprotect_local_master_key", fail_unprotect
    )
    locked = service_module.V9Service(database_path, key_path)

    assert locked.is_key_locked is True
    assert key_path.read_bytes() == key_payload
    assert _database_hash(database_path) == before


def test_locked_api_returns_stable_423_and_allows_only_recovery_artifacts(
    tmp_path,
):
    from v9.service import V9Service

    service, context, _, database_path, key_path = _service_with_record(tmp_path)
    correct_key = bytes(service._master_key)
    key_path.unlink()
    locked = V9Service(database_path, key_path)
    client = _locked_client(locked)

    blocked_write = client.post("/api/v9/claims", json={})
    blocked_read = client.get("/api/v9/evidence")

    assert blocked_write.status_code == 423
    assert blocked_read.status_code == 423
    payload = blocked_write.get_json()
    assert payload == {
        "error": "本地数据密钥不可用",
        "code": "V9_KEY_LOCKED",
        "request_id": payload["request_id"],
        "retryable": False,
        "action": "restore_key_or_backup",
    }
    assert blocked_write.headers["X-Request-ID"] == payload["request_id"]

    diagnostics = client.get("/api/v9/diagnostics/export")
    backup = client.post("/api/v9/backups", json={})
    assert diagnostics.status_code == 200
    assert backup.status_code == 201

    restored = client.post(
        "/api/v9/recovery/local-master-key",
        json={
            "key_payload_base64": base64.b64encode(correct_key).decode("ascii")
        },
    )
    assert restored.status_code == 200
    assert restored.get_json() == {"status": "ready", "code": "V9_KEY_RESTORED"}
    assert client.get(
        "/api/v9/evidence", headers=_personal_headers(context)
    ).status_code == 200


def test_local_key_recovery_rejects_noncanonical_base64_alias(tmp_path):
    from v9.service import V9Service

    service, _, _, database_path, key_path = _service_with_record(tmp_path)
    correct_key = bytes(service._master_key)
    key_path.unlink()
    locked = V9Service(database_path, key_path)
    encoded = base64.b64encode(correct_key).decode("ascii")
    alphabet = "ABCDEFGHIJKLMNOPQRSTUVWXYZabcdefghijklmnopqrstuvwxyz0123456789+/"
    alias = next(
        encoded[:-2] + candidate + "="
        for candidate in alphabet
        if candidate != encoded[-2]
        and base64.b64decode(
            encoded[:-2] + candidate + "=", validate=True
        ) == correct_key
    )

    response = _locked_client(locked).post(
        "/api/v9/recovery/local-master-key",
        json={"key_payload_base64": alias},
    )

    assert response.status_code == 400
    assert locked.is_key_locked is True
    assert not key_path.exists()


@pytest.mark.parametrize(
    "failure_stage",
    [
        "publication_record",
        "publication_outbox",
        "audit_record",
        "audit_outbox",
    ],
)
def test_publication_state_audit_and_outbox_rollback_together(
    tmp_path, failure_stage
):
    service, context, publication = _ready_publication(tmp_path)

    before = _workflow_rows(service)
    _install_crash_trigger(service.database_path, failure_stage)

    operation = lambda: service.sign_publication_item(
        context,
        publication["record_id"],
        expected_version=publication["version"],
    )
    with pytest.raises(sqlite3.DatabaseError, match="crash"):
        operation()

    assert _workflow_rows(service) == before


@pytest.mark.parametrize(
    "failure_stage",
    [
        "publication_record",
        "publication_outbox",
        "audit_record",
        "audit_outbox",
    ],
)
def test_recall_state_audit_and_outbox_rollback_together(
    tmp_path, failure_stage
):
    service, context, publication = _ready_publication(tmp_path)
    publication = service.sign_publication_item(
        context,
        publication["record_id"],
        expected_version=publication["version"],
    )
    before = _workflow_rows(service)
    _install_crash_trigger(service.database_path, failure_stage)

    with pytest.raises(sqlite3.DatabaseError, match="crash"):
        service.recall_publication_item(
            context,
            publication["record_id"],
            expected_version=publication["version"],
            reason="崩溃注入测试",
        )

    assert _workflow_rows(service) == before


def test_recalled_publication_requires_audit_mode_and_contains_recall_marks(
    tmp_path,
):
    from docx import Document
    from v9.publication import PublicationRecalled

    service, context, publication = _ready_publication(tmp_path, recalled=True)

    with pytest.raises(PublicationRecalled):
        service.export_publication(
            context, publication["record_id"], "docx"
        )

    payload, filename = service.export_publication(
        context,
        publication["record_id"],
        "docx",
        mode="audit",
    )
    document = Document(BytesIO(payload))
    body = "\n".join(paragraph.text for paragraph in document.paragraphs)
    headers = "\n".join(
        paragraph.text
        for section in document.sections
        for paragraph in section.header.paragraphs
    )

    assert "已撤回" in headers
    assert "撤回回执" in body
    assert "来源状态需要重新确认" in body
    assert "本地批准并冻结快照" in body
    assert "不代表已向外部渠道发布" in body
    assert "已撤回-审计件" in filename


def test_recalled_pdf_audit_artifact_is_marked_in_metadata_and_visible_text(
    tmp_path,
):
    import pdfplumber

    service, context, publication = _ready_publication(tmp_path, recalled=True)

    payload, filename = service.export_publication(
        context,
        publication["record_id"],
        "pdf",
        mode="audit",
    )
    with pdfplumber.open(BytesIO(payload)) as pdf:
        title = pdf.metadata.get("Title") or ""
        text = "\n".join(page.extract_text() or "" for page in pdf.pages)

    assert "已撤回-审计件" in title
    assert "已撤回" in text
    assert "撤回回执" in text
    assert "来源状态需要重新确认" in text
    assert "已撤回-审计件" in filename


def test_recalled_publication_api_returns_409_or_explicit_audit_artifact(
    tmp_path,
):
    from docx import Document
    from v9.api import create_blueprint

    service, context, publication = _ready_publication(tmp_path, recalled=True)
    app = Flask(__name__)
    app.register_blueprint(create_blueprint(lambda: service))
    client = app.test_client()
    url = f"/api/v9/publications/{publication['record_id']}/export.docx"

    blocked = client.get(
        url,
        headers=_personal_headers(context),
        query_string={"organization_id": context["organization_id"]},
    )
    assert blocked.status_code == 409
    assert blocked.get_json()["code"] == "PUBLICATION_RECALLED"
    assert blocked.headers["X-Request-ID"] == blocked.get_json()["request_id"]

    audit = client.get(
        url,
        headers=_personal_headers(context),
        query_string={
            "organization_id": context["organization_id"],
            "mode": "audit",
        },
    )
    assert audit.status_code == 200
    assert audit.mimetype == (
        "application/vnd.openxmlformats-officedocument."
        "wordprocessingml.document"
    )
    parsed = Document(BytesIO(audit.data))
    headers = "\n".join(
        paragraph.text
        for section in parsed.sections
        for paragraph in section.header.paragraphs
    )
    assert "已撤回" in headers


def test_recalled_document_cannot_bypass_publication_export_gate(tmp_path):
    from docx import Document
    from v9.publication import PublicationRecalled

    service, context, _publication = _ready_publication(tmp_path, recalled=True)
    recalled = service.list_publication_items(context)[0]
    document_id = recalled["content"]["signed_snapshot"]["receipt"]["document_id"]

    with pytest.raises(PublicationRecalled):
        service.export_document(context, document_id, "docx")

    payload, filename = service.export_document(
        context, document_id, "docx", mode="audit"
    )
    parsed = Document(BytesIO(payload))
    headers = "\n".join(
        paragraph.text
        for section in parsed.sections
        for paragraph in section.header.paragraphs
    )
    assert "已撤回" in headers
    assert "已撤回-审计件" in filename

    from v9.api import create_blueprint

    app = Flask(__name__)
    app.register_blueprint(create_blueprint(lambda: service))
    client = app.test_client()
    url = f"/api/v9/documents/{document_id}/export.docx"
    blocked = client.get(
        url,
        headers=_personal_headers(context),
        query_string={"organization_id": context["organization_id"]},
    )
    assert blocked.status_code == 409
    assert blocked.get_json()["code"] == "PUBLICATION_RECALLED"
    audit = client.get(
        url,
        headers=_personal_headers(context),
        query_string={
            "organization_id": context["organization_id"],
            "mode": "audit",
        },
    )
    assert audit.status_code == 200


def test_signed_state_is_explicitly_local_approval_not_external_publication(
    tmp_path,
):
    service, context, publication = _ready_publication(tmp_path)

    signed = service.sign_publication_item(
        context,
        publication["record_id"],
        expected_version=publication["version"],
    )
    stored = service.list_publication_items(context)[0]["content"]
    semantics = stored["signed_snapshot"]["publication_semantics"]

    assert signed["publication_scope"] == "local_approval_snapshot"
    assert signed["external_published"] is False
    assert semantics["scope"] == "local_approval_snapshot"
    assert semantics["external_published"] is False
    assert "本地批准并冻结快照" in semantics["notice"]
    assert "不代表已向外部渠道发布" in semantics["notice"]
