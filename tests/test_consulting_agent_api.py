from datetime import datetime, timezone
from io import BytesIO
import json
from pathlib import Path
import sqlite3
from zipfile import ZipFile

import pytest

import app as tracker
import consulting_agent
import report_agent

if report_agent.DOCX_AVAILABLE:
    from docx import Document


def _login_cookies(client, csrf="csrf-test-token"):
    client.set_cookie(tracker.AUTH_COOKIE, tracker.ACCESS_TOKEN)
    client.set_cookie(tracker.CSRF_COOKIE, csrf)
    return csrf


def _rss_candidate(idx=1):
    return {
        "article_id": f"rss-{idx}",
        "title": f"台海无人系统动态 {idx}",
        "summary": "公开报道显示，无人系统与跨域感知能力持续发展。",
        "source": "Defense News",
        "source_cn": "Defense News",
        "link": f"https://example.test/rss-{idx}",
        "date": datetime.now(timezone.utc).isoformat(),
        "quality_score": 82,
        "quality_level": "A",
        "quality_reasons": ["贴近防务战略分析主题"],
    }


def test_consult_session_create_requires_csrf_and_returns_capabilities(monkeypatch, tmp_path):
    monkeypatch.setattr(consulting_agent, "CONSULTING_AGENT_DB_FILE", str(tmp_path / "consult.sqlite3"))
    monkeypatch.setattr(tracker.search_adapters, "search_status", lambda: {
        "provider": "tavily",
        "web_search_enabled": False,
        "message": "实时网页搜索未启用：TAVILY_API_KEY 未配置",
    })
    tracker.app.config["TESTING"] = True
    client = tracker.app.test_client()
    csrf = _login_cookies(client)

    blocked = client.post("/api/consult/sessions", json={"instruction": "帮我做一个台海报告"})
    assert blocked.status_code == 403

    ok = client.post(
        "/api/consult/sessions",
        json={"instruction": "帮我做一个有关台海无人系统发展的战略分析报告，搜集20份来源"},
        headers={tracker.CSRF_HEADER: csrf},
    )
    data = ok.get_json()

    assert ok.status_code == 200
    assert data["session"]["topic"] == "台海无人系统发展"
    assert data["session"]["target_source_count"] == 20
    assert data["capabilities"]["web_search_enabled"] is False
    assert data["model"]["model"] == tracker.AI_CONFIG["model"]
    assert "抓取公开报告" in data["capability_notice"]


def test_consult_search_aggregates_web_rss_and_thinktanks(monkeypatch, tmp_path):
    monkeypatch.setattr(consulting_agent, "CONSULTING_AGENT_DB_FILE", str(tmp_path / "consult.sqlite3"))
    monkeypatch.setattr(tracker.search_adapters, "search_status", lambda: {
        "provider": "tavily",
        "web_search_enabled": True,
        "message": "实时网页搜索已启用",
    })
    monkeypatch.setattr(tracker.search_adapters, "search_web", lambda query, limit, **kwargs: [
        {
            "title": f"Web source {i}",
            "source": "Web",
            "url": f"https://example.test/web-{i}",
            "snippet": "公开网页资料",
            "channel": "web",
            "score": 75,
        }
        for i in range(8)
    ])
    monkeypatch.setattr(tracker, "select_quality_candidates", lambda **kwargs: (
        [_rss_candidate(i) for i in range(5)],
        {"total_scored": 5},
    ))
    monkeypatch.setattr(tracker, "THINK_TANK_DIRECTORY", [{
        "category": "PLA专项研究机构",
        "sites": [
            {
                "name": f"Think Tank {i}",
                "name_cn": f"智库{i}",
                "url": f"https://example.test/tank-{i}",
                "desc_cn": "台海无人系统长期研究报告",
            }
            for i in range(10)
        ],
    }])
    tracker.app.config["TESTING"] = True
    client = tracker.app.test_client()
    csrf = _login_cookies(client)
    created = client.post(
        "/api/consult/sessions",
        json={"instruction": "帮我做一个台海无人系统报告，搜集20份来源"},
        headers={tracker.CSRF_HEADER: csrf},
    ).get_json()["session"]

    resp = client.post(
        f"/api/consult/sessions/{created['session_id']}/search",
        json={},
        headers={tracker.CSRF_HEADER: csrf},
    )
    data = resp.get_json()

    assert resp.status_code == 200
    assert data["total"] == 20
    assert data["meta"]["target_source_count"] == 20
    assert data["meta"]["found_count"] == 13
    assert data["meta"]["target_seed_count"] == 7
    assert data["meta"]["shortfall"] == 7
    channels = {ev["channel"] for ev in data["evidence"]}
    assert {"web", "rss", "thinktank_target"} <= channels


def test_consult_search_honors_target_count_payload(monkeypatch, tmp_path):
    monkeypatch.setattr(consulting_agent, "CONSULTING_AGENT_DB_FILE", str(tmp_path / "consult.sqlite3"))
    monkeypatch.setattr(tracker.search_adapters, "search_status", lambda: {
        "provider": "public_web",
        "web_search_enabled": True,
        "message": "基础联网搜索已启用",
    })
    seen = {}

    def fake_search(query, limit, **kwargs):
        seen["limit"] = limit
        return [
            {
                "title": f"Iran missile source {i}",
                "source": "Web",
                "url": f"https://example.test/iran-{i}.pdf",
                "snippet": "伊朗导弹公开报告资料",
                "channel": "web",
                "score": 80,
            }
            for i in range(limit)
        ]

    monkeypatch.setattr(tracker.search_adapters, "search_web", fake_search)
    monkeypatch.setattr(tracker, "select_quality_candidates", lambda **kwargs: ([], {"total_scored": 0}))
    monkeypatch.setattr(tracker, "THINK_TANK_DIRECTORY", [])
    tracker.app.config["TESTING"] = True
    client = tracker.app.test_client()
    csrf = _login_cookies(client)
    created = client.post(
        "/api/consult/sessions",
        json={"instruction": "伊朗导弹战略分析报告，搜集20份来源"},
        headers={tracker.CSRF_HEADER: csrf},
    ).get_json()["session"]

    resp = client.post(
        f"/api/consult/sessions/{created['session_id']}/search",
        json={"target_count": 6},
        headers={tracker.CSRF_HEADER: csrf},
    )
    data = resp.get_json()

    assert resp.status_code == 200
    assert seen["limit"] == 6
    assert data["total"] == 6
    assert data["meta"]["target_source_count"] == 6


def test_consult_source_pack_does_not_require_ai_and_records_package(monkeypatch, tmp_path):
    monkeypatch.setattr(consulting_agent, "CONSULTING_AGENT_DB_FILE", str(tmp_path / "consult.sqlite3"))
    old_config = dict(tracker.AI_CONFIG)
    tracker.AI_CONFIG["api_key"] = ""
    tracker.app.config["TESTING"] = True
    client = tracker.app.test_client()
    csrf = _login_cookies(client)
    session = consulting_agent.create_session("帮我做一个台海无人系统报告，搜集2份智库分析")
    evidence = consulting_agent.upsert_evidence(session["session_id"], [{
        "title": "台海无人系统研究",
        "source": "RAND",
        "url": "https://example.test/report",
        "channel": "web",
        "score": 91,
        "snippet": "公开源资料",
    }])

    try:
        resp = client.post(
            f"/api/consult/sessions/{session['session_id']}/source_pack",
            json={"evidence_ids": [evidence[0]["evidence_id"]]},
            headers={tracker.CSRF_HEADER: csrf},
        )
    finally:
        tracker.AI_CONFIG.clear()
        tracker.AI_CONFIG.update(old_config)

    data = resp.get_json()

    assert resp.status_code == 200
    assert data["answer"]["kind"] == "source_pack"
    assert "报告源抓取包" in data["answer"]["content"]
    assert "不是最终战略分析报告" in data["answer"]["content"]
    assert data["answer"]["model"] == "deterministic-source-pack"


def test_consult_synthesize_requires_ai_config(monkeypatch, tmp_path):
    monkeypatch.setattr(consulting_agent, "CONSULTING_AGENT_DB_FILE", str(tmp_path / "consult.sqlite3"))
    old_config = dict(tracker.AI_CONFIG)
    tracker.AI_CONFIG["api_key"] = ""
    tracker.app.config["TESTING"] = True
    client = tracker.app.test_client()
    csrf = _login_cookies(client)
    session = consulting_agent.create_session("帮我做一个台海无人系统报告，搜集1份来源")

    try:
        resp = client.post(
            f"/api/consult/sessions/{session['session_id']}/synthesize",
            json={"evidence_ids": []},
            headers={tracker.CSRF_HEADER: csrf},
        )
    finally:
        tracker.AI_CONFIG.clear()
        tracker.AI_CONFIG.update(old_config)

    assert resp.status_code == 400
    assert "AI API Key" in resp.get_json()["error"]


def test_consult_synthesize_uses_mocked_ai_and_records_answer(monkeypatch, tmp_path):
    monkeypatch.setattr(consulting_agent, "CONSULTING_AGENT_DB_FILE", str(tmp_path / "consult.sqlite3"))
    old_config = dict(tracker.AI_CONFIG)
    tracker.AI_CONFIG.update({"api_key": "unit-key", "model": "unit-model"})
    calls = []

    def fake_call_ai(messages, temperature=None, max_tokens=None):
        calls.append({"messages": messages, "temperature": temperature, "max_tokens": max_tokens})
        return "## 目录提纲\n一、态势\n\n## 审稿意见\n来源闭环。\n\n## 综合报告\n基于[1]判断台海无人系统值得跟踪。\n\n## 来源附录\n[1] RAND"

    monkeypatch.setattr(tracker, "_call_ai", fake_call_ai)
    tracker.app.config["TESTING"] = True
    client = tracker.app.test_client()
    csrf = _login_cookies(client)
    session = consulting_agent.create_session("帮我做一个台海无人系统报告，搜集1份来源")
    evidence = consulting_agent.upsert_evidence(session["session_id"], [{
        "title": "台海无人系统研究",
        "source": "RAND",
        "url": "https://example.test/report",
        "channel": "web",
        "score": 91,
        "snippet": "公开源资料",
    }])

    try:
        resp = client.post(
            f"/api/consult/sessions/{session['session_id']}/synthesize",
            json={"evidence_ids": [evidence[0]["evidence_id"]], "writing_requirements": "突出审稿意见"},
            headers={tracker.CSRF_HEADER: csrf},
        )
    finally:
        tracker.AI_CONFIG.clear()
        tracker.AI_CONFIG.update(old_config)

    data = resp.get_json()
    prompt = "\n".join(m["content"] for m in calls[0]["messages"])

    assert resp.status_code == 200
    assert calls[0]["temperature"] == 0.35
    assert "每个关键判断必须绑定来源编号" in prompt
    assert data["answer"]["model"] == "unit-model"
    assert "秘密" not in data["answer"]["content"]
    assert "来源附录" in data["answer"]["content"]


def test_consult_revise_records_revision(monkeypatch, tmp_path):
    monkeypatch.setattr(consulting_agent, "CONSULTING_AGENT_DB_FILE", str(tmp_path / "consult.sqlite3"))
    old_config = dict(tracker.AI_CONFIG)
    tracker.AI_CONFIG.update({"api_key": "unit-key", "model": "unit-model"})
    monkeypatch.setattr(
        tracker,
        "_call_ai",
        lambda messages, temperature=None, max_tokens=None: "## 综合报告\n修订后基于[1]强化影响研判。",
    )
    tracker.app.config["TESTING"] = True
    client = tracker.app.test_client()
    csrf = _login_cookies(client)
    session = consulting_agent.create_session("帮我做一个台海无人系统报告")
    answer = consulting_agent.save_answer(
        session["session_id"],
        "## 综合报告\n原始报告",
        model="unit-model",
        kind="synthesis",
    )

    try:
        resp = client.post(
            f"/api/consult/sessions/{session['session_id']}/revise",
            json={"answer_id": answer["answer_id"], "instruction": "加强影响研判"},
            headers={tracker.CSRF_HEADER: csrf},
        )
    finally:
        tracker.AI_CONFIG.clear()
        tracker.AI_CONFIG.update(old_config)

    data = resp.get_json()

    assert resp.status_code == 200
    assert data["answer"]["kind"] == "revision"
    assert "修订后" in data["answer"]["content"]


def test_search_status_and_config_routes_mask_keys(monkeypatch, tmp_path):
    if not tracker.CRYPTO_AVAILABLE:
        pytest.skip("cryptography is unavailable")
    monkeypatch.setenv(
        "AI_CONFIG_FERNET_KEY",
        tracker.Fernet.generate_key().decode("ascii"),
    )
    monkeypatch.setattr(tracker, "_AI_CIPHER", None)
    monkeypatch.setattr(tracker, "SEARCH_CONFIG_FILE", str(tmp_path / "search_config.json"), raising=False)
    tracker.SEARCH_CONFIG.clear()
    tracker.SEARCH_CONFIG.update({})
    tracker.app.config["TESTING"] = True
    client = tracker.app.test_client()
    csrf = _login_cookies(client)

    status_resp = client.get("/api/search/status")
    status = status_resp.get_json()["status"]

    assert status_resp.status_code == 200
    assert status["online_search_enabled"] is True
    assert status["providers"]["public_web"]["enabled"] is True

    save_resp = client.post(
        "/api/search/config",
        json={"tavily_api_key": "tv-secret", "brave_api_key": "brave-secret"},
        headers={tracker.CSRF_HEADER: csrf},
    )
    saved = save_resp.get_json()["status"]

    assert save_resp.status_code == 200
    assert saved["online_search_enabled"] is True
    assert saved["providers"]["tavily"]["enabled"] is True
    assert saved["providers"]["brave"]["enabled"] is True
    assert "tv-secret" not in str(save_resp.get_json())
    assert "brave-secret" not in str(save_resp.get_json())


def test_consult_source_pack_exports_docx(monkeypatch, tmp_path):
    monkeypatch.setattr(consulting_agent, "CONSULTING_AGENT_DB_FILE", str(tmp_path / "consult.sqlite3"))
    monkeypatch.setattr(consulting_agent, "SOURCE_ARCHIVE_DIR", str(tmp_path / "source_archive"), raising=False)
    tracker.app.config["TESTING"] = True
    client = tracker.app.test_client()
    csrf = _login_cookies(client)
    session = consulting_agent.create_session("搜集2份台海无人系统智库报告")
    evidence = consulting_agent.upsert_evidence(session["session_id"], [{
        "title": "RAND 台海无人系统报告",
        "source": "RAND",
        "url": "https://example.test/rand-report",
        "channel": "web",
        "score": 91,
        "snippet": "公开报告正文摘要。",
        "payload": {"is_fetched_original": True, "document_type": "html"},
    }])
    asset = consulting_agent.archive_source_asset(session["session_id"], evidence[0], {
        "title": "RAND 台海无人系统报告",
        "url": "https://example.test/rand-report",
        "text": "公开报告正文摘要。完整正文用于资料包归档。",
        "snippet": "公开报告正文摘要。",
        "document_type": "html",
        "content_type": "text/html",
        "word_count": 20,
        "raw_bytes": b"<html><body>source</body></html>",
        "is_fetched_original": True,
    })
    evidence = consulting_agent.upsert_evidence(session["session_id"], [{
        **evidence[0],
        "payload": {**(evidence[0].get("payload") or {}), "asset_status": "archived", "asset_id": asset["asset_id"]},
    }])
    answer = consulting_agent.save_answer(
        session["session_id"],
        consulting_agent.build_source_pack(session, evidence),
        model="deterministic-source-pack",
        kind="source_pack",
        payload={"evidence_ids": [evidence[0]["evidence_id"]]},
    )

    resp = client.post(
        f"/api/consult/sessions/{session['session_id']}/export_source_pack",
        json={"answer_id": answer["answer_id"]},
        headers={tracker.CSRF_HEADER: csrf},
    )

    assert resp.status_code == 200
    assert resp.mimetype == "application/zip"
    assert ".zip" in resp.headers.get("Content-Disposition", "")
    assert resp.data[:2] == b"PK"
    with ZipFile(BytesIO(resp.data)) as zf:
        names = zf.namelist()
        assert any(name.endswith(".docx") for name in names)
        assert any(name.endswith(".pdf") for name in names)
        assert "manifest.json" in names
        assert any(name.startswith("sources/") and name.endswith(".txt") for name in names)
        manifest_text = zf.read("manifest.json").decode("utf-8")
        manifest = json.loads(manifest_text)
        assert manifest["session_id"] == session["session_id"]
        assert manifest["assets"][0]["status"] == "archived"
        assert "source_archive_path" not in manifest
        assert "local_path" not in manifest["assets"][0]
        assert "text_path" not in manifest["assets"][0]
        assert manifest["assets"][0]["archive_members"]["text"] in names
        assert str(tmp_path) not in manifest_text
        assert "failures.json" in names
        assert "needs_user_input.json" in names
        docx_name = next(name for name in names if name.endswith(".docx"))
        doc = Document(BytesIO(zf.read(docx_name)))
        texts = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
        assert texts[0].endswith("报告源资料包")
        assert "台海无人系统" in texts[0]
        assert any("公开源报告抓取与证据资料包" in text for text in texts)
        assert "目          录" in texts
        assert len(doc.tables) >= 1
        assert doc.tables[0].cell(0, 0).text == "客户指令"
        pdf_name = next(name for name in names if name.endswith(".pdf"))
        assert zf.read(pdf_name).startswith(b"%PDF")


def test_consult_asset_preview_returns_text_and_local_file(monkeypatch, tmp_path):
    monkeypatch.setattr(consulting_agent, "CONSULTING_AGENT_DB_FILE", str(tmp_path / "consult.sqlite3"))
    monkeypatch.setattr(consulting_agent, "SOURCE_ARCHIVE_DIR", str(tmp_path / "source_archive"), raising=False)
    tracker.app.config["TESTING"] = True
    client = tracker.app.test_client()
    _login_cookies(client)
    session = consulting_agent.create_session("搜集1份欧洲防空生产能力报告")
    evidence = consulting_agent.upsert_evidence(session["session_id"], [{
        "title": "European air defense production report",
        "source": "Public Institute",
        "url": "https://example.test/europe-air-defense.html",
        "channel": "web",
        "score": 88,
        "snippet": "Air defense production source.",
    }])
    asset = consulting_agent.archive_source_asset(session["session_id"], evidence[0], {
        "title": "European air defense production report",
        "url": "https://example.test/europe-air-defense.html",
        "text": "European air defense production capacity and missile stockpile replenishment.",
        "snippet": "Air defense production source.",
        "document_type": "html",
        "content_type": "text/html",
        "word_count": 12,
        "raw_bytes": b"<html><body>European air defense production capacity</body></html>",
        "is_fetched_original": True,
    })

    preview = client.get(
        f"/api/consult/sessions/{session['session_id']}/assets/{asset['asset_id']}/preview"
    )
    data = preview.get_json()

    assert preview.status_code == 200
    assert data["preview_mode"] == "document"
    assert "European air defense production" in data["text"]
    assert data["file_url"].endswith(f"/assets/{asset['asset_id']}/file")

    file_resp = client.get(data["file_url"])
    assert file_resp.status_code == 200
    assert b"European air defense production capacity" in file_resp.data


def test_consult_asset_preview_does_not_iframe_fake_pdf_placeholder(monkeypatch, tmp_path):
    monkeypatch.setattr(consulting_agent, "CONSULTING_AGENT_DB_FILE", str(tmp_path / "consult.sqlite3"))
    monkeypatch.setattr(consulting_agent, "SOURCE_ARCHIVE_DIR", str(tmp_path / "source_archive"), raising=False)
    tracker.app.config["TESTING"] = True
    client = tracker.app.test_client()
    _login_cookies(client)
    session = consulting_agent.create_session("搜集1份受限PDF报告")
    evidence = consulting_agent.upsert_evidence(session["session_id"], [{
        "title": "Restricted PDF",
        "source": "Official",
        "url": "https://example.test/restricted.pdf",
        "channel": "web",
        "score": 88,
        "snippet": "Login required.",
    }])
    private_marker = "credential=[private-value] at [private-path]"
    asset = consulting_agent.archive_source_asset(session["session_id"], evidence[0], {
        "title": "Restricted PDF",
        "url": "https://example.test/restricted.pdf",
        "text": private_marker,
        "snippet": private_marker,
        "document_type": "pdf",
        "content_type": "application/pdf",
        "word_count": 0,
        "raw_bytes": private_marker.encode("utf-8"),
        "is_fetched_original": True,
        "failure_code": "blocked",
        "diagnosis": {
            "code": "blocked",
            "label": private_marker,
            "advice": {"private": private_marker},
        },
    }, status="needs_user_input", failure_reason=private_marker)

    with sqlite3.connect(tmp_path / "consult.sqlite3") as conn:
        stored_paths = conn.execute(
            """
            SELECT local_path, text_path, metadata_path
            FROM source_assets WHERE asset_id=?
            """,
            (asset["asset_id"],),
        ).fetchone()
    assert stored_paths is not None
    for stored_path in stored_paths:
        content = Path(stored_path).read_bytes()
        assert private_marker.encode("utf-8") not in content
    assert asset["local_path"] == ""
    assert asset["text_path"] == ""
    assert asset["metadata_path"] == ""
    assert asset["checksum"] == ""

    preview = client.get(
        f"/api/consult/sessions/{session['session_id']}/assets/{asset['asset_id']}/preview"
    )
    data = preview.get_json()

    assert preview.status_code == 200
    assert data["preview_mode"] == "document"
    assert data["reader_mode"] == "blocked"
    assert data["file_is_real_pdf"] is False
    assert "用户授权" in data["text"]
    assert private_marker not in json.dumps(data, ensure_ascii=False)
    assert data["file_url"] == ""
    assert data["download_url"] == ""

    file_resp = client.get(
        f"/api/consult/sessions/{session['session_id']}/assets/{asset['asset_id']}/file"
    )
    assert file_resp.status_code == 409
    assert private_marker not in file_resp.get_data(as_text=True)


def test_consult_real_pdf_file_allows_same_origin_reader_iframe(monkeypatch, tmp_path):
    monkeypatch.setattr(consulting_agent, "CONSULTING_AGENT_DB_FILE", str(tmp_path / "consult.sqlite3"))
    monkeypatch.setattr(consulting_agent, "SOURCE_ARCHIVE_DIR", str(tmp_path / "source_archive"), raising=False)
    tracker.app.config["TESTING"] = True
    client = tracker.app.test_client()
    _login_cookies(client)
    session = consulting_agent.create_session("搜集1份真实PDF报告")
    evidence = consulting_agent.upsert_evidence(session["session_id"], [{
        "title": "Real PDF",
        "source": "Official",
        "url": "https://example.test/real.pdf",
        "channel": "web",
        "score": 90,
        "snippet": "Public PDF.",
    }])
    asset = consulting_agent.archive_source_asset(session["session_id"], evidence[0], {
        "title": "Real PDF",
        "url": "https://example.test/real.pdf",
        "text": "Real PDF extracted text.",
        "snippet": "Public PDF.",
        "document_type": "pdf",
        "content_type": "application/pdf",
        "word_count": 4,
        "raw_bytes": b"%PDF-1.6\n1 0 obj\n<<>>\nendobj\n%%EOF",
        "is_fetched_original": True,
    })

    preview = client.get(
        f"/api/consult/sessions/{session['session_id']}/assets/{asset['asset_id']}/preview"
    )
    data = preview.get_json()
    file_resp = client.get(data["file_url"])
    csp = file_resp.headers.get("Content-Security-Policy", "")

    assert preview.status_code == 200
    assert data["preview_mode"] == "pdf"
    assert data["reader_mode"] == "pdf"
    assert data["file_is_real_pdf"] is True
    assert file_resp.status_code == 200
    assert file_resp.headers["Content-Type"].startswith("application/pdf")
    assert "frame-src 'self' blob:" in csp
    assert "frame-ancestors 'self'" in csp
    assert "frame-ancestors 'none'" not in csp


def test_consult_asset_open_rejects_symlink_outside_archive(
    monkeypatch, tmp_path
):
    archive = tmp_path / "archive"
    archive.mkdir()
    outside = tmp_path / "private.txt"
    outside.write_text("must-not-be-served", encoding="utf-8")
    link = archive / "linked.txt"
    try:
        link.symlink_to(outside)
    except OSError:
        pytest.skip("symlink creation is unavailable")
    monkeypatch.setattr(
        consulting_agent,
        "source_archive_root",
        lambda _session_id: str(archive),
    )

    with pytest.raises(FileNotFoundError):
        tracker._open_consult_asset("session-1", str(link))


def test_consult_asset_open_rejects_inode_swap(monkeypatch, tmp_path):
    archive = tmp_path / "archive"
    archive.mkdir()
    candidate = archive / "asset.txt"
    replacement = archive / "replacement.txt"
    candidate.write_text("approved", encoding="utf-8")
    replacement.write_text("must-not-be-served", encoding="utf-8")
    real_open = tracker.os.open
    swapped = False

    def racing_open(path, flags, *args, **kwargs):
        nonlocal swapped
        if Path(path) == candidate and not swapped:
            swapped = True
            candidate.rename(archive / "asset-original.txt")
            replacement.replace(candidate)
        return real_open(path, flags, *args, **kwargs)

    monkeypatch.setattr(
        consulting_agent,
        "source_archive_root",
        lambda _session_id: str(archive),
    )
    monkeypatch.setattr(tracker.os, "open", racing_open)

    with pytest.raises(FileNotFoundError):
        tracker._open_consult_asset("session-1", str(candidate))
    assert swapped


def test_consult_capture_to_target_archives_success_and_marks_restricted(monkeypatch, tmp_path):
    monkeypatch.setattr(consulting_agent, "CONSULTING_AGENT_DB_FILE", str(tmp_path / "consult.sqlite3"))
    monkeypatch.setattr(consulting_agent, "SOURCE_ARCHIVE_DIR", str(tmp_path / "source_archive"), raising=False)
    tracker.SEARCH_CONFIG.clear()

    def fake_multi(queries, target_count, providers=None, config=None, **kwargs):
        return (
            [
                {
                    "title": "Red Sea UAV report",
                    "source": "RUSI",
                    "url": "https://example.test/red-sea-uav",
                    "snippet": "Public report on Red Sea UAV attacks.",
                    "channel": "web",
                    "score": 92,
                    "provider": "public_web",
                    "query": queries[0],
                    "rank": 1,
                    "document_type": "html",
                },
                {
                    "title": "Restricted Red Sea UAV PDF",
                    "source": "Paywalled Journal",
                    "url": "https://example.test/restricted-red-sea-uav.pdf",
                    "snippet": "Login required.",
                    "channel": "web",
                    "score": 90,
                    "provider": "public_web",
                    "query": queries[0],
                    "rank": 2,
                    "document_type": "pdf",
                },
            ],
            {
                "target_count": target_count,
                "queries": queries,
                "provider_stats": {"public_web": {"returned": 2, "accepted": 2}},
                "provider_errors": {},
                "deduped_count": 2,
                "rejected_low_relevance": 0,
            },
        )

    def fake_extract(url, **kwargs):
        if "restricted" in url:
            raise RuntimeError("403 Forbidden login required")
        return {
            "title": "Red Sea UAV report",
            "url": url,
            "text": "Public report text about Red Sea UAV attacks and maritime security.",
            "snippet": "Public report text about Red Sea UAV attacks.",
            "document_type": "html",
            "content_type": "text/html",
            "word_count": 180,
            "raw_bytes": b"<html>red sea uav</html>",
            "is_fetched_original": True,
        }

    monkeypatch.setattr(tracker.search_adapters, "search_web_multi", fake_multi)
    monkeypatch.setattr(tracker.search_adapters, "extract_url", fake_extract)
    monkeypatch.setattr(tracker, "_is_ssrf_safe", lambda url: (True, ""))
    tracker.app.config["TESTING"] = True
    client = tracker.app.test_client()
    csrf = _login_cookies(client)
    session = consulting_agent.create_session("搜集2份红海无人机智库报告")

    resp = client.post(
        f"/api/consult/sessions/{session['session_id']}/capture_to_target",
        json={"target_count": 2, "batch_size": 2, "max_rounds": 1, "allow_browser_render": True},
        headers={tracker.CSRF_HEADER: csrf},
    )
    data = resp.get_json()

    assert resp.status_code == 200
    assert data["target_count"] == 2
    assert data["archived_count"] == 1
    assert data["needs_user_input_count"] == 1
    assert data["status"] == "completed"

    job_resp = client.get(f"/api/consult/sessions/{session['session_id']}/capture_jobs/{data['job_id']}")
    job = job_resp.get_json()["job"]
    assert job_resp.status_code == 200
    assert job["archived_count"] == 1
    assert job["needs_user_input_count"] == 1
    assert job["stop_reason"] == "max_rounds_reached"
    assets = consulting_agent.list_source_assets(session["session_id"])
    assert {asset["status"] for asset in assets} == {"archived", "needs_user_input"}


def test_consult_capture_job_migrates_and_hides_legacy_exception_text(
    monkeypatch, tmp_path
):
    db_file = tmp_path / "consult.sqlite3"
    monkeypatch.setattr(consulting_agent, "CONSULTING_AGENT_DB_FILE", str(db_file))
    tracker.app.config["TESTING"] = True
    client = tracker.app.test_client()
    _login_cookies(client)
    session = consulting_agent.create_session("搜集1份公开来源")
    job = consulting_agent.create_capture_job(session["session_id"], target_count=1)
    private_detail = "timeout [private-path] https://api.test/?credential=[private-value]"
    with sqlite3.connect(db_file) as conn:
        conn.execute(
            "UPDATE capture_jobs SET status='failed', stop_reason=? WHERE job_id=?",
            (private_detail, job["job_id"]),
        )
        conn.commit()

    response = client.get(
        f"/api/consult/sessions/{session['session_id']}/capture_jobs/{job['job_id']}"
    )

    assert response.status_code == 200
    payload = response.get_json()
    returned = payload["job"]
    assert returned["stop_reason"] == consulting_agent.CAPTURE_FAILURE_FALLBACK
    assert returned["stop_reason_label"] == consulting_agent.CAPTURE_FAILURE_FALLBACK
    assert (
        tracker._consult_stop_reason_label(private_detail)
        == consulting_agent.CAPTURE_FAILURE_FALLBACK
    )
    assert private_detail not in json.dumps(payload, ensure_ascii=False)

    consulting_agent.init_consulting_agent_db()
    consulting_agent.init_consulting_agent_db()
    with sqlite3.connect(db_file) as conn:
        stored = conn.execute(
            "SELECT stop_reason FROM capture_jobs WHERE job_id=?", (job["job_id"],)
        ).fetchone()[0]
    assert stored == consulting_agent.CAPTURE_FAILURE_FALLBACK


def test_consult_capture_uses_browser_render_fallback(monkeypatch, tmp_path):
    monkeypatch.setattr(consulting_agent, "CONSULTING_AGENT_DB_FILE", str(tmp_path / "consult.sqlite3"))
    monkeypatch.setattr(consulting_agent, "SOURCE_ARCHIVE_DIR", str(tmp_path / "source_archive"), raising=False)
    tracker.SEARCH_CONFIG.clear()

    monkeypatch.setattr(tracker.search_adapters, "search_web_multi", lambda queries, target_count, **kwargs: ([
        {
            "title": "European air defense production report",
            "source": "Think Tank",
            "url": "https://example.test/rendered-report",
            "snippet": "Requires rendered HTML.",
            "channel": "web",
            "score": 91,
            "provider": "public_web",
            "query": queries[0],
            "rank": 1,
            "document_type": "html",
        }
    ], {"queries": queries, "deduped_count": 1, "rejected_low_relevance": 0}))
    monkeypatch.setattr(tracker.search_adapters, "extract_url", lambda url, **kwargs: (_ for _ in ()).throw(RuntimeError("empty body")))
    monkeypatch.setattr(tracker.search_adapters, "extract_url_rendered", lambda url, **kwargs: {
        "title": "European air defense production report",
        "url": url,
        "text": "Rendered public report text about European air defense production capacity.",
        "snippet": "Rendered public report text about European air defense.",
        "document_type": "html",
        "content_type": "text/html",
        "word_count": 180,
        "raw_bytes": b"<html>rendered</html>",
        "is_fetched_original": True,
    }, raising=False)
    monkeypatch.setattr(tracker, "_is_ssrf_safe", lambda url: (True, ""))
    tracker.app.config["TESTING"] = True
    client = tracker.app.test_client()
    csrf = _login_cookies(client)
    session = consulting_agent.create_session("搜集1份欧洲防空生产报告")

    resp = client.post(
        f"/api/consult/sessions/{session['session_id']}/capture_to_target",
        json={"target_count": 1, "batch_size": 1, "max_rounds": 1, "allow_browser_render": True},
        headers={tracker.CSRF_HEADER: csrf},
    )
    data = resp.get_json()

    assert resp.status_code == 200
    assert data["archived_count"] == 1
    asset = consulting_agent.list_source_assets(session["session_id"])[0]
    assert asset["status"] == "archived"
    assert asset["payload"]["extraction_method"] == "browser_render"


def test_consult_capture_defaults_to_fast_direct_mode(monkeypatch, tmp_path):
    monkeypatch.setattr(consulting_agent, "CONSULTING_AGENT_DB_FILE", str(tmp_path / "consult.sqlite3"))
    monkeypatch.setattr(consulting_agent, "SOURCE_ARCHIVE_DIR", str(tmp_path / "source_archive"), raising=False)
    tracker.SEARCH_CONFIG.clear()
    rendered_called = {"value": False}

    monkeypatch.setattr(tracker.search_adapters, "search_web_multi", lambda queries, target_count, **kwargs: ([
        {
            "title": "Iran missile inventory report",
            "source": "Think Tank",
            "url": "https://example.test/slow-render-only",
            "snippet": "Requires expensive rendered HTML.",
            "channel": "web",
            "score": 91,
            "provider": "public_web",
            "query": queries[0],
            "rank": 1,
            "document_type": "html",
        }
    ], {"queries": queries, "deduped_count": 1, "rejected_low_relevance": 0}))
    monkeypatch.setattr(tracker.search_adapters, "extract_url", lambda url, **kwargs: (_ for _ in ()).throw(RuntimeError("timeout")))

    def fake_rendered(url, **kwargs):
        rendered_called["value"] = True
        return {
            "title": "Rendered slow report",
            "url": url,
            "text": "Rendered content should only be used in deep retry mode.",
            "snippet": "Rendered content.",
            "document_type": "html",
            "content_type": "text/html",
            "word_count": 10,
            "raw_bytes": b"<html>rendered</html>",
            "is_fetched_original": True,
        }

    monkeypatch.setattr(tracker.search_adapters, "extract_url_rendered", fake_rendered, raising=False)
    monkeypatch.setattr(tracker, "_is_ssrf_safe", lambda url: (True, ""))
    tracker.app.config["TESTING"] = True
    client = tracker.app.test_client()
    csrf = _login_cookies(client)
    session = consulting_agent.create_session("搜集1份伊朗导弹库存报告")

    resp = client.post(
        f"/api/consult/sessions/{session['session_id']}/capture_to_target",
        json={"target_count": 1, "batch_size": 1, "max_rounds": 1},
        headers={tracker.CSRF_HEADER: csrf},
    )
    data = resp.get_json()

    assert resp.status_code == 200
    assert data["allow_browser_render"] is False
    assert data["archived_count"] == 0
    assert data["failed_count"] == 1
    assert data["failure_breakdown"]["timeout"] == 1
    assert data["diagnosis_cards"][0]["diagnosis"]["code"] == "timeout"
    assert data["current_phase"] == "完成本轮，仍需补抓"
    assert rendered_called["value"] is False


def test_consult_archive_failure_hides_internal_exception_text(monkeypatch):
    private_detail = "timeout at C:\\Users\\private\\source.pdf\r\nINJECTED"
    evidence = {
        "evidence_id": "evidence-safe-error",
        "title": "Public report",
        "url": "https://example.test/report",
        "source": "Example",
    }
    captured = {}

    monkeypatch.setattr(tracker, "_is_ssrf_safe", lambda _url: (True, ""))
    monkeypatch.setattr(
        tracker.search_adapters,
        "extract_url",
        lambda _url, **_kwargs: (_ for _ in ()).throw(RuntimeError(private_detail)),
    )

    def fake_record(_session_id, _evidence, reason, _url, **kwargs):
        captured.update({"reason": reason, **kwargs})
        return {"asset_id": "asset-safe-error", "status": "failed"}

    monkeypatch.setattr(consulting_agent, "record_source_asset_failure", fake_record)

    _updated, _asset, failure = tracker._consult_archive_one("session", evidence)

    assert failure["reason"] == "站点响应慢或连接超时"
    assert failure["diagnosis"]["code"] == "timeout"
    assert captured["reason"] == "站点响应慢或连接超时"
    assert private_detail not in json.dumps(failure, ensure_ascii=False)
    assert private_detail not in json.dumps(captured, ensure_ascii=False)


@pytest.mark.parametrize(
    "error,expected_status,expected_message",
    [
        (ValueError("invalid path C:\\Users\\private\\config\r\nINJECTED"), 400, "请求参数无效"),
        (KeyError("private-session-identifier"), 404, "资源不存在"),
        (RuntimeError("secret at https://private.example.test/token"), 500, "防务咨询Agent处理失败"),
    ],
)
def test_consult_error_response_does_not_reflect_exception_text(
    error, expected_status, expected_message
):
    tracker.app.config["TESTING"] = True
    with tracker.app.app_context():
        response, status = tracker._consult_error_response(error)

    assert status == expected_status
    assert response.get_json() == {"error": expected_message}
    assert str(error) not in response.get_data(as_text=True)


def test_consult_error_response_preserves_enumerated_public_value_error():
    tracker.app.config["TESTING"] = True
    with tracker.app.app_context():
        response, status = tracker._consult_error_response(ValueError("缺少可用证据"))

    assert status == 400
    assert response.get_json() == {"error": "缺少可用证据"}


@pytest.mark.parametrize(
    "endpoint,expected_message",
    [
        ("search", "实时搜索请求失败"),
        ("web_search", "联网搜索请求失败"),
    ],
)
def test_consult_search_does_not_reflect_upstream_http_error(
    monkeypatch, tmp_path, endpoint, expected_message
):
    private_detail = (
        "upstream failure C:\\Users\\private\\search.env "
        "https://private.example.test/?token=secret\r\nTRACEBACK"
    )
    upstream_response = tracker.requests.Response()
    upstream_response.status_code = 503
    upstream_response._content = private_detail.encode("utf-8")
    upstream_response.url = "https://private.example.test/?token=secret"
    upstream_error = tracker.requests.exceptions.HTTPError(
        private_detail, response=upstream_response
    )

    monkeypatch.setattr(
        consulting_agent,
        "CONSULTING_AGENT_DB_FILE",
        str(tmp_path / "consult.sqlite3"),
    )
    monkeypatch.setattr(
        tracker,
        "_masked_search_config_status",
        lambda: {
            "online_search_enabled": True,
            "web_search_enabled": True,
            "message": "联网搜索已启用",
        },
    )

    def fail_search(*_args, **_kwargs):
        raise upstream_error

    monkeypatch.setattr(tracker.search_adapters, "search_web_multi", fail_search)
    tracker.app.config["TESTING"] = True
    client = tracker.app.test_client()
    csrf = _login_cookies(client)
    session = consulting_agent.create_session("搜集1份台海无人系统报告")

    response = client.post(
        f"/api/consult/sessions/{session['session_id']}/{endpoint}",
        json={"target_count": 1},
        headers={tracker.CSRF_HEADER: csrf},
    )

    assert response.status_code == 502
    assert response.get_json() == {"error": expected_message}
    response_text = response.get_data(as_text=True)
    assert private_detail not in response_text
    assert "private.example.test" not in response_text
    assert "TRACEBACK" not in response_text


def test_consult_capture_job_does_not_persist_internal_exception_text(monkeypatch):
    private_detail = "failure at C:\\Users\\private\\capture.db\r\nINJECTED"
    updates = []
    monkeypatch.setattr(
        consulting_agent,
        "get_capture_job",
        lambda _session_id, _job_id: {
            "target_count": 1,
            "batch_size": 1,
            "max_rounds": 1,
        },
    )
    monkeypatch.setattr(
        consulting_agent,
        "get_session",
        lambda _session_id: (_ for _ in ()).throw(RuntimeError(private_detail)),
    )
    monkeypatch.setattr(
        consulting_agent,
        "update_capture_job",
        lambda _job_id, **kwargs: updates.append(kwargs),
    )
    monkeypatch.setattr(
        tracker,
        "_consult_capture_counts",
        lambda *_args: {"archived_count": 0, "target_count": 1},
    )

    tracker._run_consult_capture_job("session", "job")

    assert updates[-1]["status"] == "failed"
    assert updates[-1]["stop_reason"] == "抓取失败，原因待复核"
    assert private_detail not in json.dumps(updates, ensure_ascii=False)


def test_consult_web_search_uses_multi_provider_without_50_cap(monkeypatch, tmp_path):
    monkeypatch.setattr(consulting_agent, "CONSULTING_AGENT_DB_FILE", str(tmp_path / "consult.sqlite3"))
    tracker.SEARCH_CONFIG.clear()
    tracker.SEARCH_CONFIG.update({"tavily_api_key": "tv", "brave_api_key": "brave"})

    def fake_multi(queries, target_count, providers=None, config=None, **kwargs):
        return (
            [
                {
                    "title": f"Think tank report {i}",
                    "source": "RAND",
                    "url": f"https://example.test/report-{i}.pdf",
                    "snippet": "公开智库报告原文摘要",
                    "channel": "web",
                    "score": 90,
                    "provider": "tavily" if i % 2 == 0 else "brave",
                    "query": queries[0],
                    "rank": i + 1,
                    "document_type": "pdf",
                }
                for i in range(target_count)
            ],
            {
                "target_count": target_count,
                "queries": queries,
                "provider_stats": {"tavily": {"returned": 40}, "brave": {"returned": 40}},
                "provider_errors": {},
                "deduped_count": target_count,
            },
        )

    monkeypatch.setattr(tracker.search_adapters, "search_web_multi", fake_multi)
    tracker.app.config["TESTING"] = True
    client = tracker.app.test_client()
    csrf = _login_cookies(client)
    session = consulting_agent.create_session("搜集80份台海无人系统智库报告")

    resp = client.post(
        f"/api/consult/sessions/{session['session_id']}/web_search",
        json={"target_count": 80, "providers": ["tavily", "brave"], "include_pdf": True},
        headers={tracker.CSRF_HEADER: csrf},
    )
    data = resp.get_json()

    assert resp.status_code == 200
    assert data["total"] == 80
    assert data["meta"]["target_source_count"] == 80
    assert data["meta"]["archived_count"] == 0
    assert data["meta"]["unresolved_target_count"] >= 0
    assert data["meta"]["topic_plan"]["target_source_count"] == 80
    assert data["meta"]["next_queries"]
    assert data["meta"]["provider_stats"]["tavily"]["returned"] == 40
    assert data["evidence"][0]["payload"]["provider"] in {"tavily", "brave"}


def test_consult_extract_marks_source_as_fetched_original(monkeypatch, tmp_path):
    monkeypatch.setattr(consulting_agent, "CONSULTING_AGENT_DB_FILE", str(tmp_path / "consult.sqlite3"))
    monkeypatch.setattr(consulting_agent, "SOURCE_ARCHIVE_DIR", str(tmp_path / "source_archive"), raising=False)
    session = consulting_agent.create_session("搜集1份台海无人系统智库报告")
    evidence = consulting_agent.upsert_evidence(session["session_id"], [{
        "title": "RAND 台海无人系统报告",
        "source": "RAND",
        "url": "https://example.test/report",
        "channel": "web",
        "score": 91,
        "snippet": "待抓取正文。",
    }])
    monkeypatch.setattr(
        tracker.search_adapters,
        "extract_url",
        lambda url, **kwargs: {
            "title": "RAND 台海无人系统报告",
            "text": "公开报告正文，讨论无人系统、侦察打击链和作战概念。",
            "snippet": "公开报告正文，讨论无人系统。",
            "document_type": "html",
            "word_count": 180,
            "is_fetched_original": True,
        },
    )
    monkeypatch.setattr(tracker, "_is_ssrf_safe", lambda url: (True, ""))
    tracker.app.config["TESTING"] = True
    client = tracker.app.test_client()
    csrf = _login_cookies(client)

    resp = client.post(
        f"/api/consult/sessions/{session['session_id']}/extract",
        json={"evidence_ids": [evidence[0]["evidence_id"]]},
        headers={tracker.CSRF_HEADER: csrf},
    )
    data = resp.get_json()

    assert resp.status_code == 200
    assert data["extracted_count"] == 1
    assert data["archived_count"] == 1
    assert data["assets"][0]["status"] == "archived"
    assert Path(data["assets"][0]["text_path"]).exists()
    updated = data["evidence"][0]
    assert updated["payload"]["is_fetched_original"] is True
    assert updated["payload"]["asset_status"] == "archived"
    assert updated["payload"]["document_type"] == "html"
    assert "公开报告正文" in updated["snippet"]

    assets_resp = client.get(f"/api/consult/sessions/{session['session_id']}/assets")
    assets_data = assets_resp.get_json()
    assert assets_resp.status_code == 200
    assert assets_data["archived_count"] == 1
    assert assets_data["assets"][0]["evidence_id"] == evidence[0]["evidence_id"]


def test_consult_extract_resolves_thinktank_target_before_archiving(monkeypatch, tmp_path):
    monkeypatch.setattr(consulting_agent, "CONSULTING_AGENT_DB_FILE", str(tmp_path / "consult.sqlite3"))
    monkeypatch.setattr(consulting_agent, "SOURCE_ARCHIVE_DIR", str(tmp_path / "source_archive"), raising=False)
    session = consulting_agent.create_session("搜集1份俄罗斯电子战智库报告")
    target = consulting_agent.upsert_evidence(session["session_id"], [{
        "title": "RUSI：欧洲安全研究入口",
        "source": "RUSI",
        "url": "https://www.rusi.org/",
        "channel": "thinktank_target",
        "score": 80,
        "snippet": "站内入口",
    }])[0]

    def fake_search(queries, target_count, providers=None, config=None, **kwargs):
        return ([{
            "title": "Russia electronic warfare report",
            "source": "RUSI",
            "url": "https://www.rusi.org/russia-electronic-warfare-report",
            "channel": "web",
            "score": 91,
            "snippet": "Concrete report page",
            "provider": "public_web",
            "query": queries[0],
            "document_type": "html",
        }], {"queries": queries, "deduped_count": 1, "provider_errors": {}})

    monkeypatch.setattr(tracker.search_adapters, "search_web_multi", fake_search)
    monkeypatch.setattr(
        tracker.search_adapters,
        "extract_url",
        lambda url, **kwargs: {
            "title": "Russia electronic warfare report",
            "text": "Public report text about Russian electronic warfare capability.",
            "snippet": "Public report text about Russian electronic warfare.",
            "document_type": "html",
            "content_type": "text/html",
            "word_count": 180,
            "is_fetched_original": True,
        },
    )
    monkeypatch.setattr(tracker, "_is_ssrf_safe", lambda url: (True, ""))
    tracker.SEARCH_CONFIG.clear()
    tracker.app.config["TESTING"] = True
    client = tracker.app.test_client()
    csrf = _login_cookies(client)

    resp = client.post(
        f"/api/consult/sessions/{session['session_id']}/extract",
        json={"evidence_ids": [target["evidence_id"]]},
        headers={tracker.CSRF_HEADER: csrf},
    )
    data = resp.get_json()

    assert resp.status_code == 200
    assert data["archived_count"] == 1
    assert data["resolved_evidence"][0]["channel"] == "web"
    assert "rusi.org/russia-electronic-warfare-report" in data["resolved_evidence"][0]["url"]
    assert data["assets"][0]["status"] == "archived"


def test_consult_extract_continues_after_one_url_failure(monkeypatch, tmp_path):
    monkeypatch.setattr(consulting_agent, "CONSULTING_AGENT_DB_FILE", str(tmp_path / "consult.sqlite3"))
    monkeypatch.setattr(consulting_agent, "SOURCE_ARCHIVE_DIR", str(tmp_path / "source_archive"), raising=False)
    session = consulting_agent.create_session("搜集2份伊朗导弹智库报告")
    evidence = consulting_agent.upsert_evidence(session["session_id"], [
        {
            "title": "Iran missile assessment",
            "source": "CSIS",
            "url": "https://example.test/ok",
            "channel": "web",
            "score": 91,
            "snippet": "待抓取正文。",
        },
        {
            "title": "Blocked report",
            "source": "Think Tank",
            "url": "https://example.test/blocked",
            "channel": "web",
            "score": 88,
            "snippet": "待抓取正文。",
        },
    ])

    def fake_extract(url, **kwargs):
        if "blocked" in url:
            raise RuntimeError("403 Forbidden")
        return {
            "title": "Iran missile assessment",
            "text": "公开报告正文，讨论伊朗弹道导弹与作战运用。",
            "snippet": "公开报告正文，讨论伊朗弹道导弹。",
            "document_type": "html",
            "word_count": 180,
            "is_fetched_original": True,
        }

    monkeypatch.setattr(tracker.search_adapters, "extract_url", fake_extract)
    monkeypatch.setattr(tracker, "_is_ssrf_safe", lambda url: (True, ""))
    tracker.app.config["TESTING"] = True
    client = tracker.app.test_client()
    csrf = _login_cookies(client)

    resp = client.post(
        f"/api/consult/sessions/{session['session_id']}/extract",
        json={"evidence_ids": [row["evidence_id"] for row in evidence]},
        headers={tracker.CSRF_HEADER: csrf},
    )
    data = resp.get_json()

    assert resp.status_code == 200
    assert data["extracted_count"] == 1
    assert data["archived_count"] == 1
    assert len(data["failures"]) == 1
    assert data["failures"][0]["title"] == "Blocked report"
    assert data["failures"][0]["reason"] == "站点拒绝访问或需要授权"
    assert "403" not in str(data["failures"])


def test_consult_handoff_to_report_agent_filters_unfetched_targets(monkeypatch, tmp_path):
    monkeypatch.setattr(consulting_agent, "CONSULTING_AGENT_DB_FILE", str(tmp_path / "consult.sqlite3"))
    monkeypatch.setattr(report_agent, "REPORT_AGENT_DB_FILE", str(tmp_path / "report.sqlite3"))
    monkeypatch.setattr(consulting_agent, "SOURCE_ARCHIVE_DIR", str(tmp_path / "source_archive"), raising=False)
    session = consulting_agent.create_session("帮我做一个台海无人系统战略分析报告，搜集2份来源")
    archived_doc = {
        "title": "RAND 台海无人系统报告",
        "url": "https://example.test/rand-report.pdf",
        "text": "公开报告正文，讨论无人系统与跨域作战。",
        "snippet": "已归档公开报告正文。",
        "document_type": "pdf",
        "content_type": "application/pdf",
        "word_count": 180,
        "raw_bytes": b"%PDF-1.4 unit test",
        "is_fetched_original": True,
    }
    evidence = consulting_agent.upsert_evidence(session["session_id"], [
        {
            "title": "RAND 台海无人系统报告",
            "source": "RAND",
            "url": "https://example.test/rand-report.pdf",
            "channel": "web",
            "score": 92,
            "snippet": "已抓取公开报告正文。",
        },
        {
            "title": "RAND China Research",
            "source": "RAND",
            "url": "https://www.rand.org/topics/china.html",
            "channel": "thinktank_target",
            "score": 80,
            "snippet": "目录入口，不能作为原文证据。",
            "payload": {"is_fetched_original": False},
        },
    ])
    asset = consulting_agent.archive_source_asset(session["session_id"], evidence[0], archived_doc)
    evidence = consulting_agent.upsert_evidence(session["session_id"], [{
        **evidence[0],
        "payload": {"is_fetched_original": True, "document_type": "pdf", "asset_status": "archived", "asset_id": asset["asset_id"]},
    }, evidence[1]])
    tracker.app.config["TESTING"] = True
    client = tracker.app.test_client()
    csrf = _login_cookies(client)

    resp = client.post(
        f"/api/consult/sessions/{session['session_id']}/handoff_to_report_agent",
        json={"evidence_ids": [row["evidence_id"] for row in evidence]},
        headers={tracker.CSRF_HEADER: csrf},
    )
    data = resp.get_json()

    assert resp.status_code == 200
    assert data["imported_count"] == 1
    bundle = report_agent.get_project_bundle(data["project"]["project_id"])
    assert len(bundle["evidence"]) == 1
    assert bundle["evidence"][0]["title"] == "RAND 台海无人系统报告"
    assert "公开报告正文" in bundle["evidence"][0]["summary"]


def test_consult_handoff_requires_archived_asset_not_just_payload_flag(monkeypatch, tmp_path):
    monkeypatch.setattr(consulting_agent, "CONSULTING_AGENT_DB_FILE", str(tmp_path / "consult.sqlite3"))
    monkeypatch.setattr(report_agent, "REPORT_AGENT_DB_FILE", str(tmp_path / "report.sqlite3"))
    session = consulting_agent.create_session("帮我做一个红海无人机报告，搜集1份来源")
    evidence = consulting_agent.upsert_evidence(session["session_id"], [{
        "title": "Red Sea UAV article",
        "source": "Defense outlet",
        "url": "https://example.test/red-sea-uav",
        "channel": "web",
        "score": 88,
        "snippet": "只有摘要，没有本地归档资产。",
        "payload": {"is_fetched_original": True, "document_type": "html"},
    }])
    tracker.app.config["TESTING"] = True
    client = tracker.app.test_client()
    csrf = _login_cookies(client)

    resp = client.post(
        f"/api/consult/sessions/{session['session_id']}/handoff_to_report_agent",
        json={"evidence_ids": [evidence[0]["evidence_id"]]},
        headers={tracker.CSRF_HEADER: csrf},
    )
    data = resp.get_json()

    assert resp.status_code == 400
    assert data["skipped_count"] == 1
    assert "未归档" in data["failures"][0]["reason"]
