# -*- coding: utf-8 -*-
"""One-database acceptance path for the local V9 MVP lifecycle."""

from datetime import datetime, timezone
from io import BytesIO

from docx import Document
from flask import Flask


def test_v9_mvp_lifecycle_uses_one_database_and_preserves_recall_semantics(
    tmp_path,
):
    """News becomes evidence and ends as a recalled, auditable snapshot."""
    from v9.api import create_blueprint
    from v9.service import V9Service

    database_path = tmp_path / "v9-mvp.sqlite3"
    service = V9Service(database_path, tmp_path / ".v9-master.key")
    context = service.get_or_create_personal_context()
    service.acknowledge_personal_recovery(context["organization_id"])
    current_news = {
        "aid": "mvp-lifecycle-news",
        "title": "Carrier activity near Taiwan",
        "summary": "公开来源记录到新的活动节奏。",
        "source": "Public Source A",
        "link": "https://example.test/mvp-lifecycle-news",
        "date": datetime.now(timezone.utc).isoformat(),
        "priority": {"stars": 9},
    }
    app = Flask(__name__)
    app.register_blueprint(
        create_blueprint(
            lambda: service,
            news_provider=lambda: [current_news],
            agent_phase_executor=lambda payload: (
                f"已核对 {len(payload['evidence'])} 条证据 [E1]"
            ),
        )
    )
    client = app.test_client()
    client.environ_base.update({
        "HTTP_X_V9_CONTEXT_MODE": "personal",
        "HTTP_X_V9_ORGANIZATION_ID": context["organization_id"],
    })

    # News/source -> encrypted evidence -> alert -> case.
    rule = client.post(
        "/api/v9/alert-rules",
        json={
            "name": "航母活动",
            "keywords": ["carrier"],
            "min_stars": 7,
            "severity": "high",
        },
    )
    assert rule.status_code == 201
    materialized = client.post("/api/v9/alerts/materialize", json={})
    assert materialized.status_code == 200
    assert materialized.get_json()["created"] == 1

    evidence = client.get("/api/v9/evidence").get_json()["evidence"]
    assert len(evidence) == 1
    assert evidence[0]["content"]["title"] == current_news["title"]
    assert (
        evidence[0]["content"]["provenance"]["url"]
        == current_news["link"]
    )
    evidence_id = evidence[0]["record_id"]

    alert = client.get("/api/v9/alerts").get_json()["alerts"][0]
    converted = client.post(
        f"/api/v9/alerts/{alert['record_id']}/action",
        json={"action": "convert_case", "version": alert["version"]},
    )
    assert converted.status_code == 200
    case_id = converted.get_json()["case_id"]
    case = client.get("/api/v9/cases").get_json()["cases"][0]
    assert case["record_id"] == case_id
    assert case["content"]["evidence_ids"] == [evidence_id]

    # The task and scenario consume the same archived evidence.
    created_job = client.post(
        "/api/v9/jobs",
        json={
            "template": "rapid_assessment",
            "title": "单库快速研判",
            "instructions": f"围绕案件 {case_id} 核对来源。",
            "evidence_ids": [evidence_id],
        },
        headers={"Idempotency-Key": "mvp-lifecycle-job"},
    )
    assert created_job.status_code == 201
    job = client.get("/api/v9/jobs").get_json()["jobs"][0]
    started = client.post(
        f"/api/v9/jobs/{job['record_id']}/action",
        json={"action": "start", "version": job["version"]},
    )
    assert started.status_code == 200
    executed = client.post(
        f"/api/v9/jobs/{job['record_id']}/action",
        json={
            "action": "execute_phase",
            "version": started.get_json()["version"],
        },
    )
    assert executed.status_code == 200
    assert executed.get_json()["transition"]["phase"] == "close_read"
    executed_job = client.get("/api/v9/jobs").get_json()["jobs"][0]
    assert executed_job["content"]["evidence_ids"] == [evidence_id]
    assert "[E1]" in executed_job["content"]["stage_outputs"]["collect"]

    scenario_created = client.post(
        "/api/v9/scenarios",
        json={
            "title": "未来 72 小时三分支推演",
            "question": "活动节奏将如何演化？",
            "evidence_ids": [evidence_id],
            "assumptions": ["仅依据已归档公开来源"],
        },
    )
    assert scenario_created.status_code == 201
    scenario = client.get("/api/v9/scenarios").get_json()["scenarios"][0]
    scenario_updated = client.patch(
        f"/api/v9/scenarios/{scenario['record_id']}",
        json={
            "version": scenario["version"],
            "changes": {
                "branches": {
                    "baseline": {
                        "summary": "活动节奏维持。",
                        "triggers": ["连续公开通报"],
                        "indicators": ["公开活动频次"],
                        "counter_evidence_ids": [],
                        "confidence": 0.5,
                    },
                    "escalation": {
                        "summary": "活动节奏上升。",
                        "triggers": ["活动范围扩大"],
                        "indicators": ["新增活动区域"],
                        "counter_evidence_ids": [],
                        "confidence": 0.3,
                    },
                    "deescalation": {
                        "summary": "活动节奏下降。",
                        "triggers": ["公开活动减少"],
                        "indicators": ["通报间隔延长"],
                        "counter_evidence_ids": [],
                        "confidence": 0.2,
                    },
                },
                "team_outputs": {
                    "red": {
                        "text": "升级分支仍需更多来源支持。",
                        "evidence_ids": [evidence_id],
                    },
                    "blue": {
                        "text": "基准分支与当前材料一致。",
                        "evidence_ids": [evidence_id],
                    },
                    "judge": {
                        "text": "三分支均为情景推断，不当作事实。",
                        "evidence_ids": [evidence_id],
                    },
                },
            },
        },
    )
    assert scenario_updated.status_code == 200
    stored_scenario = client.get("/api/v9/scenarios").get_json()[
        "scenarios"
    ][0]
    assert stored_scenario["content"]["classification"] == (
        "scenario_inference"
    )
    assert stored_scenario["content"]["team_outputs"]["judge"][
        "evidence_ids"
    ] == [evidence_id]

    # Evidence-bound document -> validation -> local approval snapshot.
    document_created = client.post(
        "/api/v9/documents",
        json={
            "kind": "report",
            "stage": "ready",
            "title": "公开来源活动研判",
            "paragraphs": [
                {
                    "heading": "核心判断",
                    "text": "公开来源显示相关活动节奏出现变化。",
                    "evidence_ids": [evidence_id],
                    "source_status": "source_claim",
                    "fact_check": "passed",
                    "fact_check_note": "已与归档原文逐项核对。",
                },
                {
                    "heading": "情景边界",
                    "text": "未来走向仅为三分支推演，不能当作已发生事实。",
                    "evidence_ids": [evidence_id],
                    "source_status": "scenario_assumption",
                    "fact_check": "passed",
                },
            ],
        },
    )
    assert document_created.status_code == 201
    document = client.get("/api/v9/documents").get_json()["documents"][0]
    assert document["content"]["validation"]["ready"] is True
    assert document["content"]["validation"]["evidence_count"] == 1

    publication_created = client.post(
        "/api/v9/publications",
        json={"document_id": document["record_id"]},
    )
    assert publication_created.status_code == 201
    publication = client.get("/api/v9/publications").get_json()[
        "publications"
    ][0]
    pending = client.patch(
        f"/api/v9/publications/{publication['record_id']}",
        json={
            "version": publication["version"],
            "status": "pending_approval",
        },
    )
    assert pending.status_code == 200
    signed = client.post(
        f"/api/v9/publications/{publication['record_id']}/sign",
        json={"version": pending.get_json()["version"]},
    )
    assert signed.status_code == 200
    assert signed.get_json()["status"] == "signed"
    assert signed.get_json()["publication_scope"] == (
        "local_approval_snapshot"
    )
    assert signed.get_json()["external_published"] is False

    recalled = client.post(
        f"/api/v9/publications/{publication['record_id']}/recall",
        json={
            "version": signed.get_json()["version"],
            "reason": "新增公开材料要求重新核对判断",
        },
    )
    assert recalled.status_code == 200
    assert recalled.get_json()["status"] == "recalled"

    stored_publication = client.get("/api/v9/publications").get_json()[
        "publications"
    ][0]
    assert stored_publication["content"]["status"] == "recalled"
    assert stored_publication["content"]["signed_snapshot"]["receipt"][
        "document_id"
    ] == document["record_id"]

    # Recall closes the ordinary export path but keeps an explicitly marked
    # audit artifact available from that same stored snapshot.
    export_url = (
        f"/api/v9/publications/{publication['record_id']}/export.docx"
    )
    ordinary_export = client.get(export_url)
    assert ordinary_export.status_code == 409
    assert ordinary_export.get_json()["code"] == "PUBLICATION_RECALLED"
    assert ordinary_export.get_json()["retryable"] is False

    audit_export = client.get(f"{export_url}?mode=audit")
    assert audit_export.status_code == 200
    audit_document = Document(BytesIO(audit_export.data))
    audit_text = "\n".join(
        paragraph.text for paragraph in audit_document.paragraphs
    )
    audit_headers = "\n".join(
        paragraph.text
        for section in audit_document.sections
        for paragraph in section.header.paragraphs
    )
    assert "已撤回" in audit_text
    assert "新增公开材料要求重新核对判断" in audit_text
    assert "已撤回" in audit_headers

    audit_actions = {
        item["content"]["action"]
        for item in client.get("/api/v9/audit-events").get_json()["events"]
    }
    assert {
        "publication.created",
        "publication.moved",
        "publication.signed",
        "publication.recalled",
    } <= audit_actions

    assert database_path.is_file()
    assert service.database_path == database_path
