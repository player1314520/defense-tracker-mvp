"""
飞书云端机器人 v2.0 — 独立部署版 + RSS 自动推送
═══════════════════════════════════════════════════════
完全独立运行，不依赖 app.py，可部署到任何云平台（Railway / Render / Zeabur / VPS）。
手机发链接给飞书机器人 → 云端自动抓取+AI生成要讯 → 回复卡片，无需开电脑。
定时扫描 15 个核心 RSS 源 → 筛选高价值防务文章 → 自动推送到飞书群。

环境变量（必填）：
  FEISHU_APP_ID          飞书应用 App ID
  FEISHU_APP_SECRET      飞书应用 App Secret
  AI_API_KEY             AI 服务 API Key

环境变量（选填）：
  FEISHU_VERIFY_TOKEN    飞书事件验证 Token（推荐配置，防伪造）
  AI_BASE_URL            AI 服务地址（默认 https://api.deepseek.com）
  AI_MODEL               模型名称（默认 deepseek-chat）
  PORT                   监听端口（默认 5000）

RSS 自动推送（选填）：
  FEISHU_PUSH_CHAT_ID    推送目标群 chat_id（私聊或群聊皆可）
  PUSH_MODE              headlines = 仅标题摘要（免费）/ brief = AI生成要讯（消耗token）
  PUSH_INTERVAL          推送间隔分钟（默认 30）
  PUSH_MAX_ARTICLES      每轮最多推送文章数（默认 5）
═══════════════════════════════════════════════════════
"""

import base64, hashlib, hmac, ipaddress, json, logging, os, re, socket, sys, threading, time
from collections import OrderedDict
from concurrent.futures import ThreadPoolExecutor, as_completed
from datetime import datetime, timezone, timedelta
from difflib import SequenceMatcher
from io import BytesIO
from urllib.parse import urlparse, urljoin

import feedparser
import requests
from bs4 import BeautifulSoup
from docx import Document as DocxDocument
from docx.shared import Pt, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from flask import Flask, request, jsonify
from feishu_common import (
    ascii_fallback_name as _ascii_fallback_name,
    sanitize_feishu_filename as _sanitize_feishu_filename,
)

try:
    import trafilatura
    _HAS_TRAFILATURA = True
except ImportError:
    _HAS_TRAFILATURA = False

try:
    import pdfplumber
    _HAS_PDFPLUMBER = True
except ImportError:
    _HAS_PDFPLUMBER = False

# ════════════════════════════════════════════════════════════
# 日志 & Flask
# ════════════════════════════════════════════════════════════
logging.basicConfig(level=logging.INFO, format="%(asctime)s [%(levelname)s] %(message)s")
logger = logging.getLogger("feishu_cloud")

app = Flask(__name__)

# ════════════════════════════════════════════════════════════
# 配置
# ════════════════════════════════════════════════════════════
FEISHU_CONFIG = {
    "app_id":       os.environ.get("FEISHU_APP_ID", ""),
    "app_secret":   os.environ.get("FEISHU_APP_SECRET", ""),
    "verify_token": os.environ.get("FEISHU_VERIFY_TOKEN", ""),
}

AI_CONFIG = {
    "api_key":     os.environ.get("AI_API_KEY", ""),
    "base_url":    os.environ.get("AI_BASE_URL", "https://api.deepseek.com"),
    "model":       os.environ.get("AI_MODEL", "deepseek-chat"),
    "max_tokens":  1024,
    "temperature": 0.4,
}

FEISHU_API = "https://open.feishu.cn/open-apis"

# ════════════════════════════════════════════════════════════
# 全局线程池（替代无限 daemon thread，防止线程爆炸）
# ════════════════════════════════════════════════════════════
_worker_pool = ThreadPoolExecutor(max_workers=8, thread_name_prefix="feishu_worker")

# ════════════════════════════════════════════════════════════
# RSS 自动推送配置
# ════════════════════════════════════════════════════════════
PUSH_CONFIG = {
    "chat_id":      os.environ.get("FEISHU_PUSH_CHAT_ID", ""),
    "mode":         os.environ.get("PUSH_MODE", "headlines"),     # headlines | brief
    "interval_min": int(os.environ.get("PUSH_INTERVAL", "30")),
    "max_articles": int(os.environ.get("PUSH_MAX_ARTICLES", "5")),
}

MAX_FETCH_BYTES = 5 * 1024 * 1024
MAX_FEISHU_DOWNLOAD_BYTES = 10 * 1024 * 1024
MAX_REDIRECTS = 5

# ════════════════════════════════════════════════════════════
# 启动校验：必填环境变量
# ════════════════════════════════════════════════════════════
_REQUIRED_ENVS = {
    "FEISHU_APP_ID": FEISHU_CONFIG["app_id"],
    "FEISHU_APP_SECRET": FEISHU_CONFIG["app_secret"],
}
_missing = [k for k, v in _REQUIRED_ENVS.items() if not v]
if _missing:
    logger.critical("启动失败！缺少必填环境变量: %s", ", ".join(_missing))
    logger.critical("请在云平台（Railway / Render）的 Variables 中配置后重新部署")
    sys.exit(1)
if not AI_CONFIG["api_key"]:
    logger.warning("AI_API_KEY 未配置，机器人可启动但无法生成要讯")

# 15 个核心 RSS 源（tier 0-1，海外服务器直接可达）
RSS_FEEDS_CLOUD = [
    {"name": "War on the Rocks",      "name_cn": "战争幕后",       "url": "https://warontherocks.com/feed/",                  "focus": "strategy"},
    {"name": "The Diplomat",           "name_cn": "外交家",         "url": "https://thediplomat.com/feed/",                    "focus": "china"},
    {"name": "Breaking Defense",       "name_cn": "突破防务",       "url": "https://breakingdefense.com/feed/",                "focus": "general"},
    {"name": "Defense One",            "name_cn": "防务一号",       "url": "https://www.defenseone.com/rss/all/",              "focus": "policy"},
    {"name": "USNI News",              "name_cn": "美海军研究所",   "url": "https://news.usni.org/feed",                       "focus": "navy"},
    {"name": "SCMP Military",          "name_cn": "南早军事",       "url": "https://www.scmp.com/rss/4/feed",                  "focus": "china"},
    {"name": "Jamestown China Brief",  "name_cn": "詹姆斯敦中国简报", "url": "https://jamestown.org/programs/cb/feed/",        "focus": "china"},
    {"name": "ASPI Strategist",        "name_cn": "澳战略政策研所", "url": "https://www.aspistrategist.org.au/feed",           "focus": "china"},
    {"name": "Atlantic Council",       "name_cn": "大西洋理事会",   "url": "https://www.atlanticcouncil.org/feed/",            "focus": "strategy"},
    {"name": "Foreign Policy",         "name_cn": "外交政策",       "url": "https://foreignpolicy.com/feed/",                  "focus": "general"},
    {"name": "Defense News",           "name_cn": "防务新闻",       "url": "https://www.defensenews.com/arc/outboundfeeds/rss/","focus": "general"},
    {"name": "Naval News",             "name_cn": "海军新闻",       "url": "https://www.navalnews.com/feed/",                  "focus": "navy"},
    {"name": "NHK World Security",     "name_cn": "NHK国际安全",   "url": "https://www3.nhk.or.jp/rss/news/cat6.xml",         "focus": "japan"},
    {"name": "RFA Asia",               "name_cn": "自由亚洲电台",   "url": "https://www.rfa.org/english/rss2.xml",             "focus": "china"},
    {"name": "The Record Cyber",       "name_cn": "Recorded Future", "url": "https://therecord.media/feed",                   "focus": "cyber"},
]

# 防务关键词评分（匹配越多分越高）
_DEFENSE_PATTERNS = [
    # 对华情报（权重 3）
    (3, re.compile(r"china|chinese military|pla\b|beijing|taiwan|south china sea|xi jinping|indo.?pacific", re.I)),
    # 核战略（权重 3）
    (3, re.compile(r"nuclear|icbm|warhead|hypersonic|ballistic missile|deterren", re.I)),
    # 装备动态（权重 2）
    (2, re.compile(r"f-35|aircraft carrier|destroyer|submarine|stealth|drone|uav|missile.*system|fighter|bomber", re.I)),
    # 战略分析（权重 2）
    (2, re.compile(r"strateg|nato|aukus|quad|indopacom|alliance|military.*exercise|force posture|great power", re.I)),
    # 网络安全（权重 1）
    (1, re.compile(r"cyber|hack|espionage|intelligence|zero.day|ransomware|critical infrastructure", re.I)),
    # 突发军情（权重 2）
    (2, re.compile(r"deploy|strike|attack|conflict|escalat|incursion|intercept|provocat|mobiliz|invasion", re.I)),
]

# 已推送文章去重（URL 集合，内存中保留最近 2000 条）
_pushed_urls: set = set()
_pushed_lock = threading.Lock()

# ════════════════════════════════════════════════════════════
# 结果缓存（URL → AI生成结果，1小时过期，省钱）
# ════════════════════════════════════════════════════════════
_result_cache: OrderedDict = OrderedDict()
_cache_lock = threading.Lock()
_CACHE_TTL = 3600
_CACHE_MAX = 200

# ════════════════════════════════════════════════════════════
# 交互会话（per-chat，支持多轮优化，30分钟超时）
# ════════════════════════════════════════════════════════════
_chat_sessions: dict = {}  # chat_id → session state
_pending_filenames: dict = {}  # chat_id → 下次生成要讯时自动使用的导出文件名
_sessions_lock = threading.Lock()
_SESSION_TTL = 1800  # 30 minutes inactive → expire

# 会话可能包含用户素材、提示词和完整成稿，只在内存中保存并按 TTL 清理。
# 不提供明文磁盘持久化开关，避免聊天标识和正文落入运行卷。
_SESSION_STORE = ""


def _sessions_persist_locked():
    """Privacy boundary: full-text chat sessions are intentionally memory-only."""
    return


def _sessions_load():
    """Privacy boundary: no full-text session data is loaded from disk."""
    return


def _cache_get(url: str):
    with _cache_lock:
        if url in _result_cache:
            entry = _result_cache[url]
            if time.time() - entry["time"] < _CACHE_TTL:
                _result_cache.move_to_end(url)
                return entry["result"], entry["source_info"], entry.get("evidence")
            del _result_cache[url]
    return None, None, None


def _cache_set(url: str, result: str, source_info: dict, evidence: dict | None = None):
    with _cache_lock:
        _result_cache[url] = {
            "result": result,
            "source_info": source_info,
            "evidence": evidence,
            "time": time.time(),
        }
        while len(_result_cache) > _CACHE_MAX:
            _result_cache.popitem(last=False)


def _session_get(chat_id: str) -> dict | None:
    with _sessions_lock:
        s = _chat_sessions.get(chat_id)
        if s and time.time() - s["last_active"] < _SESSION_TTL:
            return s
        if s:
            del _chat_sessions[chat_id]
            _sessions_persist_locked()
        return None


def _session_set(chat_id: str, session: dict):
    session["last_active"] = time.time()
    with _sessions_lock:
        # 吸收 pending 文件名：用户在无会话时设置的偏好自动生效
        if chat_id in _pending_filenames and not session.get("custom_filename"):
            session["custom_filename"] = _pending_filenames.pop(chat_id)
        _chat_sessions[chat_id] = session
        _sessions_persist_locked()


def _pending_filename_set(chat_id: str, filename: str):
    with _sessions_lock:
        _pending_filenames[chat_id] = filename
        _sessions_persist_locked()


def _session_clear(chat_id: str):
    with _sessions_lock:
        if _chat_sessions.pop(chat_id, None) is not None:
            _sessions_persist_locked()


_sessions_load()


# ════════════════════════════════════════════════════════════
# AI 调用限速（信号量，最多 2 个并发）
# ════════════════════════════════════════════════════════════
_ai_semaphore = threading.Semaphore(2)

# ════════════════════════════════════════════════════════════
# 历史记录（内存，最近 100 条，重启后清空）
# ════════════════════════════════════════════════════════════
_history: list = []
_history_lock = threading.Lock()
_HISTORY_MAX = 100


def _history_add(title: str, source: str, url: str, brief: str):
    with _history_lock:
        _history.append({
            "title": title, "source": source, "url": url,
            "brief_preview": brief[:150],
            "time": time.strftime("%Y-%m-%d %H:%M:%S"),
        })
        if len(_history) > _HISTORY_MAX:
            _history.pop(0)


_BROWSER_HEADERS = {
    "User-Agent": ("Mozilla/5.0 (Windows NT 10.0; Win64; x64) "
                   "AppleWebKit/537.36 (KHTML, like Gecko) "
                   "Chrome/125.0.0.0 Safari/537.36"),
    "Accept": "text/html,application/xhtml+xml,application/xml;q=0.9,*/*;q=0.8",
    "Accept-Language": "en-US,en;q=0.9,zh-CN;q=0.8",
}

# ════════════════════════════════════════════════════════════
# 消息去重 & 签名校验
# ════════════════════════════════════════════════════════════
_seen_msg_ids: OrderedDict = OrderedDict()  # msg_id → timestamp，有序淘汰
_seen_lock = threading.Lock()
_MAX_SEEN = 500


def _is_new_message(msg_id: str) -> bool:
    if not msg_id:
        return True
    with _seen_lock:
        if msg_id in _seen_msg_ids:
            return False
        _seen_msg_ids[msg_id] = time.time()
        # 淘汰最旧的条目（FIFO，OrderedDict 保证插入顺序）
        while len(_seen_msg_ids) > _MAX_SEEN:
            _seen_msg_ids.popitem(last=False)
    return True


def _verify_signature(timestamp: str, nonce: str, body: bytes) -> bool:
    token = str(FEISHU_CONFIG.get("verify_token") or "").strip()
    incoming = request.headers.get("X-Lark-Signature", "")
    if not all((token, timestamp, nonce, incoming)):
        return False
    key = (timestamp + nonce + token).encode("utf-8") + body
    sig = hashlib.sha256(key).hexdigest()
    return hmac.compare_digest(sig, incoming)


# ════════════════════════════════════════════════════════════
# SSRF 防护
# ════════════════════════════════════════════════════════════
_SSRF_BLOCKED_NETS = [
    ipaddress.ip_network("127.0.0.0/8"),
    ipaddress.ip_network("10.0.0.0/8"),
    ipaddress.ip_network("172.16.0.0/12"),
    ipaddress.ip_network("192.168.0.0/16"),
    ipaddress.ip_network("169.254.0.0/16"),
    ipaddress.ip_network("::1/128"),
    ipaddress.ip_network("fc00::/7"),
]


def _is_blocked_addr(addr: str) -> tuple[bool, str]:
    ip = ipaddress.ip_address(addr)
    if (
        ip.is_loopback or ip.is_private or ip.is_link_local or ip.is_multicast
        or ip.is_reserved or ip.is_unspecified
    ):
        return True, str(ip)
    for net in _SSRF_BLOCKED_NETS:
        if ip in net:
            return True, str(ip)
    return False, ""


def _is_ssrf_safe(url: str) -> tuple:
    try:
        parsed = urlparse(url)
        if parsed.scheme not in ("http", "https"):
            return False, f"不允许的协议: {parsed.scheme}"
        host = parsed.hostname or ""
        if not host:
            return False, "无效的主机名"
        try:
            blocked, blocked_addr = _is_blocked_addr(host)
            if blocked:
                return False, f"禁止访问私有地址: {blocked_addr}"
        except ValueError:
            pass
        if host.lower() in {"localhost", "localhost.localdomain", "ip6-localhost"}:
            return False, f"禁止访问本地主机: {host}"
        try:
            port = parsed.port or (443 if parsed.scheme == "https" else 80)
            infos = socket.getaddrinfo(host, port, type=socket.SOCK_STREAM)
        except OSError as e:
            return False, f"DNS解析失败: {e}"
        for info in infos:
            addr = info[4][0]
            blocked, blocked_addr = _is_blocked_addr(addr)
            if blocked:
                return False, f"域名解析到私有地址: {blocked_addr}"
        return True, ""
    except Exception as e:
        return False, f"URL解析失败: {e}"


# ════════════════════════════════════════════════════════════
# 飞书 API
# ════════════════════════════════════════════════════════════
_token_lock = threading.Lock()
_token_cache = {"token": "", "expire": 0}


def _get_tenant_token() -> str:
    """获取飞书 tenant_access_token，带缓存、锁和重试。
    失败时清空缓存防止中毒，指数退避重试 3 次。"""
    with _token_lock:
        if time.time() < _token_cache["expire"] - 120:
            return _token_cache["token"]
        last_err = None
        for attempt in range(3):
            try:
                resp = requests.post(
                    f"{FEISHU_API}/auth/v3/tenant_access_token/internal",
                    json={"app_id": FEISHU_CONFIG["app_id"],
                          "app_secret": FEISHU_CONFIG["app_secret"]},
                    timeout=10,
                )
                data = resp.json()
                if data.get("code") != 0:
                    raise RuntimeError(f"飞书Token获取失败: {data.get('msg')}")
                _token_cache["token"] = data["tenant_access_token"]
                _token_cache["expire"] = time.time() + data.get("expire", 7200)
                return _token_cache["token"]
            except Exception as e:
                last_err = e
                # 清空缓存，避免后续请求使用失效 token
                _token_cache["token"] = ""
                _token_cache["expire"] = 0
                logger.warning("tenant_token 获取失败 (attempt %d/3): %s", attempt + 1, e)
                if attempt < 2:
                    time.sleep(1 * (attempt + 1))  # 1s, 2s 退避
        raise RuntimeError(f"飞书Token获取失败（3次重试后放弃）: {last_err}")


def _feishu_post(path: str, payload: dict) -> dict:
    token = _get_tenant_token()
    r = requests.post(
        f"{FEISHU_API}{path}",
        headers={"Authorization": f"Bearer {token}", "Content-Type": "application/json"},
        json=payload,
        timeout=20,
    )
    data = r.json()
    code = data.get("code", -1)
    if code != 0:
        logger.warning("飞书 API %s 返回错误 code=%s msg=%s", path, code, data.get("msg"))
    return data


def send_text(chat_id: str, text: str) -> bool:
    """发送文本消息，返回是否成功"""
    try:
        data = _feishu_post("/im/v1/messages?receive_id_type=chat_id", {
            "receive_id": chat_id,
            "msg_type": "text",
            "content": json.dumps({"text": text}),
        })
        return data.get("code") == 0
    except Exception as e:
        logger.error("send_text 失败: %s", e)
        return False


def send_card(chat_id: str, card: dict) -> bool:
    """发送卡片消息，返回是否成功"""
    try:
        data = _feishu_post("/im/v1/messages?receive_id_type=chat_id", {
            "receive_id": chat_id,
            "msg_type": "interactive",
            "content": json.dumps(card),
        })
        return data.get("code") == 0
    except Exception as e:
        logger.error("send_card 失败: %s", e)
        return False


# ════════════════════════════════════════════════════════════
# 网页抓取
# ════════════════════════════════════════════════════════════

def _read_limited_response(resp: requests.Response, max_bytes: int = MAX_FETCH_BYTES) -> requests.Response:
    length = resp.headers.get("Content-Length")
    if length:
        try:
            if int(length) > max_bytes:
                resp.close()
                raise requests.RequestException(f"响应体过大（超过 {max_bytes // 1024 // 1024}MB）")
        except ValueError:
            pass
    chunks, total = [], 0
    try:
        for chunk in resp.iter_content(chunk_size=64 * 1024):
            if not chunk:
                continue
            total += len(chunk)
            if total > max_bytes:
                resp.close()
                raise requests.RequestException(f"响应体过大（超过 {max_bytes // 1024 // 1024}MB）")
            chunks.append(chunk)
        resp._content = b"".join(chunks)
        resp._content_consumed = True
        return resp
    finally:
        resp.close()


_SSRF_HTTP_LOCAL = threading.local()


def _ssrf_http_session() -> requests.Session:
    """One direct, proxy-free Session per worker thread for untrusted URLs."""
    session = getattr(_SSRF_HTTP_LOCAL, "session", None)
    if session is None:
        session = requests.Session()
        session.trust_env = False
        _SSRF_HTTP_LOCAL.session = session
    return session


def _connected_peer_ip(resp: requests.Response) -> str:
    raw = getattr(resp, "raw", None)
    candidates = [
        getattr(getattr(raw, "_connection", None), "sock", None),
        getattr(
            getattr(getattr(getattr(raw, "_fp", None), "fp", None), "raw", None),
            "_sock",
            None,
        ),
    ]
    for candidate in candidates:
        if candidate is None:
            continue
        try:
            peer = candidate.getpeername()
            if peer and peer[0]:
                return str(peer[0])
        except (AttributeError, OSError, TypeError):
            continue
    raise requests.RequestException("无法核验远端连接地址")


def _validate_connected_peer(resp: requests.Response) -> None:
    try:
        peer_ip = _connected_peer_ip(resp)
        blocked, blocked_addr = _is_blocked_addr(peer_ip)
    except Exception:
        resp.close()
        raise
    if blocked:
        resp.close()
        raise requests.RequestException(f"连接被重绑定到私有地址: {blocked_addr}")


def _safe_get_once(url: str, headers: dict, timeout: int, max_bytes: int = MAX_FETCH_BYTES) -> requests.Response:
    current = url
    for redirect_idx in range(MAX_REDIRECTS + 1):
        safe, reason = _is_ssrf_safe(current)
        if not safe:
            raise requests.RequestException(f"URL不安全: {reason}")
        resp = _ssrf_http_session().get(
            current, headers=headers, timeout=timeout,
            allow_redirects=False, stream=True,
        )
        _validate_connected_peer(resp)
        if 300 <= resp.status_code < 400:
            location = resp.headers.get("Location")
            resp.close()
            if not location:
                return resp
            if redirect_idx >= MAX_REDIRECTS:
                raise requests.TooManyRedirects(f"重定向超过 {MAX_REDIRECTS} 次")
            current = urljoin(current, location)
            continue
        return _read_limited_response(resp, max_bytes=max_bytes)
    raise requests.TooManyRedirects(f"重定向超过 {MAX_REDIRECTS} 次")


def _fetch_url(url: str, timeout: int = 15) -> requests.Response:
    """带备用UA重试和SSRF/响应体限制的HTTP GET。"""
    uas = [
        _BROWSER_HEADERS["User-Agent"],
        ("Mozilla/5.0 (Macintosh; Intel Mac OS X 10_15_7) "
         "AppleWebKit/605.1.15 (KHTML, like Gecko) Version/17.2 Safari/605.1.15"),
    ]
    last_exc = None
    for ua in uas:
        try:
            headers = {**_BROWSER_HEADERS, "User-Agent": ua}
            r = _safe_get_once(url, headers=headers, timeout=timeout)
            if r.status_code in (403, 429):
                last_exc = requests.HTTPError(f"{r.status_code}", response=r)
                continue
            r.raise_for_status()
            return r
        except (requests.Timeout, requests.ConnectionError) as e:
            last_exc = e
            continue
        except requests.HTTPError:
            raise
    if last_exc:
        raise last_exc
    raise requests.RequestException(f"Failed to fetch {url}")


def _extract_url_content(url: str) -> dict:
    """抓取URL页面，提取标题和正文（trafilatura 优先，BeautifulSoup 兜底）"""
    r = _fetch_url(url)
    r.encoding = r.apparent_encoding or "utf-8"
    html = r.text

    title, body, pub_date, source = "", "", "", ""

    # ── trafilatura 提取（成功率高，支持微信文章、付费墙等）──
    if _HAS_TRAFILATURA:
        try:
            result = trafilatura.extract(html, include_comments=False, include_tables=False,
                                         output_format="json", with_metadata=True)
            if result:
                meta = json.loads(result)
                body = meta.get("text", "")
                title = meta.get("title", "")
                pub_date = meta.get("date", "")
                source = meta.get("sitename", "")
        except Exception as e:
            logger.debug("trafilatura failed for %s: %s", url, e)

    # ── BeautifulSoup 兜底 ──
    if not body or len(body) < 50:
        soup = BeautifulSoup(html, "html.parser")
        for tag in soup(["script", "style", "nav", "footer", "header", "aside", "form", "iframe"]):
            tag.decompose()
        if not title:
            og_title = soup.find("meta", property="og:title")
            if og_title and og_title.get("content"):
                title = og_title["content"].strip()
            if not title and soup.title:
                title = soup.title.get_text(strip=True)
            if not title:
                h1 = soup.find("h1")
                if h1:
                    title = h1.get_text(strip=True)
        article_tag = soup.find("article")
        paragraphs = (article_tag or soup).find_all("p")
        text_parts = [p.get_text(strip=True) for p in paragraphs if len(p.get_text(strip=True)) > 20]
        bs_body = "\n".join(text_parts)
        if len(bs_body) > len(body):
            body = bs_body
        if not pub_date:
            for meta_name in ["article:published_time", "datePublished", "pubdate", "date"]:
                tag = soup.find("meta", attrs={"property": meta_name}) or soup.find("meta", attrs={"name": meta_name})
                if tag and tag.get("content"):
                    pub_date = tag["content"].strip()
                    break
            if not pub_date:
                time_tag = soup.find("time")
                if time_tag and time_tag.get("datetime"):
                    pub_date = time_tag["datetime"]
        if not source:
            og_site = soup.find("meta", property="og:site_name")
            if og_site and og_site.get("content"):
                source = og_site["content"].strip()

    if not source:
        source = urlparse(url).netloc.replace("www.", "")
    if len(body) > 5000:
        body = body[:5000] + "……（内容已截断）"
    return {"title": title, "body": body, "pub_date": pub_date, "source": source, "url": url}


# ════════════════════════════════════════════════════════════
# AI 调用
# ════════════════════════════════════════════════════════════

def _is_anthropic_endpoint() -> bool:
    return "api.anthropic.com" in AI_CONFIG["base_url"].lower()


def _call_ai(messages: list, temperature: float = 0.4) -> str:
    """调用 LLM API（OpenAI兼容 / Anthropic原生）"""
    if not AI_CONFIG["api_key"]:
        raise ValueError("AI API Key 未配置")

    if _is_anthropic_endpoint():
        headers = {
            "x-api-key": AI_CONFIG["api_key"],
            "anthropic-version": "2023-06-01",
            "content-type": "application/json",
        }
        system_msg = ""
        user_msgs = []
        for m in messages:
            if m["role"] == "system":
                system_msg = (system_msg + "\n\n" + m["content"]).strip()
            else:
                user_msgs.append({"role": m["role"], "content": m["content"]})
        payload = {
            "model": AI_CONFIG["model"],
            "messages": user_msgs,
            "max_tokens": AI_CONFIG["max_tokens"],
            "temperature": temperature,
        }
        if system_msg:
            payload["system"] = system_msg
        url = AI_CONFIG["base_url"].rstrip("/") + "/v1/messages"
        resp = requests.post(url, headers=headers, json=payload, timeout=180)
        resp.raise_for_status()
        content = resp.json().get("content", [])
        if isinstance(content, list) and content:
            return content[0].get("text", "")
        return str(resp.json())

    # OpenAI 兼容（DeepSeek / Qwen / GLM / OpenRouter 等）
    headers = {
        "Authorization": f"Bearer {AI_CONFIG['api_key']}",
        "Content-Type": "application/json",
    }
    payload = {
        "model": AI_CONFIG["model"],
        "messages": messages,
        "max_tokens": AI_CONFIG["max_tokens"],
        "temperature": temperature,
    }
    base = AI_CONFIG["base_url"].rstrip("/")
    if base.endswith("/v1"):
        url = base + "/chat/completions"
    elif "/v1" in base or "/paas/v4" in base or "/compatible-mode" in base or "/api/v3" in base:
        url = base + "/chat/completions"
    else:
        url = base + "/v1/chat/completions"

    _ssl_verify = not any(h in base for h in ["simpleai.com.cn"])
    resp = requests.post(url, headers=headers, json=payload, timeout=180, verify=_ssl_verify)
    resp.raise_for_status()
    result = resp.json()
    choices = result.get("choices")
    if not choices or not isinstance(choices, list):
        logger.error("AI 响应缺少 choices: %s", str(result)[:300])
        raise ValueError("AI 返回格式异常：缺少 choices 字段")
    msg = choices[0].get("message", {})
    text = msg.get("content") or msg.get("reasoning_content") or ""
    if not text.strip():
        logger.warning("AI 返回空内容: finish_reason=%s", choices[0].get("finish_reason"))
    return text


# ════════════════════════════════════════════════════════════
# 飞书文件下载（图片 / PDF）
# ════════════════════════════════════════════════════════════

def _download_feishu_image(image_key: str) -> bytes:
    """通过飞书 API 下载图片"""
    token = _get_tenant_token()
    r = requests.get(
        f"{FEISHU_API}/im/v1/images/{image_key}",
        headers={"Authorization": f"Bearer {token}"},
        timeout=15,
        stream=True,
    )
    r.raise_for_status()
    r = _read_limited_response(r, max_bytes=MAX_FEISHU_DOWNLOAD_BYTES)
    return r.content


def _download_feishu_file(message_id: str, file_key: str) -> bytes:
    """通过飞书 API 下载消息中的文件"""
    token = _get_tenant_token()
    r = requests.get(
        f"{FEISHU_API}/im/v1/messages/{message_id}/resources/{file_key}?type=file",
        headers={"Authorization": f"Bearer {token}"},
        timeout=30,
        stream=True,
    )
    r.raise_for_status()
    r = _read_limited_response(r, max_bytes=MAX_FEISHU_DOWNLOAD_BYTES)
    return r.content


def _extract_pdf_text(pdf_bytes: bytes) -> str:
    """从 PDF 二进制数据提取文本"""
    if not _HAS_PDFPLUMBER:
        return ""
    text_parts = []
    with pdfplumber.open(BytesIO(pdf_bytes)) as pdf:
        for page in pdf.pages[:20]:
            page_text = page.extract_text()
            if page_text:
                text_parts.append(page_text)
    text = "\n".join(text_parts)
    return text[:8000] if text else ""


def _extract_docx_text(docx_bytes: bytes) -> str:
    """从 DOCX 二进制数据提取文本（python-docx）"""
    doc = DocxDocument(BytesIO(docx_bytes))
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    text = "\n".join(paragraphs)
    return text[:8000] if text else ""


def _extract_text_file(file_bytes: bytes) -> str:
    """从纯文本/Markdown文件提取文本，自动检测编码"""
    for encoding in ("utf-8-sig", "utf-8", "gbk", "gb2312", "utf-16"):
        try:
            return file_bytes.decode(encoding)[:8000]
        except (UnicodeDecodeError, LookupError):
            continue
    return file_bytes.decode("utf-8", errors="replace")[:8000]


def _call_ai_with_image(
    image_b64: str,
    media_type: str,
    prompt: str,
    system_prompt: str = "",
) -> str:
    """调用多模态 AI（图片+文字），支持 Anthropic 和 OpenAI 格式"""
    if not AI_CONFIG["api_key"]:
        raise ValueError("AI API Key 未配置")

    if _is_anthropic_endpoint():
        headers = {
            "x-api-key": AI_CONFIG["api_key"],
            "anthropic-version": "2023-06-01",
            "content-type": "application/json",
        }
        payload = {
            "model": AI_CONFIG["model"],
            "max_tokens": AI_CONFIG["max_tokens"],
            "temperature": 0.4,
            "messages": [{
                "role": "user",
                "content": [
                    {"type": "image", "source": {"type": "base64", "media_type": media_type, "data": image_b64}},
                    {"type": "text", "text": prompt},
                ]
            }],
        }
        if system_prompt:
            payload["system"] = system_prompt
        url = AI_CONFIG["base_url"].rstrip("/") + "/v1/messages"
        resp = requests.post(url, headers=headers, json=payload, timeout=180)
        resp.raise_for_status()
        content = resp.json().get("content", [])
        if isinstance(content, list) and content:
            return content[0].get("text", "")
        return str(resp.json())
    else:
        # OpenAI 兼容格式
        headers = {"Authorization": f"Bearer {AI_CONFIG['api_key']}", "Content-Type": "application/json"}
        messages = []
        if system_prompt:
            messages.append({"role": "system", "content": system_prompt})
        messages.append({
            "role": "user",
            "content": [
                {"type": "image_url", "image_url": {"url": f"data:{media_type};base64,{image_b64}"}},
                {"type": "text", "text": prompt},
            ],
        })
        payload = {
            "model": AI_CONFIG["model"],
            "max_tokens": AI_CONFIG["max_tokens"],
            "temperature": 0.4,
            "messages": messages,
        }
        base = AI_CONFIG["base_url"].rstrip("/")
        url = base + "/v1/chat/completions" if not base.endswith("/v1") else base + "/chat/completions"
        resp = requests.post(url, headers=headers, json=payload, timeout=180,
                             verify=not any(h in base for h in ["simpleai.com.cn"]))
        resp.raise_for_status()
        result = resp.json()
        choices = result.get("choices")
        if not choices or not isinstance(choices, list):
            logger.error("多模态 AI 响应缺少 choices: %s", str(result)[:300])
            raise ValueError("AI 返回格式异常：缺少 choices 字段")
        return choices[0].get("message", {}).get("content", "")


# ════════════════════════════════════════════════════════════
# 要讯 System Prompt & 用户 Prompt
# ════════════════════════════════════════════════════════════

SYSTEM_PROMPT_BRIEF_WRITE = """你是一名资深中文防务资讯编辑，长期为军事媒体撰写基于公开信息源的防务要讯（open-source defense intelligence summary）。所有素材均来自公开媒体报道，你的任务是将公开防务新闻改写为中国军事媒体常见的机关行文风格综述。

════════════════════════════════════════
【核心文风要求】中文军事媒体机关行文体
════════════════════════════════════════
必用军语词汇（自然嵌入行文）：
• 开头句式：统一使用"据XX报道，"；XX必须与信息来源行一致，不在"报道"前添加日期，禁用"近日""近期"
• 研判用语："值得警惕""值得关注""须警惕""研判""着力""亟需""显著提升""根本性威胁""现实压力"
• 建议用语："建议持续跟踪""加强""积极参与""着力构建""针对性加强""掌握战略主动""争取战略主动"
• 战略词汇："战略制高点""战略间隙""战略主动""战略支援""战略意图""根本性威胁""颠覆性威胁"
• 军事词汇："前沿部署""实战化""规模化""常态化""态势感知""反隐形""防空反导""空中安全""岛链"

避免使用：
✗ 口语词（"搞""弄""厉害"）
✗ 网络用语（"牛""卷""破防"）
✗ 主观情绪词（"震惊""愤怒""觉得"）
✗ markdown符号（#、*、-、【】以外的）

════════════════════════════════════════
【输出格式】严格按此六部分输出（参照素材1-5通用模板）
════════════════════════════════════════
第一行：事件时间：YYYY年M月D日（必须是原文明确对应的具体事件日期，不得写"近日""近期"）
第二行：价 值 点：<一句话，60字内，指出核心研判与战略意义；不得复制标题，去掉标题末尾警示词后也不得近乎原样复述>
第三行：（空行）
第四行：<标题：8-15字，主语明确，不得含中文逗号"，"或英文逗号","，只能以"值得警惕"或"值得关注"收尾>
第五行：（空行）
第六行开始：<正文：单段成文，250-350字，结构如下>
     据<与信息来源行第一条一致的具体信源>报道，<帽段：用3-4行、约80-120字简述时间+主体+动作+装备数量+地点+目的，且必须出现与事件时间一致的M月D日>。（1）<影响一：对我战略/装备/力量的直接影响，35-55字>。（2）<影响二：对区域态势/盟体/对手的影响，35-55字>。（3）<影响三：深层意图/长远威胁研判，35-55字>。建议持续跟踪<对象>的<要素一>、<要素二>及<要素三>，针对性加强<能力一>、<能力二>及<能力三>能力建设。
     也可不用编号，写成：<帽段>。<层意一>；<层意二>；<层意三>。建议……。不用编号时各层之间使用中文分号。
（空行）
倒数第二行：（信息来源：<来源一>M月D日发文《<原报道标题>》；<来源二>M月D日发文《<原报道标题>》）
末行：报送人：           电话：

信源处理要求：正文"据XX报道，"中的XX必须与信息来源行第一条名称完全一致。若公众号转述外网消息，优先采用素材中可核验的外网第一信源；无法找到第一信源时，写"据XX公众号报道，"，不得添加报道日期。凡实际引用的来源均须在信息来源行写全，包括"路透社称""XX指出"等二次来源；每条严格采用"XX M月D日发文《标题》"格式，多个来源之间使用中文分号"；"。

════════════════════════════════════════
【示范样本】参考此风格、用词、节奏
════════════════════════════════════════
样本1（标准(1)(2)(3)格式）：
事件时间：2026年3月24日
价 值 点：俄方推进低轨核能力研发，可能削弱我卫星导航与通信中继体系韧性，须持续研判技术进展及美方借题施压意图。

俄低轨核爆风险值得警惕

据美防务一号网站报道，3月24日，美参议院军事委员会举行听证会，美战略及太空司令部领导人证实俄罗斯正推进太空核能力研发。相关武器若在低轨引爆，可能无差别毁伤近地轨道航天资产。美欧围绕"核保护伞"可信度亦存在明显分歧。（1）该武器一旦实战化，将对我卫星导航、侦察预警、通信中继等战略支援能力构成根本性威胁。（2）美欧盟体协调失序，客观上为我战略运筹提供窗口期。（3）需警惕美方借此向国会争取经费、对俄施压的双重意图，须辩证研判其信息真实性与战略目的。建议持续跟踪俄太空核能力研发动态及部署进展，针对性加强我太空资产抗毁性与快速补网能力建设。

（信息来源：美防务一号网站3月26日发文《参院军事委员会主席：美国国防战略在核与太空威胁问题上"存在不足"》）
报送人：           电话：

样本2（前沿部署类）：
事件时间：2026年3月28日
价 值 点：美军以F-35A替换F-16进驻三泽，第一岛链隐形打击与态势感知能力上升，将加大我东北亚方向防空反隐形压力。

美军隐形战机前沿部署值得关注

据比利时陆军防务网报道，3月28日，美军首批F-35A隐形战斗机抵达日本三泽空军基地并开始替换F-16。美方投入超100亿美元升级相关基础设施。该机具备隐形、传感器融合及多任务能力，可执行防空压制、精确打击和盟军协同作战。（1）F-35A前沿部署将显著压缩我防空识别区反应时间，增大我周边空中安全压力。（2）三泽成为美日共用F-35平台前沿基地，明显提升美在东北亚的隐形打击与态势感知能力。（3）美方明确称此举针对中国在东海等地日益常态化的军事活动，遏华意图凸显。建议持续跟踪该机在三泽的部署规模及训练强度，针对性加强反隐形侦察、区域防空及电子对抗能力建设。

（信息来源：比利时陆军防务网3月30日发文《美国向日本部署F-35A隐形战斗机，取代F-16以应对中国威胁》）
报送人：           电话：

════════════════════════════════════════
【媒体名称中文对照表】正文和信息来源行必须使用中文名
════════════════════════════════════════
• Defense News / defensenews.com → 美国防务新闻
• Defense One / defenseone.com / defenseonc.com → 美防务一号网站
• Breaking Defense / breakingdefense.com → 美突破防务网
• USNI News / usni.org → 美海军学会新闻网
• War on the Rocks / warontherocks.com → 美岩石上的战争网
• The National Interest / nationalinterest.org → 美国家利益网
• The Diplomat / thediplomat.com → 外交学者网
• Jane's / janes.com → 简氏防务周刊
• Army Recognition / armyrecognition.com → 比利时陆军防务网
• Aviation Week / aviationweek.com → 航空周刊
• Flight Global / flightglobal.com → 全球飞行网
• Stars and Stripes / stripes.com → 美星条旗报
• Military Times / militarytimes.com → 美军事时报
• C4ISRNET / c4isrnet.com → 美指挥控制情报侦察网
• Politico / politico.com → 美政治新闻网
• The Hill / thehill.com → 美国山丘报
• Reuters / reuters.com → 路透社
• Bloomberg / bloomberg.com → 彭博社
• AP / apnews.com → 美联社
• CNN / cnn.com → 美国有线电视新闻网
• BBC / bbc.com / bbc.co.uk → 英国广播公司
• South China Morning Post / scmp.com → 香港南华早报
• Kyodo News / kyodonews.net → 日本共同社
• Yonhap / yna.co.kr → 韩联社
• Nikkei / nikkei.com → 日本经济新闻
• Financial Times / ft.com → 英国金融时报
• Jomhouri Eslami → 伊朗伊斯兰共和国报
• 其他未列出的英文媒体：根据国家+性质意译，如 "taiwanese outlet X" → "台湾X媒体"

════════════════════════════════════════
【硬性红线】违反任一条重写
════════════════════════════════════════
1. 必须严格六部分输出（事件时间/价值点/标题/正文/信息来源/报送人电话行），不得增减
2. 事件时间必须是原文明确支持的实际事件日期并具体到年月日，不得使用"近日""近期"，不得用来源发布日期或系统今日日期兜底
3. 价值点必须另行提炼战略意义，不得复制标题；去掉标题末尾"值得警惕/值得关注"后，也不得近乎原样复述标题主体
4. 标题必须为8-15字，不得含中英文逗号，只能以"值得警惕"或"值得关注"收尾
5. 正文必须单段成文；开头帽段用3-4行、约80-120字简述基本情况，并写出与事件时间一致的M月D日
6. 使用（1）（2）（3）分层时各层之间必须用句号；不用编号时写成"帽段。层意一；层意二；层意三。建议……"，各层之间用中文分号
7. 正文总字数控制在250-350字
8. 正文"据XX报道，"中的XX必须与信息来源第一条一致且使用中文媒体名称，严禁出现英文域名或英文媒体名
9. 公众号转述外网消息时优先采用可核验的外网第一信源；找不到时写"据XX公众号报道，"且不加日期
10. 信息来源每条必须写成"XX M月D日发文《标题》"，实际引用的来源全部列全，含"路透社称""XX指出"等二次来源；多个来源用中文分号分隔
11. 不得使用任何markdown符号（#、*、-、**等）
12. 必须保持PLA机关军语文风，不得口语化
13. 不得脱离原文编造事实数据、来源或日期
14. 建议必须严格采用"建议持续跟踪X的要素一、要素二、要素三，针对性加强能力一、能力二、能力三能力建设"范式
15. 末行必须输出"报送人：           电话："（留空待填）
16. 信息来源行《》内的文章标题必须翻译为中文，严禁保留英文原标题"""


def _build_user_prompt(title: str, body: str, source: str = "", url: str = "", pub_date: str = "") -> str:
    now = datetime.now()
    today_cn = f"{now.year:04d}年{now.month:02d}月{now.day:02d}日"
    date_cn = "未提供"
    pub_md = "未提供"
    dt = _brief_parse_date_value(pub_date)
    if dt:
        date_cn = f"{dt.year:04d}年{dt.month:02d}月{dt.day:02d}日"
        pub_md = f"{dt.month:02d}月{dt.day:02d}日"
    source_label = source if source else "素材中明确标注的来源"
    return f"""请根据以下导入素材，撰写一份PLA机关军语要讯（情报简报）：

════════ 导入素材 ════════
【素材标题】{title}
【信息来源】{source_label}
【来源发布日期】{date_cn}
【素材正文】
{body}
【原文链接】{url if url else "（用户导入）"}

════════ 今日日期 ════════
{today_cn}

════════ 写作任务 ════════
请输出一份要讯，严格遵循以下要求：
1. 事件时间只能填写素材原文明确支持的实际事件日期，必须具体到年月日，不得写"近日""近期"，不得把来源发布日期{date_cn}或今日日期{today_cn}当作事件时间兜底；原文未给实际事件日期时必须先核实，不得臆造
2. 价值点必须另行提炼核心研判和战略意义，不得复制标题；去掉标题末尾"值得警惕/值得关注"后，也不得近乎原样复述标题主体
3. 标题控制在8-15字，不得含中文逗号"，"或英文逗号","，只能以"值得警惕"或"值得关注"收尾
4. 正文统一以"据【中文媒体名】报道，"开头，媒体名必须与信息来源第一条一致，不得在"报道"前添加日期。若来源"{source_label}"是英文域名或英文名，必须按对照表译为中文，严禁直接写英文
5. 若素材来自公众号且转述外网消息，优先采用素材中可核验的外网第一信源；无法找到第一信源时写"据XX公众号报道，"，不加日期
6. 正文必须单段成文、250-350字；开头帽段用3-4行、约80-120字简述时间、主体、动作、装备数量、地点和目的，并写出与事件时间一致的M月D日
7. 使用（1）（2）（3）分层时各层之间用句号；不用编号时写成"帽段。层意一；层意二；层意三。建议……"并以中文分号分层；结尾建议必须采用"建议持续跟踪X的要素一、要素二、要素三，针对性加强能力一、能力二、能力三能力建设"的范式
8. 末尾信息来源每条严格填写为"XX M月D日发文《标题》"；当前来源日期线索为{pub_md}，须与原文核对后使用；凡实际引用的来源全部写全，含"路透社称""XX指出"等二次来源，多个来源之间使用中文分号"；"
9. 《》内文章标题无论原文是否为英文均须翻译为中文；素材为外文时，正文、标题、价值点和信息来源行全部写中文，专有名词如武器型号可保留
10. 使用PLA机关军语，从素材提炼对我军/对华影响，不得编造素材未提及的具体数据、来源或日期

直接输出要讯全文，不要任何解释说明。"""


# ════════════════════════════════════════════════════════════
# 卡片构建
# ════════════════════════════════════════════════════════════

_BRIEF_RELATIVE_EVENT_WORDS = (
    "近期", "近日", "日前", "最近", "当前", "本月", "上月", "本周", "上周",
    "今年", "去年", "今日", "昨日", "昨天", "明日",
)

_BRIEF_TITLE_ENDINGS = ("值得警惕", "值得关注")

_BRIEF_EVENT_DATE_RE = re.compile(
    r"(?P<year>\d{4})年(?P<month>\d{1,2})月(?P<day>\d{1,2})日"
)

_BRIEF_SOURCE_ENTRY_RE = re.compile(
    r"^(?P<name>.+?)\s*(?P<month>\d{1,2})月(?P<day>\d{1,2})日发文《(?P<title>[^》]+)》$"
)

_BRIEF_BODY_ATTRIBUTION_RE = re.compile(r"据(?P<label>[^，,。；;]{1,80}?)报道")

_BRIEF_SECONDARY_SOURCE_RE = re.compile(
    r"(?:^|[，,。；;！？!?）)])(?:同时|另据|此外|而|但)?"
    r"(?P<name>[\u4e00-\u9fffA-Za-z0-9·]{0,24}?"
    r"(?:广播公司|通讯社|电视台|研究院|研究所|新闻网|公众号|网站|周刊|杂志|媒体|智库|中心|网|报|社))"
    r"(?:称|指出|表示|披露|证实|认为|报道(?:称)?|援引)"
)

_BRIEF_CONTEXT_SOURCE_RE = re.compile(
    r"(?:另据|根据|援引|据)(?P<name>[^，,。；;]{2,40}?)"
    r"(?:的)?(?:报道|消息|声明|数据|报告)(?:显示|称|指出|披露|证实|，|,|。|；|;|$)"
)

_BRIEF_RECIPIENT_SOURCE_RE = re.compile(
    r"(?:消息人士|官员|知情人士|发言人)向(?P<name>[^，,。；;]{2,40}?)(?:表示|透露|称)"
)

_BRIEF_ATTRIBUTION_DATE_SUFFIX_RE = re.compile(
    r"(?P<date>[（(]?(?:(?:\d{4}\s*年\s*)?\d{1,2}\s*月\s*\d{1,2}\s*[日号]|"
    r"(?:\d{4}\s*[-/.]\s*)?\d{1,2}\s*[-/.]\s*\d{1,2})[）)]?)$"
)

def _brief_compact(text: str) -> str:
    return re.sub(r"[\s，,。；;：:！？!?（）()《》“”\"'、]", "", text or "")

_BRIEF_MEDIA_ALIAS_GROUPS = (
    ("美国防务新闻", "防务新闻", "Defense News", "defensenews.com"),
    ("美防务一号网站", "防务一号", "Defense One", "defenseone.com"),
    ("美突破防务网", "突破防务", "Breaking Defense", "breakingdefense.com"),
    ("美海军学会新闻网", "USNI News", "news.usni.org", "usni.org"),
    ("路透社", "Reuters", "reuters.com"),
    ("美联社", "AP", "Associated Press", "apnews.com"),
    ("彭博社", "Bloomberg", "bloomberg.com"),
    ("英国广播公司", "BBC", "bbc.com", "bbc.co.uk"),
    ("香港南华早报", "南华早报", "South China Morning Post", "scmp.com"),
    ("日本共同社", "共同社", "Kyodo News", "kyodonews.net"),
    ("韩联社", "Yonhap", "yna.co.kr"),
    ("日本经济新闻", "Nikkei", "Nikkei Asia", "nikkei.com"),
    ("英国金融时报", "金融时报", "Financial Times", "ft.com"),
)

_BRIEF_ENGLISH_MONTHS = (
    (1, "January", "Jan"), (2, "February", "Feb"), (3, "March", "Mar"),
    (4, "April", "Apr"), (5, "May", "May"), (6, "June", "Jun"),
    (7, "July", "Jul"), (8, "August", "Aug"), (9, "September", "Sep", "Sept"),
    (10, "October", "Oct"), (11, "November", "Nov"), (12, "December", "Dec"),
)

def _brief_parse_date_value(value) -> datetime | None:
    if isinstance(value, datetime):
        return value
    raw = str(value or "").strip()
    if not raw:
        return None
    try:
        return datetime.fromisoformat(raw.replace("Z", "+00:00"))
    except ValueError:
        pass
    normalized = re.sub(r"\s+", " ", raw)
    for fmt in (
        "%Y-%m-%d", "%Y/%m/%d", "%Y.%m.%d", "%Y年%m月%d日",
        "%B %d, %Y", "%b %d, %Y", "%d %B %Y", "%d %b %Y",
    ):
        try:
            return datetime.strptime(normalized, fmt)
        except ValueError:
            continue
    match = re.search(r"(?P<year>\d{4})[-/.年](?P<month>\d{1,2})[-/.月](?P<day>\d{1,2})(?:日)?", raw)
    if match:
        try:
            return datetime(int(match.group("year")), int(match.group("month")), int(match.group("day")))
        except ValueError:
            return None
    return None

def _brief_month_day_supported(text: str, month: int, day: int) -> bool:
    haystack = str(text or "")
    if re.search(rf"(?<!\d)0?{month}\s*月\s*0?{day}\s*[日号](?!\d)", haystack):
        return True
    if re.search(rf"(?<!\d)0?{month}\s*[-/.]\s*0?{day}(?!\d)", haystack):
        return True
    for number, *names in _BRIEF_ENGLISH_MONTHS:
        if number != month:
            continue
        month_names = "|".join(map(re.escape, names))
        if re.search(rf"\b(?:{month_names})\.?\s+0?{day}(?:st|nd|rd|th)?\b", haystack, re.I):
            return True
        if re.search(rf"\b0?{day}(?:st|nd|rd|th)?\s+(?:{month_names})\.?\b", haystack, re.I):
            return True
    return False

def _brief_event_date_supported(text: str, year: int, month: int, day: int,
                                publication_year: int | None = None) -> bool:
    haystack = str(text or "")
    if re.search(rf"(?<!\d){year}\s*年\s*0?{month}\s*月\s*0?{day}\s*[日号](?!\d)", haystack):
        return True
    if re.search(rf"(?<!\d){year}\s*[-/.]\s*0?{month}\s*[-/.]\s*0?{day}(?!\d)", haystack):
        return True
    for number, *names in _BRIEF_ENGLISH_MONTHS:
        if number != month:
            continue
        month_names = "|".join(map(re.escape, names))
        if re.search(rf"\b(?:{month_names})\.?\s+0?{day}(?:st|nd|rd|th)?,?\s+{year}\b", haystack, re.I):
            return True
        if re.search(rf"\b0?{day}(?:st|nd|rd|th)?\s+(?:{month_names})\.?,?\s+{year}\b", haystack, re.I):
            return True
    year_is_supported = publication_year == year or re.search(rf"(?<!\d){year}(?!\d)", haystack)
    return bool(year_is_supported and _brief_month_day_supported(haystack, month, day))

def _brief_source_aliases(name: str, url: str = "") -> set[str]:
    aliases = {str(name or "").strip()}
    host = (urlparse(url).hostname or "") if url else ""
    compact_name = _brief_compact(name).casefold()
    for group in _BRIEF_MEDIA_ALIAS_GROUPS:
        compact_group = {_brief_compact(item).casefold() for item in group}
        domains = {item.casefold() for item in group if "." in item}
        if compact_name in compact_group or any(host.casefold().endswith(domain) for domain in domains):
            aliases.update(group)
    return {item for item in aliases if item}

def _brief_name_supported_in_material(name: str, material_text: str) -> bool:
    material = _brief_compact(material_text).casefold()
    return any(
        _brief_compact(alias).casefold() in material
        for alias in _brief_source_aliases(name)
        if _brief_compact(alias)
    )

def _brief_evidence(*, material_text: str, source_name: str = "", source_title: str = "",
                    publication_date="", publication_date_verified: bool = False,
                    url: str = "") -> dict:
    parsed_publication_date = _brief_parse_date_value(publication_date)
    return {
        "material_text": str(material_text or ""),
        "source_name": str(source_name or "").strip(),
        "source_aliases": sorted(_brief_source_aliases(source_name, url)),
        "source_title": str(source_title or "").strip(),
        "publication_date": parsed_publication_date.isoformat() if parsed_publication_date else "",
        "publication_date_verified": bool(publication_date_verified and parsed_publication_date),
        "url": str(url or ""),
    }

def _brief_char_count(text: str) -> int:
    return len(re.sub(r"\s", "", text or ""))

def _parse_brief_for_validation(brief: str) -> dict:
    """严格解析六部分要讯；不使用DOCX解析器的自动填充兜底。"""
    lines = [line.rstrip() for line in (brief or "").strip().splitlines()]
    parsed = {
        "event_time": "", "value_point": "", "title": "",
        "body": "", "source": "", "reporter": "", "unexpected_lines": [],
    }
    title_lines, body_lines = [], []
    state = "meta"
    for line in lines:
        text = line.strip()
        if not text:
            if state == "meta":
                state = "title"
            elif state == "title" and title_lines:
                state = "body"
            continue
        if text.startswith("事件时间"):
            parsed["event_time"] = (
                text.split("：", 1)[1].strip()
                if "：" in text else text.replace("事件时间", "", 1).strip()
            )
            state = "meta"
            continue
        if text.startswith("价 值 点") or text.startswith("价值点"):
            parsed["value_point"] = text.split("：", 1)[1].strip() if "：" in text else ""
            state = "meta"
            continue
        if text.startswith("（信息来源") or text.startswith("(信息来源"):
            if parsed["source"]:
                parsed["unexpected_lines"].append(text)
            else:
                parsed["source"] = text
            state = "done"
            continue
        if text.startswith("报送人"):
            if state != "done" or parsed["reporter"]:
                parsed["unexpected_lines"].append(text)
            else:
                parsed["reporter"] = text
            state = "reported"
            continue
        if state in ("done", "reported"):
            parsed["unexpected_lines"].append(text)
            continue
        if state == "title":
            title_lines.append(text)
        elif state in ("meta", "body"):
            state = "body"
            body_lines.append(text)
    parsed["title"] = "".join(title_lines)
    parsed["body"] = "".join(body_lines)
    return parsed

def _parse_brief_source_entries(source: str) -> tuple[list[dict], list[str], str]:
    raw = (source or "").strip()
    raw = re.sub(r"^[（(]?\s*信息来源\s*[:：]\s*", "", raw)
    raw = re.sub(r"\s*[）)]\s*$", "", raw)
    parts = [part.strip() for part in re.split(r"[；;]", raw) if part.strip()]
    entries, invalid = [], []
    for part in parts:
        match = _BRIEF_SOURCE_ENTRY_RE.fullmatch(part)
        if not match:
            invalid.append(part)
            continue
        month, day = int(match.group("month")), int(match.group("day"))
        try:
            datetime(2000, month, day)
        except ValueError:
            invalid.append(part)
            continue
        entries.append({"name": match.group("name").strip(), "month": month, "day": day})
    return entries, invalid, raw

def _brief_body_attributions(body: str) -> list[dict]:
    attributions = []
    for match in _BRIEF_BODY_ATTRIBUTION_RE.finditer(body or ""):
        label = match.group("label").strip()
        date_match = _BRIEF_ATTRIBUTION_DATE_SUFFIX_RE.search(label)
        if date_match:
            name = label[:date_match.start()].strip()
            date = date_match.group("date")
        else:
            name, date = label, ""
        attributions.append({
            "name": name,
            "date": date,
            "has_relative_date": any(word in label for word in _BRIEF_RELATIVE_EVENT_WORDS),
        })
    return attributions

def _validate_brief_text(brief: str, evidence: dict | None = None) -> dict:
    """执行独立云端飞书发送前的轻量机械规则校验。"""
    parsed = _parse_brief_for_validation(brief)
    errors = []
    event_time = parsed["event_time"]
    value_point = parsed["value_point"]
    title = parsed["title"]
    body = parsed["body"]
    source = parsed["source"]
    reporter = parsed["reporter"]
    event_month_day = None
    event_ymd = None

    for field, label in (
        (event_time, "事件时间"), (value_point, "价值点"), (title, "标题"),
        (body, "正文"), (source, "信息来源"), (reporter, "报送人电话行"),
    ):
        if not field:
            errors.append(f"缺少{label}")

    if event_time:
        if any(word in event_time for word in _BRIEF_RELATIVE_EVENT_WORDS):
            errors.append("事件时间不得使用相对表述")
        date_match = _BRIEF_EVENT_DATE_RE.fullmatch(event_time.strip())
        if not date_match:
            errors.append("事件时间必须写成有效的YYYY年M月D日")
        else:
            try:
                year = int(date_match.group("year"))
                month = int(date_match.group("month"))
                day = int(date_match.group("day"))
                datetime(
                    year, month, day,
                )
                event_month_day = (month, day)
                event_ymd = (year, month, day)
            except ValueError:
                errors.append("事件时间不是有效日期")

    if title and value_point:
        compact_title = _brief_compact(title)
        title_core = title
        for ending in _BRIEF_TITLE_ENDINGS:
            if title_core.endswith(ending):
                title_core = title_core[:-len(ending)]
                break
        compact_core = _brief_compact(title_core)
        compact_value = _brief_compact(value_point)
        similarity = SequenceMatcher(None, compact_core, compact_value, autojunk=False).ratio() if compact_core else 0.0
        if (
            compact_title and compact_title in compact_value
            or compact_core and compact_value == compact_core
            or compact_core and compact_value.startswith(compact_core) and len(compact_value) - len(compact_core) < 8
            or compact_core and len(compact_value) <= len(compact_core) + 6 and similarity >= 0.85
        ):
            errors.append("价值点不得复制标题")

    if title:
        if not 8 <= len(title) <= 15:
            errors.append("标题必须控制在8-15字")
        if re.search(r"[，,]", title):
            errors.append("标题不得含中英文逗号")
        if not title.endswith(_BRIEF_TITLE_ENDINGS):
            errors.append("标题必须以值得警惕或值得关注收尾")

    body_chars = _brief_char_count(body)
    if body and not 250 <= body_chars <= 350:
        errors.append("正文必须控制在250-350字")

    numbered_body = body.replace("(1)", "（1）").replace("(2)", "（2）").replace("(3)", "（3）")
    markers = ("（1）", "（2）", "（3）")
    marker_counts = [numbered_body.count(marker) for marker in markers]
    hat = ""
    if body and any(marker_counts):
        if not all(count == 1 for count in marker_counts):
            errors.append("编号正文必须完整使用（1）（2）（3）分层")
        else:
            point_one = numbered_body.index("（1）")
            point_two = numbered_body.index("（2）")
            point_three = numbered_body.index("（3）")
            if not point_one < point_two < point_three:
                errors.append("正文（1）（2）（3）分层顺序错误")
            elif (
                not numbered_body[:point_two].rstrip().endswith("。")
                or not numbered_body[:point_three].rstrip().endswith("。")
            ):
                errors.append("编号各层之间必须使用句号")
            point_three_text = numbered_body[point_three + len("（3）"):]
            if "建议" in point_three_text and "。建议" not in point_three_text:
                errors.append("第（3）层与建议句之间必须使用句号")
            hat = numbered_body[:point_one].strip()
    elif body:
        suggest_start = body.rfind("。建议")
        layer_separators = [
            match.start() for match in re.finditer("；", body[:suggest_start])
        ] if suggest_start >= 0 else []
        first_layer_separator = layer_separators[-2] if len(layer_separators) >= 2 else -1
        hat_end = body.rfind("。", 0, first_layer_separator) if first_layer_separator >= 0 else -1
        if hat_end < 0 or suggest_start <= hat_end:
            errors.append("无编号正文须写成帽段。层意一；层意二；层意三。建议…")
        else:
            hat = body[:hat_end + 1].strip()
            unnumbered_analysis = body[hat_end + 1:suggest_start]
            unnumbered_layers = [part.strip() for part in unnumbered_analysis.split("；")]
            if (
                len(unnumbered_layers) < 3
                or not all(unnumbered_layers)
                or ";" in unnumbered_analysis
            ):
                errors.append("无编号正文各层之间必须使用中文分号")

    if hat:
        if not 80 <= _brief_char_count(hat) <= 120:
            errors.append("帽段必须控制在80-120字")
        if event_month_day:
            month, day = event_month_day
            if not re.search(rf"(?<!\d)0?{month}月0?{day}日", hat):
                errors.append("帽段必须出现与事件时间一致的M月D日")
    elif body:
        errors.append("无法识别帽段")

    if body and not all(
        phrase in body for phrase in ("建议持续跟踪", "针对性加强", "能力建设")
    ):
        errors.append("建议句未采用固定范式")

    source_entries, invalid_source_entries, source_raw = _parse_brief_source_entries(source)
    if parsed["unexpected_lines"]:
        errors.append("信息来源须全部写在同一行")
    if source and (invalid_source_entries or not source_entries):
        errors.append("信息来源须逐条完整填写")
    if ";" in source_raw:
        errors.append("多个信息来源须使用中文分号")

    attributions = _brief_body_attributions(body)
    secondary_sources = set()
    for pattern in (
        _BRIEF_SECONDARY_SOURCE_RE,
        _BRIEF_CONTEXT_SOURCE_RE,
        _BRIEF_RECIPIENT_SOURCE_RE,
    ):
        secondary_sources.update(
            match.group("name").strip() for match in pattern.finditer(body or "")
        )
    if body and not re.match(r"^据[^，,。；;]{1,80}?报道[，,]", body):
        errors.append("正文必须以据XX报道开头")
    source_names = {entry["name"].strip() for entry in source_entries}
    if attributions and source_entries:
        if attributions[0]["name"].strip() != source_entries[0]["name"].strip():
            errors.append("帽段据XX报道须与信息来源第一条一致")
    for attribution in attributions:
        if attribution["date"] or attribution["has_relative_date"]:
            errors.append("据XX报道归属中不得添加发文日期或相对时间")
        if source_names and attribution["name"].strip() not in source_names:
            errors.append("正文来源与信息来源行不一致")
    if source_names and any(name not in source_names for name in secondary_sources):
        errors.append("正文二次来源须全部列入信息来源行")
    if body and not attributions:
        errors.append("正文缺少据XX报道来源归属")

    if evidence is not None:
        material_text = str(evidence.get("material_text") or "")
        publication_date = _brief_parse_date_value(evidence.get("publication_date"))
        publication_verified = bool(evidence.get("publication_date_verified"))
        publication_year = publication_date.year if publication_verified and publication_date else None
        if event_ymd and not _brief_event_date_supported(
            material_text, *event_ymd, publication_year=publication_year,
        ):
            errors.append("事件时间未在原始素材中获得对应日期证据")
        if source_entries:
            first_entry = source_entries[0]
            first_name_key = _brief_compact(first_entry["name"]).casefold()
            expected_aliases = {
                _brief_compact(alias).casefold()
                for alias in evidence.get("source_aliases") or []
                if _brief_compact(alias)
            }
            primary_matches_expected = bool(expected_aliases and first_name_key in expected_aliases)
            if not (
                primary_matches_expected
                or _brief_name_supported_in_material(first_entry["name"], material_text)
            ):
                errors.append("信息来源第一条名称未在输入来源或原始素材中获得支持")
            if primary_matches_expected:
                if not publication_verified:
                    errors.append("输入来源缺少可核实的发文日期，不能生成信息来源行")
                elif (
                    first_entry["month"] != publication_date.month
                    or first_entry["day"] != publication_date.day
                ):
                    errors.append("信息来源第一条发文日期与输入来源发布日期不一致")
            elif not _brief_month_day_supported(
                material_text, first_entry["month"], first_entry["day"]
            ):
                errors.append("第一信源发文日期未在原始素材中获得支持")
            unsupported_sources = [
                entry["name"]
                for entry in source_entries[1:]
                if not (
                    _brief_name_supported_in_material(entry["name"], material_text)
                    and _brief_month_day_supported(material_text, entry["month"], entry["day"])
                )
            ]
            if unsupported_sources:
                errors.append("以下引用来源缺少名称和发文日期证据：" + "、".join(unsupported_sources[:5]))
    if reporter and not re.fullmatch(r"报送人：\s*电话：\s*", reporter):
        errors.append("报送人和电话必须留空待填")

    return {"valid": not errors, "errors": list(dict.fromkeys(errors)), "parsed": parsed}

def _brief_validation_error_text(validation: dict) -> str:
    errors = validation.get("errors") or ["格式不完整"]
    shown = errors[:3]
    suffix = f"；另有{len(errors) - len(shown)}项" if len(errors) > len(shown) else ""
    return f"要讯校验未通过：{'；'.join(shown)}{suffix}"

def _validate_brief_before_send(chat_id: str, brief: str, evidence: dict | None = None) -> bool:
    validation = _validate_brief_text(brief, evidence=evidence)
    if validation["valid"]:
        return True
    logger.warning("要讯发送被机械校验阻断: %s", " | ".join(validation["errors"]))
    send_text(chat_id, _brief_validation_error_text(validation))
    return False

def _build_brief_card(brief_text: str, source_info: dict,
                      evidence: dict | None = None) -> dict:
    validation = _validate_brief_text(brief_text, evidence=evidence)
    if not validation["valid"]:
        raise ValueError(_brief_validation_error_text(validation))
    title = (source_info.get("title") or "防务要讯")[:60]
    source = source_info.get("source", "")
    url = source_info.get("url", "")
    now = time.localtime()
    ts = f"{now.tm_mon:02d}月{now.tm_mday:02d}日 {now.tm_hour:02d}:{now.tm_min:02d}"
    preview = brief_text[:1500] + ("…" if len(brief_text) > 1500 else "")
    elements = [
        {"tag": "div", "text": {"tag": "lark_md", "content": f"**素材标题**　{title}"}},
        {"tag": "div", "text": {"tag": "lark_md", "content": f"**信源**　{source or '用户导入'}　　**时间**　{ts}"}},
        {"tag": "hr"},
        {"tag": "div", "text": {"tag": "lark_md", "content": f"**要讯全文**\n\n{preview}"}},
    ]
    if url:
        elements.append({
            "tag": "action",
            "actions": [{
                "tag": "button",
                "text": {"tag": "plain_text", "content": "查看原文"},
                "type": "default",
                "url": url,
            }]
        })
    elements.append({"tag": "hr"})
    elements.append({
        "tag": "div",
        "text": {"tag": "lark_md", "content": (
            "**交互优化**  直接回复调整要求（如「正文压缩到200字」「标题更警示」）即可继续优化\n"
            "发「导出」生成DOCX  发「查看」重新展示全文  发「完成」结束会话\n"
            "**改DOCX文件名**  发「文件名 2026年4月18日」或「把导出的文件名改成 XXX」"
        )},
    })
    return {
        "config": {"wide_screen_mode": True},
        "header": {"title": {"tag": "plain_text", "content": "防务要讯已生成"}, "template": "blue"},
        "elements": elements,
    }


def _build_error_card(msg: str) -> dict:
    return {
        "config": {"wide_screen_mode": True},
        "header": {"title": {"tag": "plain_text", "content": "生成失败"}, "template": "red"},
        "elements": [
            {"tag": "div", "text": {"tag": "lark_md", "content": msg}},
            {"tag": "div", "text": {"tag": "lark_md", "content": "**提示**：检查链接是否可访问，或AI Key环境变量是否已配置"}},
        ],
    }


# ════════════════════════════════════════════════════════════
# 消息处理
# ════════════════════════════════════════════════════════════

def _extract_url(text: str) -> str | None:
    m = re.search(r'https?://[^\s<>"{}|\\^`\[\]]+', text)
    return m.group(0) if m else None


HELP_TEXT = (
    "防务要讯机器人 v4.0（云端版）\n\n"
    "【生成要讯】\n"
    "  发送文章链接 → 自动抓取并生成要讯\n"
    "  发送文章正文（30字+）→ 直接生成\n"
    "  发送图片（新闻截图）→ AI识别生成\n"
    "  发送PDF文件 → 提取文字生成\n\n"
    "【✏️ 交互优化（生成后可继续对话）】\n"
    "  直接回复任意要求 → 按要求修改要讯\n"
    "    例：「把正文压缩到200字」\n"
    "    例：「标题改成更突出核威胁的」\n"
    "    例：「价值点重点强调对我太空资产的影响」\n"
    "  导出 / export → 生成并发送 DOCX 文件\n"
    "  查看 / show → 重新展示当前要讯全文\n"
    "  重新生成 → 基于原素材重新生成（丢弃修改）\n"
    "  文件名 XXX → 设置下次导出的DOCX文件名（也可说「把导出文件名改成XXX」）\n"
    "  完成 / 取消 → 结束优化会话\n\n"
    "【群聊使用】\n"
    "  @机器人 + 链接/文字 即可触发\n\n"
    "【情报推送】\n"
    "  订阅 / 取消订阅 / 扫描 / 状态\n"
    "  brief / headlines → 切换推送模式\n\n"
    "相同链接 1 小时内复用缓存（不重复消耗token）\n"
    "优化会话 30 分钟无操作后自动清除"
)


# ════════════════════════════════════════════════════════════
# DOCX 要讯排版（严格按模板1格式）
# ════════════════════════════════════════════════════════════

def _set_font(run, font_name: str, size_pt: float = 16, bold: bool = True):
    run.font.name = font_name
    run.font.size = Pt(size_pt)
    run.font.bold = bold
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.get_or_add_rFonts()
    rFonts.set(qn('w:eastAsia'), font_name)


def _add_para(doc, align=None, left_indent=None, first_indent=None):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    pf.line_spacing = Emu(363220)
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    if align is not None:
        p.alignment = align
    if left_indent is not None:
        pf.left_indent = Emu(left_indent)
    if first_indent is not None:
        pf.first_line_indent = Emu(first_indent)
    return p


def _parse_brief_sections(text: str) -> dict:
    """解析已通过门禁的六部分要讯；缺项时失败关闭。"""
    lines = text.strip().split('\n')
    sec = {'event_time': '', 'value_point': '', 'title': '',
           'body': '', 'source': '', 'reporter': ''}
    remaining = []
    for line in lines:
        s = line.strip()
        if not s:
            continue
        if s.startswith('事件时间'):
            sec['event_time'] = s
        elif '价' in s and '值' in s and '点' in s and '：' in s:
            sec['value_point'] = s
        elif s.startswith('（信息来源') or s.startswith('(信息来源'):
            sec['source'] = s
        elif s.startswith('报送人'):
            sec['reporter'] = s
        else:
            remaining.append(s)

    non_empty = [l for l in remaining if l]
    if non_empty:
        longest_idx = max(range(len(non_empty)), key=lambda i: len(non_empty[i]))
        sec['body'] = non_empty[longest_idx]
        title_parts = [non_empty[i] for i in range(len(non_empty)) if i != longest_idx]
        sec['title'] = ''.join(title_parts)

    missing = [name for name, value in sec.items() if not value]
    if missing:
        raise ValueError("要讯结构不完整，不能生成DOCX：" + "、".join(missing))
    return sec


def _generate_brief_docx(brief_text: str, evidence: dict | None = None) -> bytes:
    validation = _validate_brief_text(brief_text, evidence=evidence)
    if not validation["valid"]:
        raise ValueError(_brief_validation_error_text(validation))
    sec = _parse_brief_sections(brief_text)
    doc = DocxDocument()
    section = doc.sections[0]
    section.page_width = Emu(7560310)
    section.page_height = Emu(10692130)
    section.top_margin = Emu(914400)
    section.bottom_margin = Emu(914400)
    section.left_margin = Emu(1143000)
    section.right_margin = Emu(1143000)

    # 事件时间
    p = _add_para(doc)
    et = sec['event_time']
    if '：' in et:
        label, date_val = et.split('：', 1)
        _set_font(p.add_run(label + '：'), '黑体')
        _set_font(p.add_run(date_val), '楷体_GB2312')
    else:
        _set_font(p.add_run(et), '黑体')

    # 价值点
    p = _add_para(doc, left_indent=1019810, first_indent=-1019810)
    vp = sec['value_point']
    if '：' in vp:
        label, content = vp.split('：', 1)
        _set_font(p.add_run(label + '：'), '黑体')
        _set_font(p.add_run(content), '楷体_GB2312')
    else:
        _set_font(p.add_run(vp), '黑体')

    # 空行
    _add_para(doc, align=WD_ALIGN_PARAGRAPH.CENTER, left_indent=1606550, first_indent=-1612900)

    # 标题
    p = _add_para(doc, align=WD_ALIGN_PARAGRAPH.CENTER, left_indent=1606550, first_indent=-1612900)
    _set_font(p.add_run(sec['title']), '方正小标宋简体', 22)

    # 空行
    p = _add_para(doc, align=WD_ALIGN_PARAGRAPH.CENTER, left_indent=1606550, first_indent=-1612900)
    _set_font(p.add_run('      '), '方正小标宋简体', 22)

    # 正文
    p = _add_para(doc, first_indent=408305)
    _set_font(p.add_run(sec['body']), '仿宋_GB2312')

    # 信息来源
    if sec['source']:
        p = _add_para(doc, first_indent=408305)
        _set_font(p.add_run(sec['source']), '楷体_GB2312')

    # 报送人
    p = _add_para(doc, align=WD_ALIGN_PARAGRAPH.CENTER)
    _set_font(p.add_run(sec['reporter']), '楷体_GB2312')

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


# 飞书 internal server error 触发字符（实测）：引号/控制字符/中文标点/部分符号
def upload_file(file_name: str, file_data: bytes, file_type: str = "stream") -> str:
    """上传文件到飞书，返回 file_key。
    file_type 支持：opus/mp4/pdf/doc/xls/ppt/stream。
    关键点：
    1) .docx/.doc 必须用 file_type="doc"，用 stream 会 internal server error
    2) multipart filename 必须 ASCII 占位
    3) 真实文件名（中文）通过 form 字段 file_name 传递，需先清洗
    4) 失败一次后用 ASCII fallback 文件名重试
    """
    token = _get_tenant_token()
    cleaned_name = _sanitize_feishu_filename(file_name)
    lower = cleaned_name.lower()
    ext_map = {
        ".docx": ("doc", "application/vnd.openxmlformats-officedocument.wordprocessingml.document"),
        ".doc":  ("doc", "application/msword"),
        ".xlsx": ("xls", "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"),
        ".xls":  ("xls", "application/vnd.ms-excel"),
        ".pptx": ("ppt", "application/vnd.openxmlformats-officedocument.presentationml.presentation"),
        ".ppt":  ("ppt", "application/vnd.ms-powerpoint"),
        ".pdf":  ("pdf", "application/pdf"),
        ".txt":  ("stream", "text/plain"),
        ".md":   ("stream", "text/markdown"),
    }
    mime = "application/octet-stream"
    matched_ext = ".bin"
    for ext, (ft, mt) in ext_map.items():
        if lower.endswith(ext):
            file_type = ft
            mime = mt
            matched_ext = ext
            break
    ascii_placeholder = f"upload{matched_ext}"

    def _do_upload(send_name: str) -> dict:
        resp = requests.post(
            f"{FEISHU_API}/im/v1/files",
            headers={"Authorization": f"Bearer {token}"},
            data={"file_type": file_type, "file_name": send_name},
            files={"file": (ascii_placeholder, file_data, mime)},
            timeout=30,
        )
        return resp.json()

    # 第一次尝试：清洗后的中文文件名
    data = _do_upload(cleaned_name)
    last_err = data.get("msg", "")
    if data.get("code") == 0:
        return data["data"]["file_key"]

    # 第二次尝试：ASCII fallback（保留原文件名中的英文关键词）
    logger.warning("飞书上传第一次失败 (%s)，尝试 ASCII fallback", last_err)
    fallback_name = _ascii_fallback_name(matched_ext, original=file_name)
    data = _do_upload(fallback_name)
    if data.get("code") == 0:
        logger.info("飞书上传 ASCII fallback 成功: %s", fallback_name)
        return data["data"]["file_key"]

    raise RuntimeError(
        f"飞书文件上传失败: {data.get('msg', last_err)} "
        f"(file_type={file_type}, size={len(file_data)}B, "
        f"tried_name={cleaned_name!r}, fallback={fallback_name!r})"
    )


def send_file(chat_id: str, file_key: str):
    _feishu_post("/im/v1/messages?receive_id_type=chat_id", {
        "receive_id": chat_id,
        "msg_type": "file",
        "content": json.dumps({"file_key": file_key}),
    })


def _process_async(chat_id: str, text: str):
    """在子线程中处理消息，生成并回复要讯（带缓存+限速+历史）"""
    try:
        if not AI_CONFIG.get("api_key"):
            send_text(chat_id, "AI API Key 未配置，请在云平台环境变量中设置 AI_API_KEY。")
            return

        url = _extract_url(text)

        if url:
            # 缓存命中 → 直接返回（省钱）
            cached_result, cached_info, cached_evidence = _cache_get(url)
            if cached_result and cached_evidence:
                if not _validate_brief_before_send(chat_id, cached_result, cached_evidence):
                    return
                send_text(chat_id, "（命中缓存，秒回）")
                send_card(chat_id, _build_brief_card(cached_result, cached_info, cached_evidence))
                _session_set(chat_id, {
                    "source_info": cached_info,
                    "evidence": cached_evidence,
                    "current_draft": cached_result,
                    "messages": [
                        {"role": "system", "content": SYSTEM_PROMPT_BRIEF_WRITE},
                        {"role": "user", "content": _build_user_prompt(
                            title=cached_info.get("title", ""),
                            body=cached_evidence.get("material_text", ""),
                            source=cached_info.get("source", ""),
                            url=cached_info.get("url", ""),
                            pub_date=cached_info.get("pub_date", ""),
                        )},
                        {"role": "assistant", "content": cached_result},
                    ],
                })
                return

            # SSRF 防护
            safe, reason = _is_ssrf_safe(url)
            if not safe:
                send_text(chat_id, f"URL不安全，拒绝访问：{reason}")
                return
            send_text(chat_id, "正在抓取链接，请稍候（约10-30秒）...")
            extracted = _extract_url_content(url)
            body = extracted.get("body", "")
            if not body or len(body) < 50:
                send_card(chat_id, _build_error_card(f"无法提取页面正文（字符数：{len(body)}）\n链接：{url}"))
                return
            if not _brief_parse_date_value(extracted.get("pub_date")):
                send_text(chat_id, "页面未提取到可核实的发文日期，无法生成完整信息来源行。")
                return
            source_info = {
                "title": extracted.get("title", url[:60]),
                "source": extracted.get("source", ""),
                "url": url,
                "pub_date": extracted.get("pub_date", ""),
            }
            evidence = _brief_evidence(
                material_text="\n".join(filter(None, [extracted.get("title"), body])),
                source_name=extracted.get("source", ""),
                source_title=extracted.get("title", ""),
                publication_date=extracted.get("pub_date", ""),
                publication_date_verified=True,
                url=url,
            )
            prompt = _build_user_prompt(
                title=extracted.get("title", ""),
                body=body,
                source=extracted.get("source", ""),
                url=url,
                pub_date=extracted.get("pub_date", ""),
            )

        elif len(text) >= 30:
            url = None
            send_text(chat_id, "正在根据文本内容生成要讯，请稍候...")
            source_info = {"title": text[:40], "source": "用户导入", "url": ""}
            evidence = _brief_evidence(material_text=text)
            prompt = _build_user_prompt(title=text[:40], body=text)

        else:
            send_text(chat_id,
                      "请发送以下任一内容：\n"
                      "1. 文章链接（自动抓取正文）\n"
                      "2. 文章正文（30字以上）\n\n"
                      "发送「帮助」查看使用说明")
            return

        # 限速：等待信号量（最多 2 个并发 AI 调用）
        with _ai_semaphore:
            messages = [
                {"role": "system", "content": SYSTEM_PROMPT_BRIEF_WRITE},
                {"role": "user", "content": prompt},
            ]
            result = _call_ai(messages)

        if not _validate_brief_before_send(chat_id, result, evidence):
            return
        send_card(chat_id, _build_brief_card(result, source_info, evidence))

        # 保存交互会话（支持后续多轮优化）
        _session_set(chat_id, {
            "source_info": source_info,
            "evidence": evidence,
            "current_draft": result,
            "original_prompt": prompt,
            "messages": [
                {"role": "system", "content": SYSTEM_PROMPT_BRIEF_WRITE},
                {"role": "user", "content": prompt},
                {"role": "assistant", "content": result},
            ],
        })

        # 生成 DOCX 并发送
        try:
            docx_bytes = _generate_brief_docx(result, evidence)
            title_short = (source_info.get("title") or "防务要讯")[:30]
            ts = time.strftime("%Y%m%d_%H%M")
            file_name = f"要讯_{title_short}_{ts}.docx"
            file_key = upload_file(file_name, docx_bytes)
            send_file(chat_id, file_key)
        except Exception as docx_err:
            logger.error("DOCX生成/发送失败: %s", docx_err, exc_info=True)
            send_text(chat_id, f"要讯文本已生成，但DOCX发送失败：{str(docx_err)[:100]}")

        # 写入缓存 + 历史
        if url:
            _cache_set(url, result, source_info, evidence)
        _history_add(source_info.get("title", ""), source_info.get("source", ""),
                     source_info.get("url", ""), result)

    except Exception as e:
        logger.error("处理消息异常: %s", e, exc_info=True)
        send_card(chat_id, _build_error_card(f"处理异常：{str(e)[:120]}"))


def _process_image_async(chat_id: str, image_key: str):
    """处理图片消息：下载 → AI识别 → 生成要讯"""
    try:
        send_text(chat_id, "正在识别图片内容，请稍候...")
        img_bytes = _download_feishu_image(image_key)
        img_b64 = base64.b64encode(img_bytes).decode()
        # 猜测图片类型
        media_type = "image/jpeg"
        if img_bytes[:8] == b'\x89PNG\r\n\x1a\n':
            media_type = "image/png"

        ocr_prompt = (
            "逐字转录这张新闻截图中可见的全部文字，包括媒体名称、文章标题、发文日期和正文日期。"
            "看不清的字符写[无法辨认]，不得补写、推断或概括。只输出转录文本。"
        )
        with _ai_semaphore:
            ocr_text = _call_ai_with_image(
                img_b64,
                media_type,
                ocr_prompt,
                system_prompt="你是严格的OCR转录器，只记录图片中可见文字，不推断任何缺失内容。",
            )
        if len((ocr_text or "").strip()) < 30:
            send_text(chat_id, "图片可核验文字不足，无法生成要讯。")
            return
        prompt = _build_user_prompt(title="图片新闻截图", body=ocr_text)
        with _ai_semaphore:
            result = _call_ai([
                {"role": "system", "content": SYSTEM_PROMPT_BRIEF_WRITE},
                {"role": "user", "content": prompt},
            ])

        source_info = {"title": "图片识别生成", "source": "图片导入", "url": ""}
        evidence = _brief_evidence(material_text=ocr_text)
        if not _validate_brief_before_send(chat_id, result, evidence):
            return
        send_card(chat_id, _build_brief_card(result, source_info, evidence))
        _session_set(chat_id, {
            "source_info": source_info,
            "evidence": evidence,
            "current_draft": result,
            "original_prompt": prompt,
            "messages": [
                {"role": "system", "content": SYSTEM_PROMPT_BRIEF_WRITE},
                {"role": "user", "content": prompt},
                {"role": "assistant", "content": result},
            ],
        })
        _history_add("图片识别", "图片导入", "", result)
    except Exception as e:
        logger.error("图片处理异常: %s", e, exc_info=True)
        send_card(chat_id, _build_error_card(f"图片处理失败：{str(e)[:120]}"))


def _process_file_async(chat_id: str, message_id: str, file_key: str, file_name: str):
    """处理文件消息：下载文件 → 提取文字 → 生成要讯。
    支持：PDF、DOCX、DOC、TXT、MD、Markdown"""
    try:
        lower = file_name.lower()
        _PDF_EXTS = (".pdf",)
        _DOCX_EXTS = (".docx", ".doc")
        _TEXT_EXTS = (".txt", ".md", ".markdown")

        if lower.endswith(_PDF_EXTS):
            if not _HAS_PDFPLUMBER:
                send_text(chat_id, "PDF 解析库未安装，请联系管理员安装 pdfplumber")
                return
            send_text(chat_id, f"正在解析 PDF：{file_name}…")
            file_bytes = _download_feishu_file(message_id, file_key)
            text = _extract_pdf_text(file_bytes)
            source_type = "PDF导入"
        elif lower.endswith(_DOCX_EXTS):
            send_text(chat_id, f"正在解析文档：{file_name}…")
            file_bytes = _download_feishu_file(message_id, file_key)
            text = _extract_docx_text(file_bytes)
            source_type = "DOCX导入"
        elif lower.endswith(_TEXT_EXTS):
            send_text(chat_id, f"正在解析文本：{file_name}…")
            file_bytes = _download_feishu_file(message_id, file_key)
            text = _extract_text_file(file_bytes)
            source_type = "文本导入"
        else:
            ext = lower.rsplit(".", 1)[-1] if "." in lower else "未知"
            send_text(
                chat_id,
                f"不支持的文件格式（.{ext}），收到：{file_name}\n"
                "支持格式：PDF、DOCX、DOC、TXT、MD"
            )
            return

        if not text or len(text) < 50:
            send_card(chat_id, _build_error_card(
                f"文件提取文本过少（{len(text)}字），请确认文件内容不为空"
            ))
            return

        source_info = {"title": file_name, "source": source_type, "url": ""}
        evidence = _brief_evidence(material_text=text)
        prompt = _build_user_prompt(title=file_name, body=text)

        with _ai_semaphore:
            messages = [
                {"role": "system", "content": SYSTEM_PROMPT_BRIEF_WRITE},
                {"role": "user", "content": prompt},
            ]
            result = _call_ai(messages)

        if not _validate_brief_before_send(chat_id, result, evidence):
            return
        send_card(chat_id, _build_brief_card(result, source_info, evidence))
        _session_set(chat_id, {
            "source_info": source_info,
            "evidence": evidence,
            "current_draft": result,
            "original_prompt": prompt,
            "messages": [
                {"role": "system", "content": SYSTEM_PROMPT_BRIEF_WRITE},
                {"role": "user", "content": prompt},
                {"role": "assistant", "content": result},
            ],
        })
        _history_add(file_name, source_type, "", result)
    except Exception as e:
        logger.error("文件处理异常: %s", e, exc_info=True)
        send_card(chat_id, _build_error_card(f"文件处理失败：{str(e)[:120]}"))


_RENAME_INTENT_RE = re.compile(
    r'(?:把\s*)?(?:导出\s*(?:的)?\s*)?(?:docx|DOCX|word|文件|要讯)?\s*'
    r'(?:文件名|名字|标题|title|filename|文件标题)\s*'
    r'(?:改(?:成|为)?|设(?:为|成)?|命名(?:为)?|叫|为)\s*'
    r'[「『"（(\[]*\s*([^，。！？!?\n」』"）)\]]+?)\s*[」』"）)\]]*\s*$',
    re.IGNORECASE
)

def _try_extract_rename(text: str) -> str | None:
    """从自然语言中识别 "把文件名改成 XXX" 意图。"""
    t = text.strip()
    for prefix in ("文件名", "重命名", "导出名", "导出文件名"):
        if t.startswith(prefix):
            rest = t[len(prefix):].lstrip(" :：=").strip(" 「」『』\"“”()（）")
            if rest:
                return rest
    m = _RENAME_INTENT_RE.search(t)
    if m:
        return m.group(1).strip(" 「」『』\"“”()（）")
    return None


def _refine_async(chat_id: str, instruction: str):
    """多轮交互：按用户指令调整当前要讯草稿"""
    session = _session_get(chat_id)
    if not session:
        send_text(chat_id, "当前没有进行中的要讯会话，请先发送文章或链接生成要讯。")
        return
    evidence = session.get("evidence")
    if not evidence:
        send_text(chat_id, "当前会话缺少原始素材证据，请重新发送文章、文件或图片后再调整。")
        return
    try:
        send_text(chat_id, "正在按您的要求调整要讯，请稍候...")
        messages = session["messages"] + [{"role": "user", "content": instruction}]
        with _ai_semaphore:
            result = _call_ai(messages)
        if not _validate_brief_before_send(chat_id, result, evidence):
            return
        session["current_draft"] = result
        session["messages"] = messages + [{"role": "assistant", "content": result}]
        _session_set(chat_id, session)
        send_card(chat_id, _build_brief_card(result, session["source_info"], evidence))
    except Exception as e:
        logger.error("细化要讯异常: %s", e, exc_info=True)
        send_card(chat_id, _build_error_card(f"调整失败：{str(e)[:120]}"))


def _export_docx_async(chat_id: str):
    """将当前会话草稿导出为 DOCX 并发送"""
    session = _session_get(chat_id)
    if not session:
        send_text(chat_id, "当前没有进行中的要讯会话，请先发送文章或链接生成要讯。")
        return
    evidence = session.get("evidence")
    if not evidence:
        send_text(chat_id, "当前会话缺少原始素材证据，已阻止导出；请重新生成要讯。")
        return
    try:
        draft = session["current_draft"]
        if not _validate_brief_before_send(chat_id, draft, evidence):
            return
        docx_bytes = _generate_brief_docx(draft, evidence)
        custom = session.get("custom_filename")
        if custom:
            file_name = custom if custom.lower().endswith(".docx") else f"{custom}.docx"
        else:
            title_short = (session["source_info"].get("title") or "防务要讯")[:30]
            ts = time.strftime("%Y%m%d_%H%M")
            file_name = f"要讯_{title_short}_{ts}.docx"
        file_key = upload_file(file_name, docx_bytes)
        send_file(chat_id, file_key)
    except Exception as e:
        logger.error("DOCX导出失败: %s", e, exc_info=True)
        send_text(chat_id, f"DOCX导出失败：{str(e)[:100]}")


def _regenerate_async(chat_id: str):
    """基于原始素材重新生成（丢弃所有修改历史）"""
    session = _session_get(chat_id)
    if not session:
        send_text(chat_id, "当前没有进行中的要讯会话，请先发送文章或链接生成要讯。")
        return
    evidence = session.get("evidence")
    if not evidence:
        send_text(chat_id, "当前会话缺少原始素材证据，请重新发送原始材料。")
        return
    original_prompt = session.get("original_prompt")
    if not original_prompt:
        send_text(chat_id, "图片会话暂不支持重新生成，请重新发送图片。")
        return
    source_info = session["source_info"]
    try:
        send_text(chat_id, "正在重新生成要讯（丢弃所有修改历史）...")
        with _ai_semaphore:
            messages = [
                {"role": "system", "content": SYSTEM_PROMPT_BRIEF_WRITE},
                {"role": "user", "content": original_prompt},
            ]
            result = _call_ai(messages)
        if not _validate_brief_before_send(chat_id, result, evidence):
            return
        _session_clear(chat_id)
        send_card(chat_id, _build_brief_card(result, source_info, evidence))
        _session_set(chat_id, {
            "source_info": source_info,
            "evidence": evidence,
            "current_draft": result,
            "original_prompt": original_prompt,
            "messages": [
                {"role": "system", "content": SYSTEM_PROMPT_BRIEF_WRITE},
                {"role": "user", "content": original_prompt},
                {"role": "assistant", "content": result},
            ],
        })
    except Exception as e:
        logger.error("重新生成异常: %s", e, exc_info=True)
        send_card(chat_id, _build_error_card(f"重新生成失败：{str(e)[:120]}"))


# ════════════════════════════════════════════════════════════
# RSS 自动推送引擎
# ════════════════════════════════════════════════════════════

def _score_article(title: str, summary: str) -> int:
    """对文章标题+摘要进行防务关键词评分"""
    text = f"{title} {summary}"
    score = 0
    for weight, pattern in _DEFENSE_PATTERNS:
        if pattern.search(text):
            score += weight
    return score


def _parse_pub_date(entry) -> datetime:
    """解析 feedparser entry 的发布时间"""
    for attr in ("published_parsed", "updated_parsed"):
        parsed = getattr(entry, attr, None)
        if parsed:
            try:
                return datetime(*parsed[:6], tzinfo=timezone.utc)
            except Exception:
                pass
    return datetime.now(timezone.utc)


def _fetch_one_feed(feed: dict) -> list:
    """拉取单个 RSS 源，返回文章列表"""
    articles = []
    cutoff = datetime.now(timezone.utc) - timedelta(hours=24)
    try:
        headers = {**_BROWSER_HEADERS}
        r = _safe_get_once(feed["url"], headers=headers, timeout=12)
        r.raise_for_status()
        parsed = feedparser.parse(r.content)
        for entry in parsed.entries[:15]:
            title = getattr(entry, "title", "").strip()
            link = getattr(entry, "link", "").strip()
            summary = re.sub(r"<[^>]+>", "", getattr(entry, "summary", ""))[:300]
            if not title or not link:
                continue
            pub = _parse_pub_date(entry)
            if pub < cutoff:
                continue
            score = _score_article(title, summary)
            if score < 2:
                continue
            articles.append({
                "title": title, "link": link, "summary": summary,
                "source": feed["name_cn"], "source_en": feed["name"],
                "focus": feed["focus"], "score": score, "pub": pub,
            })
    except Exception as e:
        logger.debug("RSS fetch failed %s: %s", feed["name"], e)
    return articles


def _fetch_all_feeds() -> list:
    """并发拉取所有 RSS 源，合并并按分数排序"""
    all_articles = []
    with ThreadPoolExecutor(max_workers=8) as executor:
        futures = {executor.submit(_fetch_one_feed, f): f for f in RSS_FEEDS_CLOUD}
        for future in as_completed(futures, timeout=30):
            try:
                all_articles.extend(future.result())
            except Exception:
                pass
    all_articles.sort(key=lambda a: a["score"], reverse=True)
    return all_articles


def _build_headline_card(articles: list) -> dict:
    """构建 RSS 摘要推送卡片（headlines 模式）"""
    now = time.localtime()
    ts = f"{now.tm_mon:02d}月{now.tm_mday:02d}日 {now.tm_hour:02d}:{now.tm_min:02d}"
    elements = [
        {"tag": "div", "text": {"tag": "lark_md",
            "content": f"**自动扫描时间**　{ts}　｜　共发现 **{len(articles)}** 篇高价值文章"}},
        {"tag": "hr"},
    ]
    for i, art in enumerate(articles, 1):
        score_stars = "★" * min(art["score"], 5) + "☆" * max(0, 5 - art["score"])
        elements.append({"tag": "div", "text": {"tag": "lark_md",
            "content": (
                f"**{i}. [{art['title'][:50]}]({art['link']})**\n"
                f"　　{art['source']}　|　{score_stars}　|　{art['summary'][:80]}..."
            )}})
    elements.append({"tag": "hr"})
    elements.append({"tag": "div", "text": {"tag": "lark_md",
        "content": "转发任意链接给我即可生成完整要讯"}})
    return {
        "config": {"wide_screen_mode": True},
        "header": {"title": {"tag": "plain_text", "content": f"防务情报扫描 | {ts}"}, "template": "indigo"},
        "elements": elements,
    }


def _rss_push_job():
    """定时任务：拉取 RSS → 筛选 → 推送到飞书"""
    chat_id = PUSH_CONFIG["chat_id"]
    if not chat_id:
        return
    try:
        logger.info("RSS 自动推送开始...")
        articles = _fetch_all_feeds()

        # 去重：过滤已推送的
        new_articles = []
        with _pushed_lock:
            for art in articles:
                if art["link"] not in _pushed_urls:
                    new_articles.append(art)
        if not new_articles:
            logger.info("RSS 推送：无新文章")
            return

        # 取 top N
        top = new_articles[:PUSH_CONFIG["max_articles"]]

        if PUSH_CONFIG["mode"] == "brief":
            # brief 模式：逐篇生成 AI 要讯
            for art in top:
                try:
                    safe, _ = _is_ssrf_safe(art["link"])
                    if not safe:
                        continue
                    extracted = _extract_url_content(art["link"])
                    body = extracted.get("body", "")
                    if len(body) < 50 or not _brief_parse_date_value(extracted.get("pub_date")):
                        continue
                    evidence = _brief_evidence(
                        material_text="\n".join(filter(None, [extracted.get("title"), body])),
                        source_name=extracted.get("source", art["source"]),
                        source_title=extracted.get("title", art["title"]),
                        publication_date=extracted.get("pub_date", ""),
                        publication_date_verified=True,
                        url=art["link"],
                    )
                    prompt = _build_user_prompt(
                        title=extracted.get("title", art["title"]),
                        body=body,
                        source=extracted.get("source", art["source"]),
                        url=art["link"],
                        pub_date=extracted.get("pub_date", ""),
                    )
                    messages = [
                        {"role": "system", "content": SYSTEM_PROMPT_BRIEF_WRITE},
                        {"role": "user", "content": prompt},
                    ]
                    result = _call_ai(messages)
                    source_info = {
                        "title": art["title"], "source": art["source"], "url": art["link"],
                        "pub_date": extracted.get("pub_date", ""),
                    }
                    if not _validate_brief_before_send(chat_id, result, evidence):
                        continue
                    send_card(chat_id, _build_brief_card(result, source_info, evidence))
                    with _pushed_lock:
                        _pushed_urls.add(art["link"])
                    time.sleep(3)
                except Exception as e:
                    logger.error("Brief push error for %s: %s", art["title"][:30], e)
        else:
            # headlines 模式：一张卡片汇总
            send_card(chat_id, _build_headline_card(top))
            with _pushed_lock:
                for art in top:
                    _pushed_urls.add(art["link"])

        # 清理去重集合（保留最近 2000 条）
        with _pushed_lock:
            if len(_pushed_urls) > 2000:
                excess = list(_pushed_urls)[:500]
                for url in excess:
                    _pushed_urls.discard(url)

        logger.info("RSS 推送完成：%d 篇", len(top))
    except Exception as e:
        logger.error("RSS 推送异常: %s", e, exc_info=True)


_rss_scheduler_running = False

def _start_rss_scheduler():
    """启动 RSS 定时推送后台线程（幂等，不会重复启动）"""
    global _rss_scheduler_running
    if _rss_scheduler_running:
        return
    _rss_scheduler_running = True
    interval = PUSH_CONFIG["interval_min"] * 60

    def _loop():
        time.sleep(10)  # 启动后延迟 10 秒再执行首次
        while True:
            try:
                _rss_push_job()
            except Exception as e:
                logger.error("RSS scheduler error: %s", e)
            time.sleep(interval)

    t = threading.Thread(target=_loop, daemon=True)
    t.start()
    logger.info("RSS 定时推送已启动（间隔 %d 分钟）", PUSH_CONFIG["interval_min"])


# ════════════════════════════════════════════════════════════
# Flask 路由
# ════════════════════════════════════════════════════════════

@app.route("/")
def index():
    """健康检查 & 首页（云平台存活探针）"""
    return jsonify({
        "service": "feishu-defense-brief-bot",
        "status": "running",
        "version": "3.0",
        "feishu_configured": bool(FEISHU_CONFIG["app_id"] and FEISHU_CONFIG["app_secret"]),
        "ai_configured": bool(AI_CONFIG["api_key"]),
        "rss_push": {
            "enabled": bool(PUSH_CONFIG["chat_id"]),
            "mode": PUSH_CONFIG["mode"],
            "interval_min": PUSH_CONFIG["interval_min"],
            "max_articles": PUSH_CONFIG["max_articles"],
            "pushed_count": len(_pushed_urls),
        },
    })


@app.route("/health")
def health():
    return "ok", 200


@app.route("/api/history")
def api_history():
    """历史正文不通过未认证 HTTP 接口公开。"""
    return jsonify({"error": "history endpoint disabled"}), 404


@app.route("/api/cache/stats")
def api_cache_stats():
    """缓存统计"""
    with _cache_lock:
        return jsonify({"cached_urls": len(_result_cache), "max": _CACHE_MAX, "ttl_seconds": _CACHE_TTL})


@app.route("/api/feishu/webhook", methods=["POST"])
def feishu_webhook():
    """飞书事件订阅回调入口"""
    raw_body = request.get_data()
    data = request.get_json(silent=True) or {}
    verify_token = str(FEISHU_CONFIG.get("verify_token") or "").strip()
    if not verify_token:
        logger.error("webhook verification token is not configured")
        return jsonify({"code": 1, "msg": "webhook verification not configured"}), 503

    # URL 验证（首次注册 webhook 时飞书发来）
    if data.get("type") == "url_verification":
        incoming_token = str(data.get("token") or "")
        if not hmac.compare_digest(verify_token, incoming_token):
            logger.warning("URL verification token invalid")
            return jsonify({"code": 1, "msg": "invalid token"}), 403
        logger.info("URL verification OK")
        return jsonify({"challenge": data.get("challenge", "")})

    # 签名校验：仅当飞书启用加密/签名、发来 X-Lark-Signature 时强制（默认 token 模式无此头，
    # 走下面的明文 verify_token 校验，行为不变）。此前 _verify_signature 定义却从未接线，
    # 等于签名防伪形同虚设；此处按"有签名才校验"接线，签名模式下生效、token 模式零影响。
    if request.headers.get("X-Lark-Signature"):
        timestamp = request.headers.get("X-Lark-Request-Timestamp", "")
        nonce = request.headers.get("X-Lark-Request-Nonce", "")
        if not _verify_signature(timestamp, nonce, raw_body):
            logger.warning("X-Lark-Signature 校验失败，请求被拒绝")
            return jsonify({"code": 1, "msg": "invalid signature"}), 403

    # 事件处理（schema 2.0）
    header = data.get("header", {})
    event_type = header.get("event_type", "")

    # Verification Token 校验（飞书在 header.token 中传入）
    incoming_token = str(header.get("token") or "")
    if not hmac.compare_digest(verify_token, incoming_token):
        logger.warning("Verification Token 校验失败，请求被拒绝")
        return jsonify({"code": 1, "msg": "invalid token"}), 403
    event = data.get("event", {})

    if event_type != "im.message.receive_v1":
        return jsonify({"code": 0})

    message = event.get("message", {})
    msg_type = message.get("message_type", "")
    chat_id = message.get("chat_id", "")
    msg_id = message.get("message_id", "")

    if not chat_id:
        return jsonify({"code": 0})

    # 去重
    if not _is_new_message(msg_id):
        logger.info("重复消息 %s，跳过", msg_id)
        return jsonify({"code": 0})

    # ── 图片消息处理 ─────────────────────────────────
    if msg_type == "image":
        try:
            content = json.loads(message.get("content", "{}"))
            image_key = content.get("image_key", "")
            if image_key:
                _worker_pool.submit(_process_image_async, chat_id, image_key)
        except Exception:
            logger.warning("图片消息解析失败: msg_id=%s", msg_id)
        return jsonify({"code": 0})

    # ── 文件消息处理（PDF）─────────────────────────────
    if msg_type == "file":
        try:
            content = json.loads(message.get("content", "{}"))
            file_key = content.get("file_key", "")
            file_name = content.get("file_name", "")
            if file_key:
                _worker_pool.submit(_process_file_async, chat_id, msg_id, file_key, file_name)
        except Exception:
            logger.warning("文件消息解析失败: msg_id=%s", msg_id)
        return jsonify({"code": 0})

    if msg_type != "text":
        send_text(chat_id, "支持：文字/链接/图片/PDF文件。发送「帮助」查看说明。")
        return jsonify({"code": 0})

    try:
        content = json.loads(message.get("content", "{}"))
        text = content.get("text", "").strip()
    except Exception:
        return jsonify({"code": 0})

    # ── @机器人 群聊支持：剥离 @mention 占位符 ──────────
    mentions = content.get("mentions") if isinstance(content, dict) else None
    if mentions:
        # 飞书在 text 中用 @_user_N 占位，需要去掉
        for m in mentions:
            key = m.get("key", "")
            if key:
                text = text.replace(key, "").strip()
    # 群聊中如果没有 @机器人，则不处理（避免干扰正常聊天）
    chat_type = message.get("chat_type", "")
    if chat_type == "group" and not mentions:
        return jsonify({"code": 0})

    if not text:
        return jsonify({"code": 0})

    if text in ["帮助", "help", "?", "？", "/help"]:
        send_text(chat_id, HELP_TEXT)
        return jsonify({"code": 0})

    # ── RSS 推送控制指令 ────────────────────────────
    if text in ["订阅", "subscribe"]:
        PUSH_CONFIG["chat_id"] = chat_id
        _start_rss_scheduler()  # 确保定时线程已启动
        send_text(chat_id,
                  f"已开启自动情报推送到当前会话\n"
                  f"模式：{PUSH_CONFIG['mode']}　间隔：{PUSH_CONFIG['interval_min']}分钟\n"
                  f"每轮最多推送 {PUSH_CONFIG['max_articles']} 篇\n\n"
                  f"发送「扫描」立即执行一次")
        return jsonify({"code": 0})

    if text in ["取消订阅", "unsubscribe"]:
        PUSH_CONFIG["chat_id"] = ""
        send_text(chat_id, "已关闭自动情报推送")
        return jsonify({"code": 0})

    if text in ["扫描", "scan"]:
        send_text(chat_id, "正在扫描 15 个核心 RSS 源...")
        PUSH_CONFIG["chat_id"] = chat_id
        _worker_pool.submit(_rss_push_job)
        return jsonify({"code": 0})

    if text in ["状态", "status"]:
        push_status = "已开启" if PUSH_CONFIG["chat_id"] else "未开启"
        send_text(chat_id,
                  f"机器人状态 v3.0：\n"
                  f"  RSS推送：{push_status}\n"
                  f"  推送模式：{PUSH_CONFIG['mode']}\n"
                  f"  推送间隔：{PUSH_CONFIG['interval_min']} 分钟\n"
                  f"  已推送：{len(_pushed_urls)} 篇\n"
                  f"  缓存命中：{len(_result_cache)} 条\n"
                  f"  历史记录：{len(_history)} 条\n"
                  f"  AI模型：{AI_CONFIG['model']}\n"
                  f"  trafilatura：{'可用' if _HAS_TRAFILATURA else '未安装'}\n"
                  f"  PDF解析：{'可用' if _HAS_PDFPLUMBER else '未安装'}")
        return jsonify({"code": 0})

    if text in ["brief模式", "brief"]:
        PUSH_CONFIG["mode"] = "brief"
        send_text(chat_id, "已切换为 brief 模式（AI生成完整要讯，消耗token）")
        return jsonify({"code": 0})

    if text in ["headlines模式", "headlines"]:
        PUSH_CONFIG["mode"] = "headlines"
        send_text(chat_id, "已切换为 headlines 模式（仅推送标题摘要，免费）")
        return jsonify({"code": 0})

    # ── 文件名重命名意图（无论有无会话都可识别）────────────
    rename_to = _try_extract_rename(text)
    if rename_to:
        sanitized = _sanitize_feishu_filename(
            rename_to if rename_to.lower().endswith(".docx") else f"{rename_to}.docx"
        )
        existing = _session_get(chat_id)
        if existing:
            existing["custom_filename"] = sanitized
            _session_set(chat_id, existing)
            send_text(
                chat_id,
                f"已记录导出文件名为「{sanitized}」。发「导出」即可生成；"
                "若要修改要讯正文，请直接描述修改点（如「正文压缩到200字」）。"
            )
        else:
            _pending_filename_set(chat_id, sanitized)
            send_text(
                chat_id,
                f"已记录文件名「{sanitized}」（目前没有进行中的会话）。\n"
                "请接着发送文章链接或正文，生成的要讯会自动使用该文件名导出。"
            )
        return jsonify({"code": 0})

    # ── 交互会话路由（在其他命令之后、文章生成之前检查）──────
    session = _session_get(chat_id)
    if session:
        # 会话中输入 URL → 视为新文章，清除旧会话后正常处理
        if _extract_url(text):
            _session_clear(chat_id)
            # fall through to _process_async below
        elif text in ["完成", "done", "结束", "exit", "退出", "取消", "cancel"]:
            _session_clear(chat_id)
            send_text(chat_id, "已退出要讯优化会话。如需生成新要讯，请发送文章或链接。")
            return jsonify({"code": 0})
        elif text in ["导出", "export", "发docx", "DOCX", "docx", "发DOCX"]:
            _worker_pool.submit(_export_docx_async, chat_id)
            return jsonify({"code": 0})
        elif text in ["查看", "预览", "show", "preview"]:
            evidence = session.get("evidence")
            if not evidence:
                send_text(chat_id, "当前会话缺少原始素材证据，已阻止预览；请重新生成要讯。")
                return jsonify({"code": 0})
            if not _validate_brief_before_send(chat_id, session["current_draft"], evidence):
                return jsonify({"code": 0})
            send_card(
                chat_id,
                _build_brief_card(session["current_draft"], session["source_info"], evidence),
            )
            return jsonify({"code": 0})
        elif text in ["重新生成", "重写", "重来", "regenerate", "重置"]:
            _worker_pool.submit(_regenerate_async, chat_id)
            return jsonify({"code": 0})
        else:
            # 任意其他文字 → 视为细化指令
            _worker_pool.submit(_refine_async, chat_id, text)
            return jsonify({"code": 0})

    # ── 文章生成（异步线程池，立即返回 200）────────────────
    _worker_pool.submit(_process_async, chat_id, text)
    return jsonify({"code": 0})


# ════════════════════════════════════════════════════════════
# 启动
# ════════════════════════════════════════════════════════════
_scheduler_started = False

def _ensure_scheduler():
    """确保 RSS 定时推送线程只启动一次（gunicorn 多 worker 安全）"""
    global _scheduler_started
    if _scheduler_started:
        return
    _scheduler_started = True
    if PUSH_CONFIG["chat_id"]:
        _start_rss_scheduler()
    else:
        logger.info("RSS 推送未配置 chat_id，等待用户发送「订阅」指令激活")


# gunicorn 启动时自动执行
with app.app_context():
    _ensure_scheduler()


if __name__ == "__main__":
    port = int(os.environ.get("PORT", 5000))
    logger.info("=" * 50)
    logger.info("飞书云端机器人 v3.0 启动")
    logger.info("端口: %d", port)
    logger.info("飞书配置: %s", "已配置" if FEISHU_CONFIG["app_id"] else "未配置")
    logger.info("AI配置: %s (模型: %s)", "已配置" if AI_CONFIG["api_key"] else "未配置", AI_CONFIG["model"])
    logger.info("RSS推送: %s (模式: %s, 间隔: %d分钟)",
                "已配置" if PUSH_CONFIG["chat_id"] else "未配置",
                PUSH_CONFIG["mode"], PUSH_CONFIG["interval_min"])
    logger.info("Webhook: https://<你的域名>/api/feishu/webhook")
    logger.info("=" * 50)
    _ensure_scheduler()
    app.run(host="0.0.0.0", port=port)
