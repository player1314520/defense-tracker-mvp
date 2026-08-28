# -*- coding: utf-8 -*-
"""Fail-closed validation and isolated extraction for untrusted documents.

The public download limit is only the first boundary.  DOCX is a ZIP container,
so compressed and expanded sizes are checked before ``python-docx`` sees it.
PDF parsing is kept in a subprocess that the parent can terminate on timeout.
No parser exception text is returned to callers because third-party exceptions
may echo document bytes, filenames, or other untrusted input.
"""
from __future__ import annotations

import json
import os
import stat
import subprocess
import sys
import tempfile
from dataclasses import dataclass
from io import BytesIO
from pathlib import Path, PurePosixPath
from zipfile import BadZipFile, ZIP_DEFLATED, ZIP_STORED, ZipFile


MIB = 1024 * 1024


_SAFE_MESSAGES = {
    "DOCUMENT_EMPTY": "文档为空",
    "DOCUMENT_TOO_LARGE": "文档超过安全大小限制",
    "DOCX_BAD_MAGIC": "文件不是有效的 DOCX 容器",
    "DOCX_INVALID_ZIP": "DOCX 容器损坏",
    "DOCX_TOO_MANY_ENTRIES": "DOCX 条目数量超过安全限制",
    "DOCX_ENTRY_TOO_LARGE": "DOCX 单个条目超过安全限制",
    "DOCX_EXPANDED_TOO_LARGE": "DOCX 展开后超过安全限制",
    "DOCX_COMPRESSION_RATIO": "DOCX 压缩比超过安全限制",
    "DOCX_ENCRYPTED": "不接受加密 DOCX",
    "DOCX_DANGEROUS_ENTRY": "DOCX 包含危险或嵌套内容",
    "DOCX_UNSUPPORTED_COMPRESSION": "DOCX 使用了不支持的压缩格式",
    "DOCX_MISSING_CORE": "DOCX 缺少必要结构",
    "DOCX_XML_DIRECTIVE": "DOCX 包含不允许的 XML 声明",
    "DOCX_PARSER_UNAVAILABLE": "DOCX 解析器不可用",
    "DOCX_PARSE_FAILED": "DOCX 解析失败",
    "PDF_BAD_MAGIC": "文件不是有效的 PDF",
    "PDF_MISSING_EOF": "PDF 结构不完整",
    "PDF_ENCRYPTED": "不接受加密 PDF",
    "PDF_ACTIVE_CONTENT": "PDF 包含不允许的主动内容",
    "PDF_TOO_COMPLEX": "PDF 结构复杂度超过安全限制",
    "PDF_PARSER_UNAVAILABLE": "受限 PDF 解析器不可用",
    "PDF_PARSE_TIMEOUT": "PDF 解析超时",
    "PDF_PARSE_FAILED": "PDF 解析失败",
    "PDF_OUTPUT_INVALID": "PDF 解析器返回无效结果",
}


class DocumentSafetyError(RuntimeError):
    """Stable, privacy-safe document rejection exposed to application code."""

    def __init__(self, code: str):
        self.code = code
        super().__init__(f"{code}: {_SAFE_MESSAGES.get(code, '文档被安全策略拒绝')}")


@dataclass(frozen=True)
class DocxSafetyReport:
    entries: int
    compressed_bytes: int
    expanded_bytes: int


_NESTED_OR_EXECUTABLE_SUFFIXES = {
    ".7z", ".apk", ".bat", ".cmd", ".com", ".doc", ".docm", ".docx",
    ".exe", ".gz", ".hta", ".iso", ".jar", ".js", ".lnk", ".msi",
    ".pdf", ".ppt", ".pptm", ".pptx", ".ps1", ".rar", ".scr", ".tar",
    ".vbs", ".xls", ".xlsm", ".xlsx", ".xz", ".zip",
}


def _reject(code: str) -> None:
    raise DocumentSafetyError(code)


def _safe_zip_name(raw_name: str) -> str:
    if not raw_name or "\x00" in raw_name:
        _reject("DOCX_DANGEROUS_ENTRY")
    normalized = raw_name.replace("\\", "/")
    path = PurePosixPath(normalized)
    if normalized.startswith("/") or path.is_absolute() or ".." in path.parts:
        _reject("DOCX_DANGEROUS_ENTRY")
    if path.parts and ":" in path.parts[0]:
        _reject("DOCX_DANGEROUS_ENTRY")
    return path.as_posix()


def validate_docx(
    content: bytes,
    *,
    max_input_bytes: int = 18 * MIB,
    max_entries: int = 1024,
    max_entry_bytes: int = 16 * MIB,
    max_expanded_bytes: int = 64 * MIB,
    max_compression_ratio: float = 200.0,
) -> DocxSafetyReport:
    """Validate a DOCX container without trusting filenames or expansion sizes."""
    if not isinstance(content, (bytes, bytearray, memoryview)) or not content:
        _reject("DOCUMENT_EMPTY")
    payload = bytes(content)
    if len(payload) > max_input_bytes:
        _reject("DOCUMENT_TOO_LARGE")
    if not payload.startswith(b"PK\x03\x04"):
        _reject("DOCX_BAD_MAGIC")

    try:
        archive = ZipFile(BytesIO(payload))
    except (BadZipFile, OSError, ValueError):
        _reject("DOCX_INVALID_ZIP")

    with archive:
        try:
            infos = archive.infolist()
        except (BadZipFile, OSError, ValueError):
            _reject("DOCX_INVALID_ZIP")
        if not infos or len(infos) > max_entries:
            _reject("DOCX_TOO_MANY_ENTRIES")

        expanded = 0
        compressed = 0
        seen_names: set[str] = set()
        xml_names: list[str] = []
        for info in infos:
            name = _safe_zip_name(info.filename)
            folded = name.casefold()
            if folded in seen_names:
                _reject("DOCX_DANGEROUS_ENTRY")
            seen_names.add(folded)

            unix_mode = (info.external_attr >> 16) & 0xFFFF
            if stat.S_ISLNK(unix_mode):
                _reject("DOCX_DANGEROUS_ENTRY")
            if info.flag_bits & 0x1:
                _reject("DOCX_ENCRYPTED")
            if info.compress_type not in {ZIP_STORED, ZIP_DEFLATED}:
                _reject("DOCX_UNSUPPORTED_COMPRESSION")
            if info.is_dir():
                continue

            suffix = PurePosixPath(name).suffix.casefold()
            if (
                suffix in _NESTED_OR_EXECUTABLE_SUFFIXES
                or "/embeddings/" in f"/{folded}"
                or folded.endswith("vbaproject.bin")
                or folded.endswith("oleobject.bin")
            ):
                _reject("DOCX_DANGEROUS_ENTRY")

            if info.file_size < 0 or info.file_size > max_entry_bytes:
                _reject("DOCX_ENTRY_TOO_LARGE")
            expanded += info.file_size
            compressed += max(0, info.compress_size)
            if expanded > max_expanded_bytes:
                _reject("DOCX_EXPANDED_TOO_LARGE")
            if info.file_size:
                ratio = info.file_size / max(1, info.compress_size)
                if ratio > max_compression_ratio:
                    _reject("DOCX_COMPRESSION_RATIO")
            if suffix in {".xml", ".rels"} or folded == "[content_types].xml":
                xml_names.append(info.filename)

        required = {"[content_types].xml", "word/document.xml"}
        if not required.issubset(seen_names):
            _reject("DOCX_MISSING_CORE")

        # XML external entities and DTDs are unnecessary in OOXML and create an
        # avoidable parser-expansion surface.  Total reads remain bounded above.
        try:
            for name in xml_names:
                xml = archive.read(name)
                upper = xml.upper()
                if b"<!DOCTYPE" in upper or b"<!ENTITY" in upper:
                    _reject("DOCX_XML_DIRECTIVE")
        except DocumentSafetyError:
            raise
        except (BadZipFile, OSError, RuntimeError, ValueError):
            _reject("DOCX_INVALID_ZIP")

    return DocxSafetyReport(
        entries=len(infos),
        compressed_bytes=compressed,
        expanded_bytes=expanded,
    )


def extract_docx_text_safe(
    content: bytes,
    *,
    max_chars: int = 60_000,
    include_tables: bool = True,
    max_input_bytes: int = 18 * MIB,
) -> str:
    """Validate then extract DOCX text without exposing parser exceptions."""
    validate_docx(content, max_input_bytes=max_input_bytes)
    try:
        from docx import Document
    except ImportError as exc:
        raise DocumentSafetyError("DOCX_PARSER_UNAVAILABLE") from exc
    try:
        doc = Document(BytesIO(bytes(content)))
        parts: list[str] = []
        chars = 0

        def append(value: str) -> bool:
            nonlocal chars
            value = (value or "").strip()
            if not value:
                return True
            remaining = max_chars - chars
            if remaining <= 0:
                return False
            parts.append(value[:remaining])
            chars += min(len(value), remaining) + 1
            return chars < max_chars

        for paragraph in doc.paragraphs:
            if not append(paragraph.text):
                break
        if include_tables and chars < max_chars:
            for table in doc.tables:
                for row in table.rows:
                    line = " | ".join(
                        text for text in ((cell.text or "").strip() for cell in row.cells) if text
                    )
                    if not append(line):
                        break
                if chars >= max_chars:
                    break
        return "\n".join(parts)[:max_chars]
    except DocumentSafetyError:
        raise
    except Exception:
        raise DocumentSafetyError("DOCX_PARSE_FAILED") from None


_PDF_ACTIVE_MARKERS = (
    b"/JAVASCRIPT", b"/JS", b"/LAUNCH", b"/OPENACTION", b"/AA",
    b"/EMBEDDEDFILE", b"/RICHMEDIA", b"/XFA",
)


def validate_pdf(
    content: bytes,
    *,
    max_input_bytes: int = 18 * MIB,
    max_object_markers: int = 100_000,
) -> None:
    """Apply cheap structural limits before handing bytes to a PDF parser."""
    if not isinstance(content, (bytes, bytearray, memoryview)) or not content:
        _reject("DOCUMENT_EMPTY")
    payload = bytes(content)
    if len(payload) > max_input_bytes:
        _reject("DOCUMENT_TOO_LARGE")
    if b"%PDF-" not in payload[:1024]:
        _reject("PDF_BAD_MAGIC")
    if b"%%EOF" not in payload[-8192:]:
        _reject("PDF_MISSING_EOF")
    upper = payload.upper()
    if b"/ENCRYPT" in upper:
        _reject("PDF_ENCRYPTED")
    if any(marker in upper for marker in _PDF_ACTIVE_MARKERS):
        _reject("PDF_ACTIVE_CONTENT")
    if upper.count(b" OBJ") > max_object_markers:
        _reject("PDF_TOO_COMPLEX")


def _pdf_worker_command(
    temp_path: str,
    max_pages: int,
    max_total_pages: int,
    max_chars: int,
    result_path: str | None = None,
) -> list[str]:
    override = (os.environ.get("DEFENSE_TRACKER_PDF_WORKER_PYTHON") or "").strip()
    if override:
        candidate = Path(override).expanduser()
        if not candidate.is_absolute() or not candidate.is_file():
            _reject("PDF_PARSER_UNAVAILABLE")
        prefix = [str(candidate), "-I", str(Path(__file__).resolve()), "--pdf-worker"]
    elif getattr(sys, "frozen", False):
        # PyInstaller onedir re-enters this exact signed executable in a
        # worker-only mode. launcher.py dispatches the private argument before
        # Flask or WebView is imported.
        if not result_path:
            _reject("PDF_PARSER_UNAVAILABLE")
        prefix = [sys.executable, "--defense-tracker-pdf-worker"]
    else:
        prefix = [sys.executable, "-I", str(Path(__file__).resolve()), "--pdf-worker"]
    command = prefix + [
        temp_path,
        str(max_pages),
        str(max_total_pages),
        str(max_chars),
    ]
    if result_path:
        command.append(result_path)
    return command


def extract_pdf_text_isolated(
    content: bytes,
    *,
    max_pages: int = 30,
    max_total_pages: int = 500,
    max_chars: int = 60_000,
    timeout_seconds: int = 12,
    max_input_bytes: int = 18 * MIB,
) -> str:
    """Extract PDF text in a time-bounded subprocess on Windows and POSIX."""
    validate_pdf(content, max_input_bytes=max_input_bytes)
    max_pages = max(1, min(int(max_pages), max_total_pages))
    max_chars = max(1, min(int(max_chars), 250_000))
    timeout_seconds = max(1, min(int(timeout_seconds), 60))

    temp_path = ""
    result_path = ""
    try:
        with tempfile.NamedTemporaryFile(prefix="defense-pdf-", suffix=".pdf", delete=False) as handle:
            temp_path = handle.name
            handle.write(bytes(content))
        try:
            os.chmod(temp_path, 0o600)
        except OSError:
            pass

        uses_frozen_self_worker = bool(
            getattr(sys, "frozen", False)
            and not (os.environ.get("DEFENSE_TRACKER_PDF_WORKER_PYTHON") or "").strip()
        )
        if uses_frozen_self_worker:
            with tempfile.NamedTemporaryFile(
                prefix="defense-pdf-result-", suffix=".json", delete=False
            ) as result_handle:
                result_path = result_handle.name
            try:
                os.chmod(result_path, 0o600)
            except OSError:
                pass
        command = _pdf_worker_command(
            temp_path,
            max_pages,
            max_total_pages,
            max_chars,
            result_path or None,
        )
        creationflags = subprocess.CREATE_NO_WINDOW if os.name == "nt" else 0
        try:
            completed = subprocess.run(
                command,
                stdin=subprocess.DEVNULL,
                stdout=subprocess.PIPE,
                stderr=subprocess.DEVNULL,
                check=False,
                timeout=timeout_seconds,
                encoding="utf-8",
                errors="strict",
                creationflags=creationflags,
            )
        except subprocess.TimeoutExpired:
            _reject("PDF_PARSE_TIMEOUT")
        except (OSError, UnicodeError, ValueError):
            _reject("PDF_PARSER_UNAVAILABLE")

        worker_output = completed.stdout
        if result_path:
            try:
                worker_output = Path(result_path).read_text(encoding="utf-8")
            except (OSError, UnicodeError):
                _reject("PDF_OUTPUT_INVALID")
        if len(worker_output) > max_chars * 8 + 4096:
            _reject("PDF_PARSE_FAILED")
        try:
            result = json.loads(worker_output)
        except (json.JSONDecodeError, TypeError, ValueError):
            _reject("PDF_OUTPUT_INVALID")
        if not isinstance(result, dict):
            _reject("PDF_OUTPUT_INVALID")
        if completed.returncode != 0 or result.get("ok") is not True:
            code = str(result.get("code") or "PDF_PARSE_FAILED")
            if code not in _SAFE_MESSAGES or not code.startswith("PDF_"):
                code = "PDF_PARSE_FAILED"
            _reject(code)
        text = result.get("text")
        if not isinstance(text, str):
            _reject("PDF_OUTPUT_INVALID")
        return text[:max_chars]
    finally:
        if temp_path:
            try:
                os.unlink(temp_path)
            except OSError:
                pass
        if result_path:
            try:
                os.unlink(result_path)
            except OSError:
                pass


def _apply_worker_limits() -> None:
    if os.name == "nt":
        return
    try:
        import resource

        memory = 768 * MIB
        resource.setrlimit(resource.RLIMIT_AS, (memory, memory))
        resource.setrlimit(resource.RLIMIT_CPU, (15, 15))
        resource.setrlimit(resource.RLIMIT_NOFILE, (64, 64))
        resource.setrlimit(resource.RLIMIT_FSIZE, (2 * MIB, 2 * MIB))
    except (ImportError, OSError, ValueError):
        # The parent still enforces input/output limits and a hard timeout.
        return


def _pdf_worker_main(argv: list[str]) -> int:
    output_path: Path | None = None
    if len(argv) == 6 and argv[0] == "--pdf-worker-output":
        output_path = Path(argv[5])
        if (
            not output_path.name.startswith("defense-pdf-result-")
            or output_path.suffix != ".json"
            or not output_path.is_file()
            or output_path.stat().st_size != 0
        ):
            return 2
        argv = ["--pdf-worker", *argv[1:5]]
    if len(argv) != 5 or argv[0] != "--pdf-worker":
        return 2

    def emit(payload: dict[str, object]) -> None:
        serialized = json.dumps(payload, ensure_ascii=True)
        if output_path is None:
            print(serialized)
        else:
            output_path.write_text(serialized, encoding="utf-8")

    try:
        path = Path(argv[1])
        max_pages = int(argv[2])
        max_total_pages = int(argv[3])
        max_chars = int(argv[4])
        _apply_worker_limits()
        import pdfplumber

        parts: list[str] = []
        chars = 0
        with pdfplumber.open(str(path)) as pdf:
            if len(pdf.pages) > max_total_pages:
                raise DocumentSafetyError("PDF_TOO_COMPLEX")
            for page in pdf.pages[:max_pages]:
                page_text = page.extract_text() or ""
                remaining = max_chars - chars
                if remaining <= 0:
                    break
                page_text = page_text[:remaining]
                parts.append(page_text)
                chars += len(page_text) + 1
        text = "\n".join(parts)[:max_chars]
        emit({"ok": True, "text": text})
        return 0
    except ImportError:
        emit({"ok": False, "code": "PDF_PARSER_UNAVAILABLE"})
        return 3
    except DocumentSafetyError as exc:
        emit({"ok": False, "code": exc.code})
        return 3
    except Exception:
        emit({"ok": False, "code": "PDF_PARSE_FAILED"})
        return 3


if __name__ == "__main__":  # pragma: no cover - exercised through subprocess
    raise SystemExit(_pdf_worker_main(sys.argv[1:]))
