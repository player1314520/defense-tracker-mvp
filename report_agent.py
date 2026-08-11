# -*- coding: utf-8 -*-
"""Lightweight report Agent core for Alpha 1.0.

This module owns project/evidence/draft persistence and prompt construction.
Flask routes pass in news selection and LLM call functions so this file stays
free of app.py imports.
"""
import hashlib
import json
import os
import re
import sqlite3
import threading
import uuid
from datetime import datetime, timezone
from io import BytesIO
from pathlib import Path

from state import DATA_DIR

try:
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor
    from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    Document = None
    Pt = None
    Cm = None
    RGBColor = None
    WD_TABLE_ALIGNMENT = None
    WD_CELL_VERTICAL_ALIGNMENT = None
    WD_ALIGN_PARAGRAPH = None
    WD_LINE_SPACING = None
    OxmlElement = None
    qn = None


REPORT_AGENT_DB_FILE = os.path.join(DATA_DIR, "report_agent.sqlite3")
_DB_LOCK = threading.Lock()
REPORT_AGENT_WRITING_SPEC_FILE = os.environ.get("REPORT_AGENT_WRITING_SPEC_FILE", "")
_DEFAULT_DEFENSETRACKER_SOD_FILE = os.path.join(
    str(Path(__file__).resolve().parent),
    "docs",
    "defensetracker_sod_writing.md",
)
_WRITING_SPEC_MAX_CHARS = 6000

REPORT_TYPE_DEFAULTS = {
    "strategic": {"target_count": 12, "time_window_days": 14, "label": "战略分析报告"},
    "daily": {"target_count": 5, "time_window_days": 1, "label": "每日简报"},
    "weekly": {"target_count": 8, "time_window_days": 7, "label": "周报汇编"},
    "short_topic": {"target_count": 5, "time_window_days": 2, "label": "专题短报"},
}

DEFAULT_AGENT_VOICE = "strategic_analysis"
NEWSPAPER_VOICE = "newspaper"
AGENT_VOICES = (DEFAULT_AGENT_VOICE, NEWSPAPER_VOICE)
MIN_EVIDENCE_TARGET = 1
FORBIDDEN_CLASSIFICATION_TERMS = ("绝密", "机密", "秘密")
LONG_REPORT_MIN_RATIO = 0.85
CHINA_FOCUSED_SOURCE_CATEGORY_IDS = {"china_zone", "pla_research", "us_eu_china_analysis"}
NON_CHINA_SCOPE_MARKERS = {
    "iran": ("伊朗", "iran", "tehran", "irgc", "revolutionary guard"),
    "middle_east": ("中东", "波斯湾", "以色列", "加沙", "胡塞", "真主党", "middle east", "persian gulf", "israel", "gaza", "houthi", "hezbollah"),
    "russia": ("俄罗斯", "俄乌", "乌克兰", "russia", "ukraine"),
    "india": ("印度", "india"),
    "korea": ("朝鲜", "半岛", "dprk", "north korea", "korea"),
}
CHINA_SCOPE_MARKERS = ("中国", "台海", "台湾", "南海", "解放军", "pla", "china", "taiwan", "south china sea")
DEFENSE_TOPIC_MARKERS = {
    "missile": ("导弹", "弹道", "巡航导弹", "高超声速", "missile", "ballistic", "cruise missile", "hypersonic"),
    "nuclear": ("核", "核力量", "nuclear", "warhead", "deterrence"),
    "uav": ("无人机", "无人系统", "uav", "drone", "unmanned"),
}

_SOD_WRITING_FALLBACK = """DefenseTracker SOD/SOP写作要求（内置摘要）：
1. 先将客户命题解构为可回答的研究问题，归一到技术—能力—规则三维分析框架。
2. 资料采集遵循OSINT五源体系，优先官方文件、权威智库、企业披露、防务媒体、学术期刊。
3. 关键判断必须经过三角交叉验证，避免单源孤证和未经核实传闻。
4. 报告结构应包含研究问题、方法框架、证据矩阵、能力态势、战略影响、风险预警、对策建议、来源附录。
5. 章节遵循“章引言—主体分析—关键案例—综合研判—启示建议”五要素。
6. 段落遵循PARA结构：Point论断、Argument论证、References证据、Application战略含义。
7. 关键判断建立FACT-DATA-CITE证据链，引用官方文件、GAO/CRS、智库报告、企业披露或防务媒体。
8. 数据呈现优先使用表格、矩阵、时间线；避免模糊量化和断裂论证。
9. 文风采取严谨军事理论/战略研究风格，减少英语化表达，缩略语首次出现给出中文释义。
10. 证据不足处必须标注待核实或后续跟踪，不得补写虚构数据、来源和结论。"""


def sanitize_report_text(value: str) -> str:
    text = str(value or "")
    for term in FORBIDDEN_CLASSIFICATION_TERMS:
        text = text.replace(term, "")
    text = re.sub(r"[ \t]+([，。；：、,.!?！？])", r"\1", text)
    text = re.sub(r"([（(])\s+", r"\1", text)
    text = re.sub(r"\s+([）)])", r"\1", text)
    return text.strip()


def has_forbidden_classification_terms(value: str) -> bool:
    return any(term in (value or "") for term in FORBIDDEN_CLASSIFICATION_TERMS)


def report_word_count(value: str) -> int:
    text = sanitize_report_text(value)
    text = re.sub(r"https?://\S+", "", text)
    text = re.sub(r"[#*_`|\\\-\s，。；：、,.!?！？（）()《》“”\"'：:]+", "", text)
    return len(text)


def _parse_chinese_number_under_100(text: str) -> int | None:
    if not text:
        return None
    digits = {"零": 0, "一": 1, "二": 2, "两": 2, "三": 3, "四": 4, "五": 5,
              "六": 6, "七": 7, "八": 8, "九": 9}
    if text in digits:
        return digits[text]
    if text == "十":
        return 10
    if "十" in text:
        left, right = text.split("十", 1)
        tens = digits.get(left, 1 if left == "" else None)
        ones = digits.get(right, 0 if right == "" else None)
        if tens is not None and ones is not None:
            return tens * 10 + ones
    return None


def extract_target_word_count(*values) -> int | None:
    text = " ".join(str(v or "") for v in values)
    patterns = [
        r"(\d+(?:\.\d+)?)\s*万\s*字",
        r"([一二两三四五六七八九十]{1,4})\s*万\s*字",
        r"(?:不少于|至少|正文不少于|目标|字数|要|约|达到|扩写到|变成)?\s*(\d{4,6})\s*字",
    ]
    for idx, pattern in enumerate(patterns):
        match = re.search(pattern, text)
        if not match:
            continue
        value = match.group(1)
        if idx == 0:
            return int(float(value) * 10000)
        if idx == 1:
            parsed = _parse_chinese_number_under_100(value)
            return parsed * 10000 if parsed else None
        return int(value)
    return None


def recommended_max_tokens_for_target(target_word_count: int | None) -> int | None:
    if not target_word_count:
        return None
    return max(4096, min(24000, int(target_word_count * 1.6)))


def report_quality_payload(content: str, target_word_count: int | None = None) -> dict:
    cleaned = sanitize_report_text(content)
    count = report_word_count(cleaned)
    min_count = int(target_word_count * LONG_REPORT_MIN_RATIO) if target_word_count else 0
    return {
        "word_count": count,
        "target_word_count": int(target_word_count or 0),
        "min_required_word_count": min_count,
        "word_count_ok": bool(not target_word_count or count >= min_count),
        "forbidden_terms_removed": cleaned != (content or ""),
        "forbidden_terms_present": has_forbidden_classification_terms(cleaned),
    }


def assert_report_exportable(draft: dict, project: dict | None = None):
    payload = draft.get("payload") or {}
    content = draft.get("content") or ""
    # 字数目标只从用户意图（已存 payload 或客户需求）解析，绝不从报告正文 content 解析：
    # 否则正文里一句"可扩写至8000字/全文约3万字"会被误判为目标，把用户从未要求字数的报告永久锁死导出。
    target = int(
        payload.get("target_word_count")
        or extract_target_word_count((project or {}).get("client_request", ""))
        or 0
    )
    quality = report_quality_payload(content, target)
    if quality["forbidden_terms_present"]:
        raise ValueError("报告仍包含禁止使用的涉密等级字眼，已阻断导出")
    if target and not quality["word_count_ok"]:
        raise ValueError(
            f"当前正文约{quality['word_count']}字，低于目标字数{target}字"
            f"（最低要求{quality['min_required_word_count']}字），请继续生成或扩写后再导出"
        )
    return quality


def _writing_spec_candidates() -> list[str]:
    candidates = [
        REPORT_AGENT_WRITING_SPEC_FILE,
        _DEFAULT_DEFENSETRACKER_SOD_FILE,
        os.path.join(Path(__file__).resolve().parent, "docs", "report_agent_sod_writing.md"),
    ]
    return [p for p in candidates if p]


def _extract_relevant_sod_lines(markdown: str) -> str:
    keep_sections = ("SOD-1", "SOD-2", "SOD-4", "SOD-5", "SOD-6", "SOD-10")
    keep_keywords = (
        "技术—能力—规则", "三维", "OSINT", "五源", "三角", "交叉验证",
        "章节", "五要素", "PARA", "FACT-DATA-CITE", "引用规范", "数据可视化",
        "论证语气", "写作禁忌", "检查清单", "必须", "避免", "Step",
    )
    lines: list[str] = []
    capture = False
    for raw in markdown.splitlines():
        line = raw.rstrip()
        if line.startswith("## "):
            capture = any(token in line for token in keep_sections)
        if not capture:
            continue
        if (
            line.startswith(("## ", "### ", "#### ", "- ", "|"))
            or any(keyword in line for keyword in keep_keywords)
        ):
            lines.append(line)
        if sum(len(x) + 1 for x in lines) >= _WRITING_SPEC_MAX_CHARS:
            break
    return "\n".join(lines).strip()


def load_report_writing_requirements() -> str:
    for path in _writing_spec_candidates():
        if not os.path.exists(path):
            continue
        try:
            with open(path, "r", encoding="utf-8") as f:
                text = f.read().strip()
        except UnicodeDecodeError:
            with open(path, "r", encoding="utf-8-sig") as f:
                text = f.read().strip()
        if not text:
            continue
        excerpt = text if len(text) <= _WRITING_SPEC_MAX_CHARS else _extract_relevant_sod_lines(text)
        excerpt = (excerpt or text[:_WRITING_SPEC_MAX_CHARS]).strip()
        return f"DefenseTracker SOD/SOP写作要求来源：{path}\n{excerpt[:_WRITING_SPEC_MAX_CHARS]}"
    return _SOD_WRITING_FALLBACK


def _now() -> str:
    return datetime.now(timezone.utc).isoformat()


def _format_chinese_date(value) -> str:
    return f"{value.year:04d}年{value.month:02d}月{value.day:02d}日"


def _json_dumps(value) -> str:
    return json.dumps(value, ensure_ascii=False, default=str)


def _json_loads(value, default):
    if not value:
        return default
    try:
        return json.loads(value)
    except Exception:
        return default


def _new_id(prefix: str) -> str:
    return f"{prefix}_{uuid.uuid4().hex[:12]}"


def _connect():
    os.makedirs(os.path.dirname(REPORT_AGENT_DB_FILE), exist_ok=True)
    conn = sqlite3.connect(REPORT_AGENT_DB_FILE, timeout=10)
    conn.row_factory = sqlite3.Row
    conn.execute("PRAGMA journal_mode=WAL")      # 多读单写并发，缓解 database is locked
    conn.execute("PRAGMA busy_timeout=10000")     # 写争用自动重试 10s 而非立即报错
    conn.execute("PRAGMA synchronous=NORMAL")
    return conn


def init_report_agent_db():
    with _DB_LOCK, _connect() as conn:
        conn.executescript("""
        CREATE TABLE IF NOT EXISTS projects (
            project_id TEXT PRIMARY KEY,
            title TEXT NOT NULL,
            report_type TEXT NOT NULL,
            topic TEXT NOT NULL DEFAULT '',
            client_request TEXT NOT NULL DEFAULT '',
            time_window_days INTEGER NOT NULL,
            target_count INTEGER NOT NULL,
            voice TEXT NOT NULL DEFAULT 'strategic_analysis',
            status TEXT NOT NULL DEFAULT 'active',
            created_at TEXT NOT NULL,
            updated_at TEXT NOT NULL
        );
        CREATE TABLE IF NOT EXISTS evidence (
            project_id TEXT NOT NULL,
            evidence_id TEXT NOT NULL,
            article_id TEXT NOT NULL,
            title TEXT NOT NULL,
            summary TEXT NOT NULL DEFAULT '',
            source TEXT NOT NULL DEFAULT '',
            source_cn TEXT NOT NULL DEFAULT '',
            link TEXT NOT NULL DEFAULT '',
            date TEXT NOT NULL DEFAULT '',
            quality_score INTEGER NOT NULL DEFAULT 0,
            quality_level TEXT NOT NULL DEFAULT '',
            quality_reasons_json TEXT NOT NULL DEFAULT '[]',
            brief_hits_json TEXT NOT NULL DEFAULT '[]',
            payload_json TEXT NOT NULL,
            selected INTEGER NOT NULL DEFAULT 0,
            created_at TEXT NOT NULL,
            updated_at TEXT NOT NULL,
            PRIMARY KEY(project_id, evidence_id),
            FOREIGN KEY(project_id) REFERENCES projects(project_id)
        );
        CREATE TABLE IF NOT EXISTS drafts (
            draft_id TEXT PRIMARY KEY,
            project_id TEXT NOT NULL,
            kind TEXT NOT NULL,
            content TEXT NOT NULL,
            model TEXT NOT NULL DEFAULT '',
            source_draft_id TEXT NOT NULL DEFAULT '',
            payload_json TEXT NOT NULL DEFAULT '{}',
            created_at TEXT NOT NULL,
            FOREIGN KEY(project_id) REFERENCES projects(project_id)
        );
        CREATE TABLE IF NOT EXISTS events (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            project_id TEXT NOT NULL,
            event_type TEXT NOT NULL,
            payload_json TEXT NOT NULL,
            created_at TEXT NOT NULL,
            FOREIGN KEY(project_id) REFERENCES projects(project_id)
        );
        CREATE TABLE IF NOT EXISTS draft_jobs (
            job_id TEXT PRIMARY KEY,
            project_id TEXT NOT NULL,
            status TEXT NOT NULL DEFAULT 'queued',
            draft_id TEXT NOT NULL DEFAULT '',
            error TEXT NOT NULL DEFAULT '',
            request_json TEXT NOT NULL DEFAULT '{}',
            created_at TEXT NOT NULL,
            updated_at TEXT NOT NULL,
            FOREIGN KEY(project_id) REFERENCES projects(project_id)
        );
        """)
        columns = {row["name"] for row in conn.execute("PRAGMA table_info(projects)").fetchall()}
        if "client_request" not in columns:
            conn.execute("ALTER TABLE projects ADD COLUMN client_request TEXT NOT NULL DEFAULT ''")
        draft_columns = {row["name"] for row in conn.execute("PRAGMA table_info(drafts)").fetchall()}
        if "payload_json" not in draft_columns:
            conn.execute("ALTER TABLE drafts ADD COLUMN payload_json TEXT NOT NULL DEFAULT '{}'")
        conn.commit()


def _project_from_row(row) -> dict:
    keys = set(row.keys())
    return {
        "project_id": row["project_id"],
        "title": row["title"],
        "report_type": row["report_type"],
        "topic": row["topic"],
        "client_request": row["client_request"] if "client_request" in keys else "",
        "time_window_days": row["time_window_days"],
        "target_count": row["target_count"],
        "voice": row["voice"],
        "status": row["status"],
        "created_at": row["created_at"],
        "updated_at": row["updated_at"],
    }


# ── 证据一句话研判（确定性拼装，零 AI）：把已算出的打分信号翻成人话，供分析师 2 秒 triage。
# 规则/启发式产物，非情报结论——前端标注"规则研判·供参考"；全局规则5：规则/AI 输出永远不是来源。
_VERDICT_DIM_LABELS = {
    "source": "高权威信源",
    "topic": "高度契合选题",
    "density": "信息密度高",
    "novelty": "含首发/新型信号",
    "writability": "可直接改写要讯",
}


def _evidence_recency_hint(date_str: str) -> str:
    try:
        age_h = (datetime.now(timezone.utc) - datetime.fromisoformat(date_str)).total_seconds() / 3600
    except Exception:
        return ""
    if age_h < 0:
        return ""
    if age_h < 6:
        return "6小时内"
    if age_h < 24:
        return "24小时内"
    if age_h < 72:
        return "近3天"
    return ""


def _compose_evidence_verdict(ev: dict) -> str:
    """从已算出的打分信号确定性拼一句 ≤40 字研判，帮助快速 triage。不调用任何 AI。"""
    dims = ev.get("dims") or (ev.get("payload") or {}).get("quality", {}).get("dims") or {}
    hits = [h for h in (ev.get("brief_hits") or []) if h]
    reasons = [r for r in (ev.get("quality_reasons") or []) if r]
    source_type = ev.get("source_type") or "公开信息"
    level = ev.get("quality_level") or ""

    # 1) 主信号：优先取最高维；无维度（智库卡 / 转入原文）退回 source_type / 理由
    if dims:
        top_dim = max(dims, key=lambda k: dims.get(k) or 0)
        main = _VERDICT_DIM_LABELS.get(top_dim, "防务相关")
    elif source_type == "智库/报告源":
        main = f"权威智库源·{reasons[1]}" if len(reasons) > 1 else "权威智库/报告源"
    elif source_type == "已抓取公开报告/原文":
        main = "已获取原文可溯源"
    elif reasons:
        main = reasons[0]
    else:
        main = "基础防务相关"

    segs = [main]
    if hits:
        segs.append("命中" + "、".join(hits[:2]))
    else:
        rec = _evidence_recency_hint(ev.get("date") or "")
        if rec:
            segs.append(rec)

    line = "，".join(segs)
    if level:
        line = f"{level}级·{line}"
    if len(line) > 40:
        line = line[:39] + "…"
    return line


def _evidence_from_row(row) -> dict:
    payload = _json_loads(row["payload_json"], {})
    quality = payload.get("quality") or {}
    ev = {
        "project_id": row["project_id"],
        "evidence_id": row["evidence_id"],
        "article_id": row["article_id"],
        "title": row["title"],
        "summary": row["summary"],
        "source": row["source"],
        "source_cn": row["source_cn"],
        "link": row["link"],
        "date": row["date"],
        "quality_score": row["quality_score"],
        "quality_level": row["quality_level"],
        "quality_reasons": _json_loads(row["quality_reasons_json"], []),
        "brief_hits": _json_loads(row["brief_hits_json"], []),
        "source_type": payload.get("source_type") or "公开信息",
        # 已算出但此前被丢弃的 5 维明细 / 惩罚项——透出以支撑"打分透明化"
        "dims": quality.get("dims") or {},
        "penalties": quality.get("penalties") or payload.get("quality_penalties") or [],
        # 多源印证：新闻刷新时 _dedup_articles 已把同标题跨源合并进 _sources（0=不适用，如智库目录源）
        "corroboration_count": len(payload.get("_sources") or []),
        "corroborating_sources": payload.get("_sources") or [],
        "selected": bool(row["selected"]),
        "payload": payload,
        "created_at": row["created_at"],
        "updated_at": row["updated_at"],
    }
    ev["verdict_line"] = _compose_evidence_verdict(ev)
    return ev


def _draft_from_row(row) -> dict:
    payload = _json_loads(row["payload_json"], {}) if "payload_json" in row.keys() else {}
    return {
        "draft_id": row["draft_id"],
        "project_id": row["project_id"],
        "kind": row["kind"],
        "content": sanitize_report_text(row["content"]),
        "model": row["model"],
        "source_draft_id": row["source_draft_id"],
        "payload": payload,
        "created_at": row["created_at"],
    }


def _event_from_row(row) -> dict:
    return {
        "id": row["id"],
        "project_id": row["project_id"],
        "event_type": row["event_type"],
        "payload": _json_loads(row["payload_json"], {}),
        "created_at": row["created_at"],
    }


def _log_event(conn, project_id: str, event_type: str, payload=None):
    conn.execute(
        "INSERT INTO events(project_id, event_type, payload_json, created_at) VALUES (?, ?, ?, ?)",
        (project_id, event_type, _json_dumps(payload or {}), _now()),
    )


def _clean_text(value: str) -> str:
    return re.sub(r"\s+", " ", (value or "").strip())


def _derive_topic_from_request(client_request: str) -> str:
    text = _clean_text(client_request)
    text = re.sub(r"[，,；;。]?\s*(搜集|收集|检索|抓取|找|列出)?\s*\d{1,3}\s*(份|条|个)?\s*(信息源|证据|资料|来源|智库|报告源|报告)?", "", text)
    text = re.sub(r"^[请麻烦]*帮[我忙]?(做|写|生成|出)?(一个|一份|1份)?", "", text)
    text = re.sub(r"^(请|麻烦|做|写|生成|出)(一个|一份|1份)?", "", text)
    text = re.sub(r"^(有关|关于|围绕)", "", text)
    text = re.sub(r"(的)?(高价值)?(防务)?(战略)?(深度)?(综合)?(分析)?报告$", "", text)
    text = re.sub(r"[，。；;,.!！?？：:\-—]+", "", text).strip()
    return text or "综合防务态势"


def _extract_requested_count(client_request: str) -> int | None:
    text = _clean_text(client_request)
    patterns = [
        r"(?:搜集|收集|检索|抓取|找|列出)\s*(\d{1,6})\s*(?:份|条|个)?\s*(?:信息源|证据|资料|来源|智库|报告源|报告)?",
        r"(\d{1,6})\s*(?:份|条|个)\s*(?:信息源|证据|资料|来源|智库|报告源|报告)",
    ]
    for pattern in patterns:
        match = re.search(pattern, text)
        if match:
            value = int(match.group(1))
            return max(MIN_EVIDENCE_TARGET, value)
    return None


def _derive_topic_from_title(title: str) -> str:
    text = _clean_text(title)
    text = re.sub(r"(的)?(高价值)?(防务)?(战略)?(深度)?(综合)?(分析)?报告$", "", text)
    return text.strip("：: -—") or "综合防务态势"


def create_project(title: str = "", report_type: str = "strategic", topic: str = "",
                   time_window_days: int | None = None,
                   target_count: int | None = None, voice: str = DEFAULT_AGENT_VOICE,
                   client_request: str = "") -> dict:
    title = (title or "").strip()
    client_request = _clean_text(client_request)
    report_type = (report_type or "strategic").strip()
    if report_type not in REPORT_TYPE_DEFAULTS:
        raise ValueError("无效报告类型")
    if not topic and client_request:
        topic = _derive_topic_from_request(client_request)
    if not title and topic:
        title = f"{topic}战略分析报告"
    if not title and client_request:
        topic = _derive_topic_from_request(client_request)
        title = f"{topic}战略分析报告"
    if not title:
        raise ValueError("缺少项目标题或客户需求")
    if not topic:
        topic = _derive_topic_from_title(title)
    defaults = REPORT_TYPE_DEFAULTS[report_type]
    time_window_days = int(time_window_days or defaults["time_window_days"])
    requested_count = _extract_requested_count(client_request) if client_request else None
    target_count = int(target_count or requested_count or defaults["target_count"])
    target_count = max(MIN_EVIDENCE_TARGET, target_count)
    voice = voice or DEFAULT_AGENT_VOICE
    project_id = _new_id("rp")
    now = _now()
    init_report_agent_db()
    with _DB_LOCK, _connect() as conn:
        conn.execute(
            """
            INSERT INTO projects
            (project_id, title, report_type, topic, client_request, time_window_days, target_count, voice, status, created_at, updated_at)
            VALUES (?, ?, ?, ?, ?, ?, ?, ?, 'active', ?, ?)
            """,
            (project_id, title, report_type, (topic or "").strip(), client_request, time_window_days, target_count, voice, now, now),
        )
        _log_event(conn, project_id, "project_created", {"report_type": report_type, "client_request": client_request})
        conn.commit()
    return get_project(project_id)


def get_project(project_id: str) -> dict:
    init_report_agent_db()
    with _DB_LOCK, _connect() as conn:
        row = conn.execute("SELECT * FROM projects WHERE project_id=?", (project_id,)).fetchone()
    if not row:
        raise KeyError("项目不存在")
    return _project_from_row(row)


def list_projects(limit: int = 50) -> list[dict]:
    init_report_agent_db()
    limit = max(1, min(int(limit or 50), 100))
    with _DB_LOCK, _connect() as conn:
        rows = conn.execute(
            "SELECT * FROM projects ORDER BY updated_at DESC LIMIT ?", (limit,)
        ).fetchall()
    return [_project_from_row(r) for r in rows]


def _candidate_evidence_id(candidate: dict) -> tuple[str, str]:
    article_id = (candidate.get("article_id") or "").strip()
    if not article_id:
        basis = "|".join([
            candidate.get("link", ""),
            candidate.get("source", ""),
            candidate.get("title", ""),
            candidate.get("date", ""),
        ])
        article_id = hashlib.sha256(basis.encode("utf-8", errors="ignore")).hexdigest()[:20]
    return article_id, article_id


def upsert_project_evidence(project_id: str, candidates: list[dict]) -> list[dict]:
    get_project(project_id)
    init_report_agent_db()
    now = _now()
    ids = []
    with _DB_LOCK, _connect() as conn:
        for item in candidates or []:
            article_id, evidence_id = _candidate_evidence_id(item)
            ids.append(evidence_id)
            reasons = item.get("quality_reasons") or item.get("quality", {}).get("reasons") or []
            hits = item.get("brief_hits") or item.get("quality", {}).get("brief_hits") or []
            conn.execute(
                """
                INSERT INTO evidence
                (project_id, evidence_id, article_id, title, summary, source, source_cn, link, date,
                 quality_score, quality_level, quality_reasons_json, brief_hits_json, payload_json,
                 selected, created_at, updated_at)
                VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, 0, ?, ?)
                ON CONFLICT(project_id, evidence_id) DO UPDATE SET
                    title=excluded.title,
                    summary=excluded.summary,
                    source=excluded.source,
                    source_cn=excluded.source_cn,
                    link=excluded.link,
                    date=excluded.date,
                    quality_score=excluded.quality_score,
                    quality_level=excluded.quality_level,
                    quality_reasons_json=excluded.quality_reasons_json,
                    brief_hits_json=excluded.brief_hits_json,
                    payload_json=excluded.payload_json,
                    updated_at=excluded.updated_at
                """,
                (
                    project_id,
                    evidence_id,
                    article_id,
                    item.get("title", ""),
                    item.get("summary", ""),
                    item.get("source", ""),
                    item.get("source_cn", ""),
                    item.get("link", ""),
                    item.get("date", ""),
                    int(item.get("quality_score") or item.get("quality", {}).get("total") or 0),
                    item.get("quality_level") or item.get("quality", {}).get("level") or "",
                    _json_dumps(reasons),
                    _json_dumps(hits),
                    _json_dumps(item),
                    now,
                    now,
                ),
            )
        conn.execute("UPDATE projects SET updated_at=? WHERE project_id=?", (now, project_id))
        _log_event(conn, project_id, "evidence_collected", {"count": len(ids)})
        conn.commit()
    return get_project_evidence(project_id, ids)


def _source_topic_text(project: dict) -> str:
    return " ".join([
        project.get("topic") or "",
        project.get("title") or "",
        project.get("client_request") or "",
    ]).lower()


def _source_topic_scopes(text: str) -> set[str]:
    scopes = set()
    for scope, markers in NON_CHINA_SCOPE_MARKERS.items():
        if any(marker.lower() in text for marker in markers):
            scopes.add(scope)
    if any(marker.lower() in text for marker in CHINA_SCOPE_MARKERS):
        scopes.add("china")
    for scope, markers in DEFENSE_TOPIC_MARKERS.items():
        if any(marker.lower() in text for marker in markers):
            scopes.add(scope)
    return scopes


def _source_topic_terms(text: str) -> list[str]:
    terms = []
    marker_groups = [CHINA_SCOPE_MARKERS, *NON_CHINA_SCOPE_MARKERS.values(), *DEFENSE_TOPIC_MARKERS.values()]
    for markers in marker_groups:
        for marker in markers:
            marker_l = marker.lower()
            if marker_l in text and marker_l not in terms:
                terms.append(marker_l)
    for word in re.findall(r"[a-z][a-z0-9\-]{2,}", text):
        if word not in terms:
            terms.append(word)
    return terms


def _is_china_focused_source(category_id: str, haystack: str) -> bool:
    if category_id in CHINA_FOCUSED_SOURCE_CATEGORY_IDS:
        return True
    return any(marker in haystack for marker in ("china", "中国", "pla", "解放军", "台海", "taiwan"))


def _source_relevance_score(category_id: str, haystack: str, terms: list[str], scopes: set[str]) -> int | None:
    score = 0
    direct_hits = sum(1 for term in terms if term and term in haystack)
    score += direct_hits * 4

    non_china_scopes = scopes.intersection(NON_CHINA_SCOPE_MARKERS)
    if non_china_scopes and "china" not in scopes:
        if _is_china_focused_source(category_id, haystack):
            non_china_hit = any(
                marker.lower() in haystack
                for scope in non_china_scopes
                for marker in NON_CHINA_SCOPE_MARKERS[scope]
            )
            if not non_china_hit:
                return None
        for scope in non_china_scopes:
            if any(marker.lower() in haystack for marker in NON_CHINA_SCOPE_MARKERS[scope]):
                score += 12

    if "missile" in scopes and any(marker.lower() in haystack for marker in DEFENSE_TOPIC_MARKERS["missile"]):
        score += 8
    if "nuclear" in scopes and any(marker.lower() in haystack for marker in DEFENSE_TOPIC_MARKERS["nuclear"]):
        score += 5
    if "uav" in scopes and any(marker.lower() in haystack for marker in DEFENSE_TOPIC_MARKERS["uav"]):
        score += 5
    if "china" in scopes and _is_china_focused_source(category_id, haystack):
        score += 8
    if category_id in {"global_media", "official", "missile_mideast_research"}:
        score += 2

    if terms and direct_hits == 0 and non_china_scopes:
        return None
    return score


def build_source_candidates(project: dict, thinktank_directory: list[dict],
                            limit: int = 10, level_fn=None) -> list[dict]:
    """Create reviewable evidence cards for curated think-tank/report sources."""
    topic = (project.get("topic") or project.get("title") or "综合防务态势").strip()
    query = _source_topic_text({**project, "topic": topic})
    terms = _source_topic_terms(query)
    scopes = _source_topic_scopes(query)
    rows = []
    for category in thinktank_directory or []:
        category_name = category.get("category") or category.get("category_en") or "智库/报告源"
        category_id = category.get("id") or ""
        for site in category.get("sites") or []:
            name = site.get("name") or site.get("name_cn") or "公开源"
            name_cn = site.get("name_cn") or name
            desc = site.get("desc_cn") or site.get("desc_en") or category.get("desc") or ""
            url = site.get("url") or ""
            haystack = f"{name} {name_cn} {desc} {category_name} {category_id}".lower()
            relevance = _source_relevance_score(category_id, haystack, terms, scopes)
            if relevance is None:
                continue
            score = 1 + relevance
            if category_id in {"official", "global_media", "missile_mideast_research"}:
                score += 2
            if category_id in {"pla_research", "us_eu_china_analysis"} and "china" in scopes:
                score += 2
            source_qscore = min(90, 72 + score * 3)
            rows.append((score, {
                "article_id": f"source-{hashlib.sha256((url or name).encode('utf-8', errors='ignore')).hexdigest()[:20]}",
                "title": f"{name_cn}：{desc or category_name}",
                "summary": f"围绕“{topic}”可检索该智库/报告源的公开报告、专题数据库、政策分析与长期研究积累。类别：{category_name}。",
                "source": name,
                "source_cn": name_cn,
                "link": url,
                "date": "",
                "quality_score": source_qscore,
                # 分级由分数驱动（默认回退 "A" 保持既有调用/测试不变；报告池路径注入 level_fn 取真实分级）
                "quality_level": level_fn(source_qscore) if level_fn else "A",
                "quality_reasons": ["智库/报告源", category_name],
                "brief_hits": [],
                "source_type": "智库/报告源",
            }))
    rows.sort(key=lambda item: (-item[0], item[1]["source"]))
    return [item for _, item in rows[:max(1, int(limit or 10))]]


def get_project_evidence(project_id: str, evidence_ids: list[str] | None = None) -> list[dict]:
    init_report_agent_db()
    with _DB_LOCK, _connect() as conn:
        if evidence_ids:
            placeholders = ",".join("?" for _ in evidence_ids)
            rows = conn.execute(
                f"SELECT * FROM evidence WHERE project_id=? AND evidence_id IN ({placeholders})",
                [project_id, *evidence_ids],
            ).fetchall()
            by_id = {r["evidence_id"]: _evidence_from_row(r) for r in rows}
            return [by_id[eid] for eid in evidence_ids if eid in by_id]
        rows = conn.execute(
            "SELECT * FROM evidence WHERE project_id=? ORDER BY quality_score DESC, updated_at DESC",
            (project_id,),
        ).fetchall()
    return [_evidence_from_row(r) for r in rows]


def save_draft(project_id: str, kind: str, content: str, model: str = "",
               source_draft_id: str = "", payload: dict | None = None) -> dict:
    get_project(project_id)
    kind = (kind or "draft").strip()
    if kind not in {"outline", "draft", "revision"}:
        raise ValueError("无效草稿类型")
    content = sanitize_report_text(content or "")
    if not content:
        raise ValueError("草稿内容为空")
    payload = payload or {}
    draft_id = _new_id("df")
    with _DB_LOCK, _connect() as conn:
        conn.execute(
            """
            INSERT INTO drafts(draft_id, project_id, kind, content, model, source_draft_id, payload_json, created_at)
            VALUES (?, ?, ?, ?, ?, ?, ?, ?)
            """,
            (draft_id, project_id, kind, content, model or "", source_draft_id or "", _json_dumps(payload), _now()),
        )
        conn.execute("UPDATE projects SET updated_at=? WHERE project_id=?", (_now(), project_id))
        _log_event(conn, project_id, "draft_saved", {"draft_id": draft_id, "kind": kind, **payload})
        conn.commit()
    return get_draft(draft_id)


def get_draft(draft_id: str) -> dict:
    init_report_agent_db()
    with _DB_LOCK, _connect() as conn:
        row = conn.execute("SELECT * FROM drafts WHERE draft_id=?", (draft_id,)).fetchone()
    if not row:
        raise KeyError("草稿不存在")
    return _draft_from_row(row)


# ── 草稿生成 job（阻塞式 AI 迁出请求线程，避免双次 _call_ai 超 gunicorn timeout 丢稿）──

def _draft_job_from_row(row) -> dict:
    return {
        "job_id": row["job_id"],
        "project_id": row["project_id"],
        "status": row["status"],
        "draft_id": row["draft_id"] or "",
        "error": row["error"] or "",
        "request": _json_loads(row["request_json"], {}),
        "created_at": row["created_at"],
        "updated_at": row["updated_at"],
    }


def create_draft_job(project_id: str, request: dict | None = None) -> dict:
    get_project(project_id)
    init_report_agent_db()
    job_id = _new_id("dj")
    now = _now()
    with _DB_LOCK, _connect() as conn:
        conn.execute(
            "INSERT INTO draft_jobs(job_id, project_id, status, request_json, created_at, updated_at) "
            "VALUES (?, ?, 'queued', ?, ?, ?)",
            (job_id, project_id, _json_dumps(request or {}), now, now),
        )
        conn.commit()
    return get_draft_job(project_id, job_id)


def update_draft_job(job_id: str, status: str | None = None, draft_id: str | None = None,
                     error: str | None = None) -> dict:
    init_report_agent_db()
    updates = ["updated_at=?"]
    params = [_now()]
    for field, value in (("status", status), ("draft_id", draft_id), ("error", error)):
        if value is not None:
            updates.append(f"{field}=?")
            params.append(value)
    params.append(job_id)
    with _DB_LOCK, _connect() as conn:
        conn.execute(f"UPDATE draft_jobs SET {', '.join(updates)} WHERE job_id=?", params)
        conn.commit()
        row = conn.execute("SELECT project_id FROM draft_jobs WHERE job_id=?", (job_id,)).fetchone()
    if not row:
        raise KeyError("草稿任务不存在")
    return get_draft_job(row["project_id"], job_id)


def get_draft_job(project_id: str, job_id: str) -> dict:
    init_report_agent_db()
    with _DB_LOCK, _connect() as conn:
        row = conn.execute(
            "SELECT * FROM draft_jobs WHERE project_id=? AND job_id=?", (project_id, job_id)
        ).fetchone()
    if not row:
        raise KeyError("草稿任务不存在")
    return _draft_job_from_row(row)


def get_project_drafts(project_id: str) -> list[dict]:
    init_report_agent_db()
    with _DB_LOCK, _connect() as conn:
        rows = conn.execute(
            "SELECT * FROM drafts WHERE project_id=? ORDER BY created_at DESC", (project_id,)
        ).fetchall()
    return [_draft_from_row(r) for r in rows]


def get_project_events(project_id: str, limit: int | None = None) -> list[dict]:
    init_report_agent_db()
    with _DB_LOCK, _connect() as conn:
        if limit:
            # 有界读取最近 N 条（events 表无 TTL/上限，长寿命项目会单调累积）
            rows = conn.execute(
                "SELECT * FROM events WHERE project_id=? ORDER BY id DESC LIMIT ?", (project_id, int(limit))
            ).fetchall()
            rows = list(reversed(rows))
        else:
            rows = conn.execute(
                "SELECT * FROM events WHERE project_id=? ORDER BY id ASC", (project_id,)
            ).fetchall()
    return [_event_from_row(r) for r in rows]


# ── 报纸式交付（voice=newspaper）：从已生成报告 markdown 确定性抽 front-matter，零 AI，不改生成契约 ──
def estimate_reading_minutes(word_count: int) -> int:
    """中文约 500 字/分钟；至少 1 分钟。"""
    return max(1, round(int(word_count or 0) / 500))


def _newspaper_toc(content: str) -> list:
    toc = []
    for ln in (content or "").splitlines():
        if re.match(r"^\s*###", ln):
            continue
        m = re.match(r"^\s*##\s+(.+?)\s*$", ln)
        if m and m.group(1).strip():
            toc.append(m.group(1).strip())
    return toc


def _newspaper_section(content: str, keyword: str) -> str:
    """取含 keyword 的二级标题段落正文（到下一个二级标题止）。"""
    out, capturing = [], False
    for ln in (content or "").splitlines():
        m = re.match(r"^\s*##\s+(.*)$", ln)
        is_h2 = bool(m) and not re.match(r"^\s*###", ln)
        if is_h2:
            if capturing:
                break
            capturing = keyword in m.group(1)
            continue
        if capturing:
            out.append(ln)
    return "\n".join(out).strip()


def _newspaper_cards(content: str, limit: int = 5) -> list:
    """从『核心判断』段 best-effort 抽 3-5 条判断卡（启发式，非结构化时按句切分）。"""
    section = _newspaper_section(content, "核心判断")
    if not section:
        return []
    cards = []
    for ln in section.splitlines():
        s = re.sub(r"^[\-\*\d\.、）)（(\s]+", "", ln.strip())
        s = re.sub(r"^判断[一二三四五六七八九十\d]+[：:、.]?\s*", "", s)
        if len(s) >= 8:
            cards.append(s[:60])
        if len(cards) >= limit:
            return cards
    if not cards:
        for seg in re.split(r"[。；]", section):
            seg = seg.strip()
            if len(seg) >= 8:
                cards.append(seg[:60])
            if len(cards) >= limit:
                break
    return cards[:limit]


def build_newspaper_front_matter(project: dict, content: str, issue_date: str = "") -> dict:
    project = project or {}
    wc = report_word_count(content)
    toc = _newspaper_toc(content)
    label = REPORT_TYPE_DEFAULTS.get(project.get("report_type"), {}).get(
        "label", project.get("report_type") or "战略报告")
    topic = project.get("topic") or project.get("title") or "综合防务态势"
    date = issue_date or datetime.now(timezone.utc).strftime("%Y.%m.%d")
    return {
        "issue": f"VOL.{date}",
        "byline": f"{label} · {topic}",
        "word_count": wc,
        "reading_minutes": estimate_reading_minutes(wc),
        "section_count": len(toc),
        "toc": toc,
        "cards": _newspaper_cards(content),
    }


def get_project_bundle(project_id: str) -> dict:
    project = get_project(project_id)
    drafts = get_project_drafts(project_id)
    bundle = {
        "project": project,
        "evidence": get_project_evidence(project_id),
        "drafts": drafts,
        # bundle 是前端每次操作后刷新的热路径：只带最近 100 条事件，避免随项目寿命线性膨胀
        "events": get_project_events(project_id, limit=100),
    }
    # 报纸式 front-matter（确定性、零 AI）：取最新报告草稿（非大纲）计算
    latest = next((d for d in drafts if d.get("kind") in ("draft", "revision")), None)
    if latest and latest.get("content"):
        fm = build_newspaper_front_matter(project, latest["content"])
        fm["active"] = (project.get("voice") == NEWSPAPER_VOICE)
        bundle["newspaper"] = fm
    return bundle


def _evidence_lines(evidence: list[dict]) -> str:
    lines = []
    for idx, ev in enumerate(evidence, 1):
        reasons = "、".join(ev.get("quality_reasons") or []) or "基础防务相关"
        source_type = ev.get("source_type") or "公开信息"
        lines.append(
            f"{idx}. 【{source_type}｜{ev.get('source') or '公开来源'}】{ev.get('title')}\n"
            f"   时间：{ev.get('date') or '未知'}\n"
            f"   链接：{ev.get('link') or '无'}\n"
            f"   质量：{ev.get('quality_level') or '-'}级/{ev.get('quality_score') or 0}分；理由：{reasons}\n"
            f"   摘要：{ev.get('summary') or '无摘要'}"
        )
    return "\n".join(lines)


def _project_line(project: dict) -> str:
    label = REPORT_TYPE_DEFAULTS.get(project.get("report_type"), {}).get("label", project.get("report_type"))
    topic = project.get("topic") or "综合防务态势"
    client_request = project.get("client_request") or "未填写"
    return (
        f"报告类型：{label}\n项目标题：{project.get('title')}\n主题：{topic}\n"
        f"客户需求：{client_request}\n目标证据源数量：{project.get('target_count')}"
    )


def build_outline_messages(project: dict, evidence: list[dict], voice: str = DEFAULT_AGENT_VOICE) -> list[dict]:
    writing_requirements = load_report_writing_requirements()
    return [
        {
            "role": "system",
            "content": (
                "你是顶尖防务战略分析报告Agent，面向客户的一句话委托交付高价值防务战略分析报告。"
                "这不是要讯、不是新闻简报、不是素材汇编；必须围绕战略问题、能力态势、长期影响和风险预警建立分析框架。"
                "所有判断必须能追溯到公开源证据、智库报告源或报告线索，不得编造素材未提供的具体数据。"
                "报告写作方法优先遵循DefenseTracker SOD/SOP手册。"
            ),
        },
        {
            "role": "user",
            "content": (
                "请基于以下项目和证据池生成“目录提纲大纲”，输出中文Markdown。\n\n"
                f"{_project_line(project)}\n\n写作要求：\n{writing_requirements}\n\n证据池：\n{_evidence_lines(evidence)}\n\n"
                "大纲必须包含：研究问题、核心判断、证据矩阵、能力与部署态势、战略影响、风险预警、后续跟踪、来源附录。"
            ),
        },
    ]


def build_draft_messages(project: dict, evidence: list[dict], outline: str = "",
                         voice: str = DEFAULT_AGENT_VOICE, review_notes: str = "") -> list[dict]:
    outline_part = outline.strip() if outline else "请先形成战略分析框架，再生成完整报告。"
    review_part = review_notes.strip() if review_notes else "无额外审稿意见。"
    writing_requirements = load_report_writing_requirements()
    target_word_count = extract_target_word_count(project.get("client_request", ""), outline, review_notes)
    target_part = (
        f"目标字数：正文约{target_word_count}字，最低不得少于{int(target_word_count * LONG_REPORT_MIN_RATIO)}字；"
        "必须分章充分展开，不得输出短版摘要。"
        if target_word_count else
        "目标字数：按客户问题复杂度生成完整战略分析报告，不得写成短摘要。"
    )
    return [
        {
            "role": "system",
            "content": (
                "你是一名资深防务战略分析师，负责把公开源信息、智库材料和报告线索整合为高价值防务战略分析报告。"
                "报告要有明确问题意识、证据链、战略判断和可执行跟踪建议；不要写成要讯、新闻简报或列表汇编。"
                "所有判断必须能够从证据池追溯，不得虚构数据、来源和结论。"
                "报告写作方法优先遵循DefenseTracker SOD/SOP手册。"
                "严禁使用任何涉密等级标识字眼，报告必须定位为公开源研究成果。"
            ),
        },
        {
            "role": "user",
            "content": (
                "请生成可直接交付客户审阅的中文Markdown高价值防务战略分析报告。\n\n"
                f"{_project_line(project)}\n{target_part}\n\n写作要求：\n{writing_requirements}\n\n目录提纲大纲：\n{outline_part}\n\n审稿意见：\n{review_part}\n\n"
                f"证据池：\n{_evidence_lines(evidence)}\n\n"
                "输出结构：# 标题、## 执行摘要、## 核心判断、## 证据与来源、## 能力态势分析、"
                "## 战略影响研判、## 风险预警、## 后续跟踪建议、## 来源附录。"
            ),
        },
    ]


def build_revision_messages(project: dict, draft: dict, instruction: str) -> list[dict]:
    target_word_count = extract_target_word_count(
        instruction,
        project.get("client_request", ""),
    ) or int((draft.get("payload") or {}).get("target_word_count") or 0)
    target_part = (
        f"目标字数：正文约{target_word_count}字，最低不得少于{int(target_word_count * LONG_REPORT_MIN_RATIO)}字。"
        if target_word_count else
        "目标字数：保持完整战略分析报告体量，不得压缩为短摘要。"
    )
    return [
        {
            "role": "system",
            "content": "你是防务战略分析报告Agent的审稿助手，只根据用户指令修订草稿，保持战略分析深度、证据链和来源附录。严禁使用任何涉密等级标识字眼。",
        },
        {
            "role": "user",
            "content": (
                f"项目：{project.get('title')}\n{target_part}\n修订要求：{instruction}\n\n"
                f"当前草稿：\n{draft.get('content')}"
            ),
        },
    ]


def build_expansion_messages(project: dict, evidence: list[dict], current_content: str,
                             target_word_count: int, outline: str = "",
                             review_notes: str = "") -> list[dict]:
    return [
        {
            "role": "system",
            "content": (
                "你是顶尖防务战略分析报告扩写助手。必须输出完整替换稿，不要只续写。"
                "严禁使用任何涉密等级标识字眼；只使用公开源证据与证据池材料。"
            ),
        },
        {
            "role": "user",
            "content": (
                f"项目：{project.get('title')}\n目标字数：约{target_word_count}字，最低不得少于{int(target_word_count * LONG_REPORT_MIN_RATIO)}字。\n"
                f"目录提纲：\n{outline or '沿用现有章节框架'}\n\n审稿意见：\n{review_notes or '扩写为完整报告'}\n\n"
                f"证据池：\n{_evidence_lines(evidence)}\n\n当前短稿：\n{sanitize_report_text(current_content)}\n\n"
                "请按DefenseTracker SOD/SOP格式输出完整Markdown替换稿：保留已有判断，扩充分章论证、证据矩阵、技术—能力—规则分析、风险预警和后续跟踪建议。"
            ),
        },
    ]


def _set_run_font(run, east_asia: str, size_pt: float | None = None,
                  bold: bool | None = None, latin: str = "Times New Roman",
                  color: str | None = None):
    run.font.name = latin
    if size_pt:
        run.font.size = Pt(size_pt)
    if bold is not None:
        run.bold = bold
    if color and RGBColor:
        run.font.color.rgb = RGBColor.from_string(color)
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:eastAsia"), east_asia)


def _set_style_font(style, east_asia: str, size_pt: float | None = None,
                    latin: str = "Times New Roman", bold: bool | None = None,
                    color: str | None = None):
    style.font.name = latin
    if size_pt:
        style.font.size = Pt(size_pt)
    if bold is not None:
        style.font.bold = bold
    if color and RGBColor:
        style.font.color.rgb = RGBColor.from_string(color)
    rpr = style._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:eastAsia"), east_asia)


def _format_paragraph(paragraph, alignment=None, first_line: bool = False,
                      left_indent_pt: float | None = None, before_pt: float = 0,
                      after_pt: float = 0, line_pt: float = 28.5):
    pf = paragraph.paragraph_format
    if alignment is not None:
        paragraph.alignment = alignment
    if first_line:
        pf.first_line_indent = Pt(32)
    if left_indent_pt is not None:
        pf.left_indent = Pt(left_indent_pt)
    pf.space_before = Pt(before_pt)
    pf.space_after = Pt(after_pt)
    pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    pf.line_spacing = Pt(line_pt)


def _set_paragraph_bottom_border(paragraph, color: str = "000000", size: str = "6"):
    ppr = paragraph._p.get_or_add_pPr()
    existing = ppr.find(qn("w:pBdr"))
    if existing is not None:
        ppr.remove(existing)
    border = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), "3")
    bottom.set(qn("w:color"), color)
    border.append(bottom)
    ppr.append(border)


def _set_cell_shading(cell, fill: str):
    tcpr = cell._tc.get_or_add_tcPr()
    shd = tcpr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tcpr.append(shd)
    shd.set(qn("w:fill"), fill)


def _set_cell_margins(cell, top: int = 100, start: int = 140, bottom: int = 100, end: int = 140):
    tcpr = cell._tc.get_or_add_tcPr()
    margins = tcpr.find(qn("w:tcMar"))
    if margins is None:
        margins = OxmlElement("w:tcMar")
        tcpr.append(margins)
    for side, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = margins.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            margins.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def _set_table_borders(table, color: str = "8A8F98", size: str = "4"):
    tblpr = table._tbl.tblPr
    borders = tblpr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tblpr.append(borders)
    for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
        node = borders.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            borders.append(node)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), size)
        node.set(qn("w:color"), color)


def _add_page_field(paragraph):
    run = paragraph.add_run()
    _set_run_font(run, "仿宋_GB2312", 11)
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.text = " PAGE \\* MERGEFORMAT "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, separate, text, end])


def _configure_report_furniture(doc, title: str):
    section = doc.sections[0]
    header = section.header
    header.is_linked_to_previous = False
    hp = header.paragraphs[0]
    hp.text = ""
    hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _format_paragraph(hp, after_pt=0, line_pt=16)
    _set_run_font(hp.add_run("OSINT 战略研究报告 · DefenseTracker SOD/SOP"), "宋体", 9, color="5B6472")
    _set_paragraph_bottom_border(hp, "B7C0CC", "4")

    footer = section.footer
    footer.is_linked_to_previous = False
    fp = footer.paragraphs[0]
    fp.text = ""
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _format_paragraph(fp, after_pt=0, line_pt=14)
    _set_run_font(fp.add_run("— "), "仿宋_GB2312", 11)
    _add_page_field(fp)
    _set_run_font(fp.add_run(" —"), "仿宋_GB2312", 11)


def _configure_defensetracker_docx(doc):
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(3.7)
    section.bottom_margin = Cm(3.5)
    section.left_margin = Cm(2.8)
    section.right_margin = Cm(2.6)

    styles = doc.styles
    _set_style_font(styles["Normal"], "仿宋_GB2312", 16)
    styles["Normal"].paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    styles["Normal"].paragraph_format.line_spacing = Pt(28.5)
    styles["Normal"].paragraph_format.space_after = Pt(0)
    _set_style_font(styles["Title"], "方正小标宋简体", 26)
    _set_style_font(styles["Heading 1"], "黑体", 16, bold=True, color="0F172A")
    _set_style_font(styles["Heading 2"], "黑体", 16, bold=True, color="1F2937")
    _set_style_font(styles["Heading 3"], "楷体_GB2312", 16, bold=True, color="374151")


def _clean_inline_markdown(text: str) -> str:
    text = sanitize_report_text(text or "")
    text = re.sub(r"\[([^\]]+)\]\(([^)]+)\)", r"\1（\2）", text)
    text = re.sub(r"(\*\*|__|`)", "", text)
    return sanitize_report_text(text)


def _is_markdown_table_separator(line: str) -> bool:
    raw = line.strip()
    # 纯 "---" 是水平分割线而非表格分隔行：真正的表格分隔行必含 |（如 |---|---| 或 ---|---）。
    # 否则正文中"某系统 A|B 说明"后跟一行 --- 会被误判为单行表格，吞掉该行正文与水平线语义。
    if "|" not in raw:
        return False
    stripped = raw.strip("|").strip()
    return bool(stripped) and set(stripped) <= set("-:| ") and "-" in stripped


def _parse_markdown_table(lines: list[str]) -> list[list[str]]:
    rows = []
    for line in lines:
        if _is_markdown_table_separator(line):
            continue
        cells = [_clean_inline_markdown(c) for c in line.strip().strip("|").split("|")]
        if any(cells):
            rows.append(cells)
    width = max((len(row) for row in rows), default=0)
    return [row + [""] * (width - len(row)) for row in rows]


def _markdown_blocks(content: str) -> list[dict]:
    lines = content.splitlines()
    blocks: list[dict] = []
    i = 0
    while i < len(lines):
        line = lines[i].rstrip()
        stripped = line.strip()
        if not stripped:
            i += 1
            continue

        if "|" in stripped and i + 1 < len(lines) and _is_markdown_table_separator(lines[i + 1]):
            table_lines = [stripped, lines[i + 1].strip()]
            i += 2
            while i < len(lines) and "|" in lines[i].strip() and lines[i].strip():
                table_lines.append(lines[i].strip())
                i += 1
            rows = _parse_markdown_table(table_lines)
            if rows:
                blocks.append({"type": "table", "rows": rows})
            continue

        heading = re.match(r"^(#{1,6})\s+(.+)$", stripped)
        if heading:
            blocks.append({
                "type": "heading",
                "level": len(heading.group(1)),
                "text": _clean_inline_markdown(heading.group(2)),
            })
            i += 1
            continue

        list_item = re.match(r"^([-*+]|\d+[.)])\s+(.+)$", stripped)
        if list_item:
            items = []
            while i < len(lines):
                item_match = re.match(r"^([-*+]|\d+[.)])\s+(.+)$", lines[i].strip())
                if not item_match:
                    break
                items.append(_clean_inline_markdown(item_match.group(2)))
                i += 1
            blocks.append({"type": "list", "items": items})
            continue

        blocks.append({"type": "paragraph", "text": _clean_inline_markdown(stripped)})
        i += 1
    return blocks


def _same_title(a: str, b: str) -> bool:
    clean = lambda value: re.sub(r"[\s#《》“”\"'：:，,。.\-—]+", "", value or "")
    return bool(clean(a)) and clean(a) == clean(b)


def _heading_display_text(text: str) -> str:
    return "摘    要" if text.strip() == "摘要" else text.strip()


def _add_title_page(doc, project: dict, front_matter: dict | None = None):
    title = sanitize_report_text(project.get("title") or "防务战略分析报告")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(92)
    p.paragraph_format.space_after = Pt(10)
    run = p.add_run(title)
    _set_run_font(run, "方正小标宋简体", 26, color="0F172A")
    _set_paragraph_bottom_border(p, "0F172A", "12")

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_before = Pt(12)
    subtitle.paragraph_format.space_after = Pt(18)
    _set_run_font(subtitle.add_run("（军事理论研究报告）"), "仿宋_GB2312", 16, color="334155")

    tagline = doc.add_paragraph()
    tagline.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tagline.paragraph_format.space_after = Pt(34)
    _set_run_font(tagline.add_run("基于公开源证据链的技术—能力—规则三维综合研判"), "楷体_GB2312", 14, color="475569")

    meta_rows = [
        ("写作规范", "DefenseTracker SOD/SOP · 论文式战略研究版式"),
        ("分析框架", "技术—能力—规则 · PARA · FACT-DATA-CITE"),
        ("报告类型", REPORT_TYPE_DEFAULTS.get(project.get("report_type"), {}).get("label", "战略分析报告")),
        ("研究主题", project.get("topic") or title),
        ("生成时间", _format_chinese_date(datetime.now())),
    ]
    if front_matter:  # 报纸式 masthead：与屏上 newspaper 视图同款
        meta_rows.extend([
            ("期号", front_matter.get("issue") or ""),
            ("全文字数", f"约 {front_matter.get('word_count', 0)} 字"),
            ("预计阅读", f"约 {front_matter.get('reading_minutes', 1)} 分钟 · {front_matter.get('section_count', 0)} 节"),
        ])
    table = doc.add_table(rows=len(meta_rows), cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    _set_table_borders(table, "CBD5E1")
    for row_idx, (label, value) in enumerate(meta_rows):
        left, right = table.cell(row_idx, 0), table.cell(row_idx, 1)
        _set_cell_margins(left)
        _set_cell_margins(right)
        _set_cell_shading(left, "E2E8F0")
        if row_idx % 2:
            _set_cell_shading(right, "F8FAFC")
        left.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        right.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
        _set_run_font(left.paragraphs[0].add_run(label), "黑体", 11, bold=True, color="0F172A")
        _set_run_font(right.paragraphs[0].add_run(sanitize_report_text(str(value))), "仿宋_GB2312", 11, color="1F2937")

    if front_matter and front_matter.get("cards"):  # 报纸式：核心判断卡（与屏上同款）
        cards_head = doc.add_paragraph()
        cards_head.paragraph_format.space_before = Pt(28)
        cards_head.paragraph_format.space_after = Pt(6)
        _set_run_font(cards_head.add_run("核心判断卡"), "黑体", 13, bold=True, color="0F766E")
        for card in front_matter["cards"]:
            cp = doc.add_paragraph()
            cp.paragraph_format.space_after = Pt(4)
            _set_run_font(cp.add_run("• " + sanitize_report_text(str(card))), "仿宋_GB2312", 11, color="1F2937")

    note = doc.add_paragraph()
    note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    note.paragraph_format.space_before = Pt(34)
    _set_run_font(note.add_run("本报告仅基于公开来源资料生成；关键事实须以来源索引和证据附录闭环核验。"), "楷体_GB2312", 11, color="64748B")

    doc.add_page_break()


def _content_headings(blocks: list[dict], title: str) -> list[dict]:
    headings = []
    for block in blocks:
        if block.get("type") != "heading":
            continue
        if block.get("level") == 1 and _same_title(block.get("text", ""), title):
            continue
        headings.append(block)
    return headings


def _add_toc_page(doc, headings: list[dict]):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(10)
    run = p.add_run("目          录")
    _set_run_font(run, "黑体", 22, bold=True, color="0F172A")
    _set_paragraph_bottom_border(p, "0F172A", "8")

    hint = doc.add_paragraph()
    hint.alignment = WD_ALIGN_PARAGRAPH.CENTER
    hint.paragraph_format.space_after = Pt(12)
    _set_run_font(hint.add_run("按章节层级列示，便于审稿与快速定位。"), "楷体_GB2312", 10.5, color="64748B")

    for heading in headings:
        level = max(1, int(heading.get("level", 2)) - 1)
        item = doc.add_paragraph()
        _format_paragraph(item, left_indent_pt=(level - 1) * 22, after_pt=2, line_pt=24)
        run = item.add_run(_heading_display_text(heading.get("text", "")))
        _set_run_font(
            run,
            "黑体" if level == 1 else "仿宋_GB2312",
            14 if level == 1 else 12.5,
            bold=(level == 1),
            color="0F172A" if level == 1 else "334155",
        )

    doc.add_page_break()


def _render_report_heading(doc, text: str, markdown_level: int):
    display = _heading_display_text(text)
    level = 1 if markdown_level <= 2 else min(markdown_level - 1, 3)
    p = doc.add_paragraph(style=f"Heading {level}")
    align = WD_ALIGN_PARAGRAPH.CENTER if level == 1 else WD_ALIGN_PARAGRAPH.LEFT
    _format_paragraph(p, alignment=align, before_pt=14 if level == 1 else 9, after_pt=6, line_pt=30 if level == 1 else 28.5)
    font = "黑体" if level <= 2 else "楷体_GB2312"
    size = 16 if level == 1 else 15
    run = p.add_run(display)
    _set_run_font(run, font, size, bold=True, color="0F172A" if level == 1 else "1F2937")
    if level == 1:
        _set_paragraph_bottom_border(p, "CBD5E1", "6")


def _render_report_paragraph(doc, text: str):
    p = doc.add_paragraph()
    _format_paragraph(p, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY, first_line=True, after_pt=2)
    _set_run_font(p.add_run(text), "仿宋_GB2312", 16, color="111827")


def _render_report_list(doc, items: list[str]):
    for item in items:
        p = doc.add_paragraph()
        _format_paragraph(p, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY, left_indent_pt=32, after_pt=2)
        _set_run_font(p.add_run(f"●  {item}"), "仿宋_GB2312", 16, color="111827")


def _render_report_table(doc, rows: list[list[str]]):
    # 净化：剔除全空行、按最宽行补齐列数，避免畸形 markdown 表格导致 add_table(cols=0)/单元格越界崩溃
    rows = [list(r) for r in (rows or []) if any((c or "").strip() for c in r)]
    if not rows:
        return
    width = max((len(r) for r in rows), default=0)
    if width == 0:
        return
    rows = [r + [""] * (width - len(r)) for r in rows]
    caption = doc.add_paragraph()
    _format_paragraph(caption, alignment=WD_ALIGN_PARAGRAPH.CENTER, before_pt=6, after_pt=4, line_pt=20)
    _set_run_font(caption.add_run("表  证据/研判矩阵"), "楷体_GB2312", 10.5, color="64748B")
    table = doc.add_table(rows=len(rows), cols=width)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    _set_table_borders(table, "94A3B8")
    for r_idx, row in enumerate(rows):
        for c_idx, cell_text in enumerate(row):
            cell = table.cell(r_idx, c_idx)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            _set_cell_margins(cell, top=110, bottom=110, start=140, end=140)
            if r_idx == 0:
                _set_cell_shading(cell, "E2E8F0")
            elif r_idx % 2:
                _set_cell_shading(cell, "F8FAFC")
            para = cell.paragraphs[0]
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = para.add_run(cell_text)
            _set_run_font(run, "黑体" if r_idx == 0 else "仿宋_GB2312", 10.5, bold=(r_idx == 0), color="0F172A")
    doc.add_paragraph()


def _render_report_blocks(doc, blocks: list[dict], title: str):
    for block in blocks:
        btype = block.get("type")
        if btype == "heading":
            if block.get("level") == 1 and _same_title(block.get("text", ""), title):
                continue
            _render_report_heading(doc, block.get("text", ""), int(block.get("level", 2)))
        elif btype == "table":
            _render_report_table(doc, block.get("rows") or [])
        elif btype == "list":
            _render_report_list(doc, block.get("items") or [])
        elif btype == "paragraph" and block.get("text"):
            _render_report_paragraph(doc, block["text"])


def _add_source_index(doc, evidence: list[dict] | None):
    if not evidence:
        return
    doc.add_page_break()
    _render_report_heading(doc, "附录A：来源索引", 2)
    table = doc.add_table(rows=1, cols=5)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    _set_table_borders(table, "94A3B8")
    headers = ["序号", "来源", "资料标题", "质量", "链接"]
    for idx, header in enumerate(headers):
        cell = table.cell(0, idx)
        _set_cell_margins(cell, top=100, bottom=100, start=100, end=100)
        _set_cell_shading(cell, "E2E8F0")
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        _set_run_font(cell.paragraphs[0].add_run(header), "黑体", 9.5, bold=True, color="0F172A")
    for idx, ev in enumerate(evidence, 1):
        title = sanitize_report_text(ev.get("title") or "未命名资料")
        source = sanitize_report_text(ev.get("source_cn") or ev.get("source") or "公开来源")
        link = ev.get("link") or ""
        row = table.add_row().cells
        values = [
            str(idx),
            source,
            title,
            ev.get("quality_level") or str(ev.get("quality_score") or ""),
            link or "无",
        ]
        for c_idx, value in enumerate(values):
            cell = row[c_idx]
            _set_cell_margins(cell, top=90, bottom=90, start=90, end=90)
            if idx % 2 == 0:
                _set_cell_shading(cell, "F8FAFC")
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER if c_idx in (0, 3) else WD_ALIGN_PARAGRAPH.LEFT
            _set_run_font(cell.paragraphs[0].add_run(sanitize_report_text(str(value))), "仿宋_GB2312", 8.5, color="111827")


def build_report_docx(project: dict, draft: dict, evidence: list[dict] | None = None) -> BytesIO:
    if not DOCX_AVAILABLE:
        raise RuntimeError("python-docx 未安装")
    assert_report_exportable(draft, project)
    content = sanitize_report_text(draft.get("content") or "")
    if not content:
        raise ValueError("草稿内容为空")

    doc = Document()
    _configure_defensetracker_docx(doc)
    title = sanitize_report_text(project.get("title") or "防务战略分析报告")
    _configure_report_furniture(doc, title)
    blocks = _markdown_blocks(content)
    front_matter = build_newspaper_front_matter(project, content) if project.get("voice") == NEWSPAPER_VOICE else None
    _add_title_page(doc, project, front_matter)
    _add_toc_page(doc, _content_headings(blocks, title))
    _render_report_blocks(doc, blocks, title)
    _add_source_index(doc, evidence)

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf
